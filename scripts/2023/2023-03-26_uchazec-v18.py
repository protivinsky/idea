###  IMPORTS  ###
# region
import os
import sys
from venv import create
import pyreadstat
import pandas as pd
import numpy as np
import re
import requests
from urllib.request import urlopen
from io import StringIO
from statsmodels.stats.weightstats import DescrStatsW
from sklearn.linear_model import LinearRegression
import statsmodels.api as sm
import statsmodels.formula.api as smf
import scipy
import seaborn as sns
import matplotlib as mpl
import matplotlib.pyplot as plt

plt.rcParams['figure.figsize'] = 12, 6
import dbf
import json
import itertools

from docx import Document
from docx.shared import Mm, Pt

from libs.utils import *
from libs.plots import *
from libs.extensions import *
from libs import uchazec
from libs.maths import *
import reportree as rt

import importlib
import stata_setup

stata_setup.config('C:\\Program Files\\Stata17', 'mp')
stata = importlib.import_module('pystata.stata')
st = stata.run

import warnings
warnings.simplefilter(action='ignore', category=FutureWarning)

###  LOAD AND PREPARE DATA  ###
# region

generate_data = False
postprocess_data = False

if generate_data:

    for year in [17, 21]:
        print(f'loading 20{year}')

        # load and process data
        df, variable_labels, value_labels = uchazec.loader(year=year)

        # save data for stata -- this is apparently broken in windows - crashes on value labels...
        #   - and now it works fine?
        df.to_stata(f'{uchazec.data_root}/uchazec/uch{year}.dta', write_index=False, version=118,
                    variable_labels=variable_labels, value_labels=value_labels)

        # apply value labels and save for python
        for c in df.columns:
            if c in value_labels.keys():
                df[c] = df[c].map(value_labels[c]).astype('category')
        df.to_parquet(f'temp/uchazec/uch{year}.parquet')

if postprocess_data:
    print('Regenerating from cache.')
    df17 = pd.read_parquet('temp/uchazec/uch17.parquet')
    df21 = pd.read_parquet('temp/uchazec/uch21.parquet')

    df21 = uchazec.add_isced(df21)
    df17 = uchazec.add_isced(df17)

    isc17 = uchazec.get_per_isced(df17)
    isc21 = uchazec.get_per_isced(df21)

    df17, rep17 = uchazec.filter_data(df17)
    df21, rep21 = uchazec.filter_data(df21)

    df17 = uchazec.add_variables(df17)
    df21 = uchazec.add_variables(df21)
    df21 = uchazec.add_ss_typ_g(df21)
    df17 = uchazec.add_ss_typ_g(df17)

    pf17 = uchazec.get_per_app(df17)
    pf21 = uchazec.get_per_app(df21)

    po17 = uchazec.get_per_app(df17, ped_col='ped_obor')
    po21 = uchazec.get_per_app(df21, ped_col='ped_obor')

    pedf_vs = df21[df21['pedf']][['fak_nazev', 'vs_nazev']].drop_duplicates().sort_values('vs_nazev').reset_index(
        drop=False)

    print('Storing final datasets.')
    df17.to_parquet('temp/uchazec/final/uch17.parquet')
    df21.to_parquet('temp/uchazec/final/uch21.parquet')
    pf17.to_parquet('temp/uchazec/final/pf17.parquet')
    pf21.to_parquet('temp/uchazec/final/pf21.parquet')
    po17.to_parquet('temp/uchazec/final/po17.parquet')
    po21.to_parquet('temp/uchazec/final/po21.parquet')
    isc17.to_parquet('temp/uchazec/final/isc17.parquet')
    isc21.to_parquet('temp/uchazec/final/isc21.parquet')
    pedf_vs.to_parquet('temp/uchazec/final/pedf_vs.parquet')
    with open('temp/uchazec/final/rep17.json', 'w') as f:
        json.dump(rep17, f)
    with open('temp/uchazec/final/rep21.json', 'w') as f:
        json.dump(rep21, f)
    with open('temp/uchazec/final/var_labels.json', 'w') as f:
        json.dump(uchazec.variable_labels_all, f)

    print('Writing Stata datasets.')
    df17.to_stata(f'temp/uchazec/stata/uch17.dta', write_index=False, version=118,
                  variable_labels=uchazec.variable_labels_all)
    df21.to_stata(f'temp/uchazec/stata/uch21.dta', write_index=False, version=118,
                  variable_labels=uchazec.variable_labels_all)
    pf17.to_stata(f'temp/uchazec/stata/pf17.dta', write_index=False, version=118)
    pf21.to_stata(f'temp/uchazec/stata/pf21.dta', write_index=False, version=118)
    po17.to_stata(f'temp/uchazec/stata/po17.dta', write_index=False, version=118)
    po21.to_stata(f'temp/uchazec/stata/po21.dta', write_index=False, version=118)
    isc17.to_stata(f'temp/uchazec/stata/isc17.dta', write_index=False, version=118)
    isc21.to_stata(f'temp/uchazec/stata/isc21.dta', write_index=False, version=118)


def pedf_app(ff, c, ped_col='pedf'):
    foo = ff.groupby(c)[[f'prihl_{ped_col}_bool', 'prihl_bool']].sum().reset_index()
    foo[f'{ped_col}_rel'] = np.round(100 * foo[f'prihl_{ped_col}_bool'] / foo['prihl_bool'], 1)
    return foo


print('Reading final datasets.')
df17 = pd.read_parquet('temp/uchazec/final/uch17.parquet')
df21 = pd.read_parquet('temp/uchazec/final/uch21.parquet')
pf17 = pd.read_parquet('temp/uchazec/final/pf17.parquet')
pf21 = pd.read_parquet('temp/uchazec/final/pf21.parquet')
po17 = pd.read_parquet('temp/uchazec/final/po17.parquet')
po21 = pd.read_parquet('temp/uchazec/final/po21.parquet')
isc17 = pd.read_parquet('temp/uchazec/final/isc17.parquet')
isc21 = pd.read_parquet('temp/uchazec/final/isc21.parquet')
pedf_vs = pd.read_parquet('temp/uchazec/final/pedf_vs.parquet')
with open('temp/uchazec/final/rep17.json', 'r') as f:
    rep17 = json.load(f)
with open('temp/uchazec/final/rep21.json', 'r') as f:
    rep21 = json.load(f)

# df = df21
# df.head(n=1000).show()
# endregion

plt.rcParams['figure.figsize'] = 16, 9
# font_size()

figs = []
rps = {}

for y in [17, 21]:
    ff = eval(f'po{y}')
    ped_col = 'ped_obor'

    df = eval(f'df{y}')
    isc = eval(f'isc{y}').copy()

    foo = df[df['prijat']].copy()
    akivar = 'aki2'
    foo_id = foo[['id', akivar]].drop_duplicates()['id'].value_counts()
    foo_id_mult = foo_id[foo_id > 1].index

    foo = foo[foo['id'].isin(foo_id_mult)].copy()
    bar = foo.groupby(['id', akivar])['zaps_zapsal'].sum().reset_index()
    bar['pref'] = np.where(bar['zaps_zapsal'] > 0.5, 1, -1)

    aki2s = isc['aki2'].drop_duplicates().dropna().sort_values().values
    barpr = bar.pivot(index='id', columns=akivar, values='pref').fillna(0).astype(int)
    barpr['min'] = barpr.min(axis=1)
    barpr['max'] = barpr.max(axis=1)
    barpr = barpr[(barpr['min'] == -1) & (barpr['max'] == 1)].drop(columns=['min', 'max'])


    def comp_prefs(data, akis):
        assert data.shape[1] == len(akis)
        prefs = np.full([len(akis), len(akis)], 0)
        tots = np.full([len(akis), len(akis)], 0)

        def inner(x):
            idx = np.outer(x > 0, x < 0)
            prefs[idx] += 1
            tots[idx] += 1
            tots[idx.T] += 1

        np.apply_along_axis(inner, 1, data.values)
        return prefs, tots

    prefs, tots = comp_prefs(barpr, aki2s)
    prefs = pd.DataFrame(prefs, index=aki2s, columns=aki2s)
    tots = pd.DataFrame(tots, index=aki2s, columns=aki2s)

    aki2_dict = isc[['aki2', 'isced2']].drop_duplicates().set_index('aki2')['isced2'].to_dict()
    aki2_dict_short = {k: v if len(v) < 20 else v[:18] + '...' for k, v in aki2_dict.items()}

    rp = pd.merge(tots.stack().rename('tot').reset_index(), prefs.stack().rename('pref').reset_index())
    rp['ppct'] = 100 * rp['pref'] / rp['tot']
    rp['x'] = rp['level_1'].map(aki2_dict_short)
    rp['y'] = rp['level_0'].map(aki2_dict_short)
    rp = rp.dropna(subset='ppct')
    rp['sqrt_tot'] = np.sqrt(rp['tot'])

    isced2s = [aki2_dict_short[x] for x in aki2s]
    rp['x'] = pd.Categorical(rp['x'], categories=isced2s, ordered=True)
    rp['y'] = pd.Categorical(rp['y'], categories=isced2s, ordered=True)
    rps[y] = rp
    fig, ax = plt.subplots()
    sns.scatterplot(data=rp, x='x', y='y', hue='ppct', palette='RdYlBu', size='sqrt_tot', sizes=(5, 3000))
    for _, row in rp.iterrows():
        color, alpha = ('black', 0.8) if 25 < row['ppct'] < 75 else ('white', 1)
        ax.text(x=row['x'], y=row['y'], s=f'{row["ppct"]:.0f}', va='center', ha='center', alpha=alpha, color=color)
    ax.set(xlabel=None, ylabel=None, title='Preference při zápisu do širších oborových skupin')
    plt.xticks(rotation=30, ha='right')
    plt.legend([], [], frameon=False)
    fig.tight_layout()
    figs.append(fig)

prefs

# # prefs is ready for processing via bradley-terry
# n = 5
# p + p[:, None]

def bradley_terry(X, tol=1e-5, max_steps=1000):
    assert X.shape[0] == X.shape[1]
    if isinstance(X, pd.DataFrame):
        X = X.values
    n = X.shape[0]
    p = np.ones(shape=n, dtype=np.float_) / n
    error = np.inf
    i = 0
    W = np.sum(X, axis=1)
    # Parameters
    while error > tol and i < max_steps:
        p_new = W / np.sum((X + X.T) / (p + p[:, None]), axis=1)
        error = np.max(np.abs((p_new - p)))
        p = p_new
        i += 1
    if i == max_steps:
        print(f'Convergence was not reached, last error {error} after {i} steps.')
    # Variance
    P = p[:, None] / (p[:, None] + p)
    I = P * P.T - np.diag(np.diag(P))
    return p, np.linalg.inv(I)

p = bradley_terry(prefs)
{aki2_dict[f'{i + 1:02d}']: x for i, x in enumerate(p)}


def bradley_terry(X, tol=1e-6, max_iter=1000):
    X = np.asarray(X)  # Add this line to convert the input to a numpy array
    n = X.shape[0]
    p = np.ones(shape=n, dtype=np.float_) / n
    for i in range(max_iter):
        p_old = p.copy()
        logit_diffs = np.subtract.outer(p, p)
        P_ij = np.exp(logit_diffs) / (1 + np.exp(logit_diffs))
        ll = -np.sum(X * np.log(P_ij))
        print(f'Step {i}, log-likelihood {ll:.3g}, params {p}')
        H_ij = (X + X.T) * P_ij * P_ij.T - np.diag(np.sum((X + X.T) * P_ij * P_ij.T, axis=1))
        grad = np.sum(X - (X + X.T) * P_ij, axis=1)
        p = p - np.linalg.inv(H_ij) @ grad
        p = p / np.sum(p)

        if np.linalg.norm(p - p_old) < tol:
            break

    return p


def bradley_terry(X, tol=1e-6, max_iter=1000):
    X = np.asarray(X)  # Add this line to convert the input to a numpy array
    n = X.shape[0]
    p = np.zeros(shape=n, dtype=np.float_)
    for i in range(max_iter):
        p_old = p.copy()
        logit_diffs = np.subtract.outer(p, p)
        P_ij = np.exp(logit_diffs) / (1 + np.exp(logit_diffs))
        ll = -np.sum(X * np.log(P_ij))
        print(f'Step {i}, log-likelihood {ll:.3g}, params {p}')
        H_ij = (X + X.T) * P_ij * P_ij.T - np.diag(np.sum((X + X.T) * P_ij * P_ij.T, axis=1))
        grad = np.sum(X - (X + X.T) * P_ij, axis=1)
        p[1:] = p[1:] - np.linalg.inv(H_ij[1:, 1:]) @ grad[1:]

        if np.linalg.norm(p - p_old) < tol:
            break

    logit_diffs = np.subtract.outer(p, p)
    P_ij = np.exp(logit_diffs) / (1 + np.exp(logit_diffs))
    # H_ij = (X + X.T) * P_ij * P_ij.T - np.diag(np.sum((X + X.T) * P_ij * P_ij.T, axis=1))

    return p

np.linalg.inv(H_ij[])

p = bradley_terry(prefs, max_iter=10)
p = p / p.sum()
X = prefs
X

X = np.array([[0, 90], [10, 0]])
p = bradley_terry(X)
logit_diffs = np.subtract.outer(p, p)
P_ij = np.exp(logit_diffs) / (1 + np.exp(logit_diffs))

{aki2_dict[f'{i+1:02}']: f'{x:.3g}' for i, x in enumerate(p)}

np.sqrt(np.diag(- np.linalg.inv(H_ij[1:, 1:])))

def bradley_terry_gpt(X, max_iter=1000, tol=1e-6):
    n = X.shape[0]

    # Initialize logit-scale strengths (parameters) as zeros
    params = np.zeros(n)

    # Compute total comparisons per item
    comp_sums = X.sum(axis=1) + X.sum(axis=0)

    for _ in range(max_iter):
        params_old = params.copy()

        # Calculate P_ij matrix
        logit_diffs = np.subtract.outer(params, params)
        P_ij = np.exp(logit_diffs) / (1 + np.exp(logit_diffs))

        # Update parameters
        P_ij_row_sums = P_ij.sum(axis=1)
        params += np.log(X.sum(axis=1) / P_ij_row_sums) - np.log(comp_sums / P_ij.sum(axis=0))

        # Check for convergence
        if np.linalg.norm(params - params_old) < tol:
            break

    # # Calculate the Fisher information matrix
    # I = np.zeros((n, n))
    # np.fill_diagonal(I, (P_ij * (1 - P_ij)).sum(axis=1))
    # I -= P_ij * (1 - P_ij)
    #
    # # Calculate the inverse of the Fisher information matrix (variance-covariance matrix)
    # cov_matrix = np.linalg.inv(I)

    return params  #, cov_matrix

gpt_p = bradley_terry_gpt(prefs.values)
P = np.exp(gpt_p[:, None]) / (np.exp(gpt_p[:, None]) + np.exp(gpt_p))


np.subtract.outer(p, p)

X = prefs
prefs
aki2s
aki2_dict

rt.Leaf(figs, title='Revealed preferences', num_cols=1).show()

rp17 = rps[17]
rp21 = rps[21]
rp21.show()

cols = ['level_1', 'ppct', 'sqrt_tot']
rp17 = rp17[rp17['level_0'] == '01'][cols]
rp17['year'] = 'Rok 2017'
rp21 = rp21[rp21['level_0'] == '01'][cols]
rp21['year'] = 'Rok 2021'
rp = pd.concat([rp17, rp21])
rp['label'] = rp['level_1'].map(aki2_dict)
rp['label'] = pd.Categorical(rp['label'], categories=[aki2_dict[x] for x in aki2s], ordered=True)
rp['label_short'] = rp['level_1'].map(aki2_dict_short)
rp['label_short'] = pd.Categorical(rp['label_short'], categories=[aki2_dict_short[x] for x in aki2s], ordered=True)
rp['year'] = pd.Categorical(rp['year'])

# font_size_word()

plt.rcParams['figure.figsize'] = 10, 7
fig, ax = plt.subplots()
sns.scatterplot(data=rp, x='year', y='label', hue='ppct', palette='RdYlBu', size='sqrt_tot', sizes=(500, 2000))
for _, row in rp.iterrows():
    color, alpha = ('black', 0.8) if 24 < row['ppct'] < 76 else ('white', 1)
    ax.text(x=row['year'], y=row['label'], s=f'{row["ppct"]:.0f}', va='center', ha='center', alpha=alpha, color=color)
ax.set(xlabel=None, ylabel=None)
fig.suptitle('Podíl zapsaných do pedagogických oborů')
ax.set(xlim=(-1, 2))
#plt.xticks(rotation=30, ha='right')
plt.legend([], [], frameon=False)
fig.tight_layout()

fig.show()
rp.show()


###  SETUP THE ENVIRONMENT  ###
# region
# font sizes for importing to word
def font_size(small=9, medium=11, big='large'):
    plt.rc('font', size=small)  # controls default text sizes
    plt.rc('axes', titlesize=big)  # fontsize of the axes title
    plt.rc('axes', labelsize=medium)  # fontsize of the x and y labels
    plt.rc('xtick', labelsize=small)  # fontsize of the tick labels
    plt.rc('ytick', labelsize=small)  # fontsize of the tick labels
    plt.rc('legend', fontsize=small)  # legend fontsize
    plt.rc('figure', titlesize=big)  # fontsize of the figure title


def font_size_word():
    font_size(small=12, medium=14, big=18)


def savefig(ax_or_fig, title):
    fig_path = f'D:\\projects\\idea\\code\\output\\uchazec\\{title}.png'
    fig = ax_or_fig if isinstance(ax_or_fig, plt.Figure) else ax_or_fig.get_figure()
    fig.savefig(fig_path)
    return fig_path


font_size_word()
CHART_WIDTH = Mm(158)
# plt.style.use('dark_background')
# endregion

###  LOAD AND PREPARE DATA  ###
# region

generate_data = False
postprocess_data = False

if generate_data:

    for year in [17, 21]:
        print(f'loading 20{year}')

        # load and process data
        df, variable_labels, value_labels = uchazec.loader(year=year)

        # save data for stata -- this is apparently broken in windows - crashes on value labels...
        #   - and now it works fine?
        df.to_stata(f'{uchazec.data_root}/uchazec/uch{year}.dta', write_index=False, version=118,
                    variable_labels=variable_labels, value_labels=value_labels)

        # apply value labels and save for python
        for c in df.columns:
            if c in value_labels.keys():
                df[c] = df[c].map(value_labels[c]).astype('category')
        df.to_parquet(f'temp/uchazec/uch{year}.parquet')

if postprocess_data:
    print('Regenerating from cache.')
    df17 = pd.read_parquet('temp/uchazec/uch17.parquet')
    df21 = pd.read_parquet('temp/uchazec/uch21.parquet')

    df21 = uchazec.add_isced(df21)
    df17 = uchazec.add_isced(df17)

    isc17 = uchazec.get_per_isced(df17)
    isc21 = uchazec.get_per_isced(df21)

    df17, rep17 = uchazec.filter_data(df17)
    df21, rep21 = uchazec.filter_data(df21)

    df17 = uchazec.add_variables(df17)
    df21 = uchazec.add_variables(df21)
    df21 = uchazec.add_ss_typ_g(df21)
    df17 = uchazec.add_ss_typ_g(df17)

    pf17 = uchazec.get_per_app(df17)
    pf21 = uchazec.get_per_app(df21)

    po17 = uchazec.get_per_app(df17, ped_col='ped_obor')
    po21 = uchazec.get_per_app(df21, ped_col='ped_obor')

    pedf_vs = df21[df21['pedf']][['fak_nazev', 'vs_nazev']].drop_duplicates().sort_values('vs_nazev').reset_index(
        drop=False)

    print('Storing final datasets.')
    df17.to_parquet('temp/uchazec/final/uch17.parquet')
    df21.to_parquet('temp/uchazec/final/uch21.parquet')
    pf17.to_parquet('temp/uchazec/final/pf17.parquet')
    pf21.to_parquet('temp/uchazec/final/pf21.parquet')
    po17.to_parquet('temp/uchazec/final/po17.parquet')
    po21.to_parquet('temp/uchazec/final/po21.parquet')
    isc17.to_parquet('temp/uchazec/final/isc17.parquet')
    isc21.to_parquet('temp/uchazec/final/isc21.parquet')
    pedf_vs.to_parquet('temp/uchazec/final/pedf_vs.parquet')
    with open('temp/uchazec/final/rep17.json', 'w') as f:
        json.dump(rep17, f)
    with open('temp/uchazec/final/rep21.json', 'w') as f:
        json.dump(rep21, f)
    with open('temp/uchazec/final/var_labels.json', 'w') as f:
        json.dump(uchazec.variable_labels_all, f)

    print('Writing Stata datasets.')
    df17.to_stata(f'temp/uchazec/stata/uch17.dta', write_index=False, version=118,
                  variable_labels=uchazec.variable_labels_all)
    df21.to_stata(f'temp/uchazec/stata/uch21.dta', write_index=False, version=118,
                  variable_labels=uchazec.variable_labels_all)
    pf17.to_stata(f'temp/uchazec/stata/pf17.dta', write_index=False, version=118)
    pf21.to_stata(f'temp/uchazec/stata/pf21.dta', write_index=False, version=118)
    po17.to_stata(f'temp/uchazec/stata/po17.dta', write_index=False, version=118)
    po21.to_stata(f'temp/uchazec/stata/po21.dta', write_index=False, version=118)
    isc17.to_stata(f'temp/uchazec/stata/isc17.dta', write_index=False, version=118)
    isc21.to_stata(f'temp/uchazec/stata/isc21.dta', write_index=False, version=118)


def pedf_app(ff, c, ped_col='pedf'):
    foo = ff.groupby(c)[[f'prihl_{ped_col}_bool', 'prihl_bool']].sum().reset_index()
    foo[f'{ped_col}_rel'] = np.round(100 * foo[f'prihl_{ped_col}_bool'] / foo['prihl_bool'], 1)
    return foo


print('Reading final datasets.')
df17 = pd.read_parquet('temp/uchazec/final/uch17.parquet')
df21 = pd.read_parquet('temp/uchazec/final/uch21.parquet')
pf17 = pd.read_parquet('temp/uchazec/final/pf17.parquet')
pf21 = pd.read_parquet('temp/uchazec/final/pf21.parquet')
po17 = pd.read_parquet('temp/uchazec/final/po17.parquet')
po21 = pd.read_parquet('temp/uchazec/final/po21.parquet')
isc17 = pd.read_parquet('temp/uchazec/final/isc17.parquet')
isc21 = pd.read_parquet('temp/uchazec/final/isc21.parquet')
pedf_vs = pd.read_parquet('temp/uchazec/final/pedf_vs.parquet')
with open('temp/uchazec/final/rep17.json', 'r') as f:
    rep17 = json.load(f)
with open('temp/uchazec/final/rep21.json', 'r') as f:
    rep21 = json.load(f)

# df = df21
# df.head(n=1000).show()
# endregion

###  GENERATE REPORT  ###
# region

# EMPTY DOC
# region
doc = Document('template.docx')
doc._body.clear_content()

h = uchazec.HNumbering()

section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(25.4)
section.right_margin = Mm(25.4)
section.top_margin = Mm(25.4)
section.bottom_margin = Mm(25.4)

doc.add_heading('IDEA: Uchazeč', 0)

# table of contents field
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc.add_paragraph('')
paragraph = doc.add_paragraph()
run = paragraph.add_run()
fldChar = OxmlElement('w:fldChar')  # creates a new element
fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
instrText.text = 'TOC \\o "1-2" \\h \\z \\u'  # change 1-3 depending on heading levels you need

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:t')
fldChar3.text = "Right-click to update field."
fldChar2.append(fldChar3)

fldChar4 = OxmlElement('w:fldChar')
fldChar4.set(qn('w:fldCharType'), 'end')

r_element = run._r
r_element.append(fldChar)
r_element.append(instrText)
r_element.append(fldChar2)
r_element.append(fldChar4)
p_element = paragraph._p


def add_hyperlink(paragraph, url, text, color='0000FF', underline=True):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Add ? Remove underlining if it is requested
    if underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        # u.set(docx.oxml.shared.qn('w:val'), 'none')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


# endregion

# 1. PŘÍPRAVA DAT
# region
doc.add_heading(f'{h.h1} Příprava dat', 1)
m5_17 = np.sum(df17['typ_st'] == 'Magisterský') / df17.shape[0]
m5_21 = np.sum(df21['typ_st'] == 'Magisterský') / df21.shape[0]

doc.add_paragraph(
    f'Analýza je založena na anonymizovaných datech z databáze Uchazeč, která obsahuje informaci o všech přihláškách podaných všemi přihlášených na vysoké školy pro roky 2017 a 2021. Do analýzy jsou zahrnuti pouze žáci maturující v daném ročníku, jedná se tedy o přihlášky do bakalářských nebo (pětiletých) magisterských programů (tyto programy jsou stále časté především na lékařských, právnických a pedagogických fakultách, jedná o {m5_17 * 100:.1f} % v roce 2017 a o {m5_21 * 100:.1f} % v roce 2021), zahrnuti nejsou uchazeči o navazující magisterské programy.')

for y in [17, 21]:
    doc.add_heading(f'{h.h2} Data 20{y}', 2)
    for r in eval(f'rep{y}'):
        doc.add_paragraph(r, style='Bullet List')

doc.add_heading(f'{h.h2} Limity dat', 2)
doc.add_paragraph('V datech se vyskytují různé problémy, které komplikují analýzu:')
doc.add_paragraph('')
doc.add_paragraph(
    'Stejný student má často více přihlášek, které se v různých údajích (například o SŠ, odkud se hlásí) mohou lišit (týká se přibližně 1 % uchazečů). Tedy například jedna přihláška uvádí "SOŠ", druhá "Není SŠ", případně odlišná lokalita SŠ nebo chybějící údaje). Při agregaci na úroveň jednotlivých respondentů se snažíme identifikovat nejspolehlivější údaj (ignorujeme nevyplněné hodnoty; pokud na některé přihlášce byla uvedena informace o střední škole, uvažujeme tu). V některých ojedinělých případech nelze spolehlivě rozhodnout (například 4-leté a 8-leté gymnázium nebo různá informace o kraji), pak zvolíme pouze jednu z těchto hodnot.',
    style='Bullet List')
doc.add_paragraph('')
p = doc.add_paragraph(
    'Informace o výsledku přijímacího řízení je v datech z roku 2017 kódována mírně odlišným způsobem a některé hodnoty již ',
    style='Bullet List')
add_hyperlink(p,
              'http://stistko.uiv.cz/katalog/cslnk.asp?idc=MCPR&ciselnik=V%FDsledek+p%F8ij%EDmac%EDho+%F8%EDzen%ED&aap=on',
              'aktuální číselník')
p.add_run(' neobsahuje. Lze je dohledat ve ')
add_hyperlink(p,
              'http://stistko.uiv.cz/katalog/ciselnik11x.asp?idc=MCPR&ciselnik=V%FDsledek+p%F8ij%EDmac%EDho+%F8%EDzen%ED&aap=on&poznamka=',
              'starších souborech')
p.add_run(
    ', tím se opět komplikuje napojování dat, protože starší číselníky nejsou dostupné jako standardizovaný XML soubor.')
doc.add_paragraph('')
p = doc.add_paragraph(
    'Studijní programy a obory jsou v datech pro rok 2017 kódovány odlišným způsobem než v datech pro rok 2021, proto není možné uchazeče o pedagogické studium snadno identifikovat na úrovni oborů nebo programů. Je možné použít klasicikaci oborů podle CZ-ISCED-F 2013 – data pro rok 2021 již obsahují ISCED-F kód jako součást kódu studijního programu, pro starší data existuje převodník původních kódů programů a fakult na ISCED-F kód. Tento převod je však pouze přibližný, stejně tak i rozdělení oborů do studijních programů jednotlivých fakult: příkladem je uchazeč, který se hlasí na kombinaci oborů ',
    style='Bullet List')
r = p.add_run('Biologie se zaměřením na vzdělávání')
r.italic = True
p.add_run(' a ')
r = p.add_run('Geografie a kartografie se zaměřením na vzdělávání')
r.italic = True
p.add_run(' na Přírodovědecké fakultě Masarykovy univerzity. Fakulta tento program zařadila do studijního programu ')
r = p.add_run('Ekologická a evoluční biologie')
r.italic = True
p.add_run(', takže podle převodníku se jedná o ISCED-F obor ')
r = p.add_run('Biologie')
r.italic = True
p.add_run(
    '. Informace o oborech lze tedy také použít jako určité proxy pro uchazeče o studium pedagogiky, avšak zejména u starších dat pro rok 2017 je nutné být si vědom limitů tohoto přiřazení.')
# endregion

# 2. PŘIHLÁŠKY NA VYSOKÉ ŠKOLY
# region
doc.add_heading(f'{h.h1} Přihlášky na vysoké školy', 1)

for y in [17, 21]:
    ff = eval(f'pf{y}')
    tot_prihlasek = eval(f'df{y}').shape[0]
    tot_uchazecu = ff.shape[0]
    ff_chlapci = ff[ff['gender'] == 'Chlapec']
    ff_divky = ff[ff['gender'] == 'Dívka']
    tot_chlapcu = ff_chlapci.shape[0]
    tot_divek = ff_divky.shape[0]

    # celkové statistiky
    avg_prihlasek = ff['prihl'].mean()
    max_prihlasek = ff['prihl'].max()
    med_prihlasek = ff['prihl'].median()

    # chlapci
    avg_prihlasek_chlapci = ff_chlapci['prihl'].mean()
    max_prihlasek_chlapci = ff_chlapci['prihl'].max()
    med_prihlasek_chlapci = ff_chlapci['prihl'].median()

    # dívky
    avg_prihlasek_divky = ff_divky['prihl'].mean()
    max_prihlasek_divky = ff_divky['prihl'].max()
    med_prihlasek_divky = ff_divky['prihl'].median()

    doc.add_heading(f'{h.h2} Přihlášky 20{y}', 2)
    doc.add_paragraph(
        f'V roce 20{y} bylo podáno {tot_prihlasek:,} přihlášek na vysoké školy od {tot_uchazecu:,} českých uchazečů, z toho se hlásilo {tot_chlapcu:,} chlapců ({100 * tot_chlapcu / tot_uchazecu:.3g} %) a {tot_divek:,} dívek ({100 * tot_divek / tot_uchazecu:.3g} %). Průměrný uchazeč si podal {avg_prihlasek:.1f} přihlášek (chlapci {avg_prihlasek_chlapci:.1f}, dívky {avg_prihlasek_divky:.1f}), mediánový uchazeč {med_prihlasek:.0f} (chlapci {med_prihlasek_chlapci:.0f}, dívky {med_prihlasek_divky:.0f}), maximální počet podaných přihlášek byl {max_prihlasek:.0f} (chlapci {max_prihlasek_chlapci:.0f}, dívky {max_prihlasek_divky:.0f}). Četnost počtu přihlášek jednotlivých uchazečů ukazuje následující graf.')
    doc.add_paragraph('')

    # fig, ax = plt.subplots()
    # ax = sns.histplot(data=ff, y='prihl', discrete=True, binwidth=0.8, shrink=0.8, ec=None)
    # ax.set(xlabel='Četnost', ylabel='Počet přihlášek', title='Počet přihlášek jednotlivých uchazečů')
    # ax.set_ylim(0, 13)
    # ax.invert_yaxis()
    # fig.tight_layout()

    # v procentech, rozdělení chlapci a dívky
    gender_order = ['Chlapec', 'Dívka']
    gender_colors = ['tab:blue', 'tab:red']

    fig, ax = plt.subplots()
    ax = sns.histplot(data=ff, y='prihl', discrete=True, binwidth=0.8, shrink=0.8, ec=None, hue='gender',
                      stat='percent', multiple='stack', hue_order=gender_order, palette=gender_colors)
    ax.set(xlabel='Četnost', ylabel='Počet přihlášek', title='Počet přihlášek jednotlivých uchazečů')
    ax.set_ylim(0, 13)
    ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
    ax.get_legend().set_title('Gender')
    ax.invert_yaxis()
    fig.tight_layout()

    figpath = savefig(fig, f'cetnost_prihlasek_{y}')

    doc.add_picture(figpath, width=CHART_WIDTH)
    doc.add_paragraph('')

doc.add_paragraph(
    'Dívky si podávají více přihlášek oproti chlapcům. Další analýza také ukáže, že dívky si zpravidla vybírají vysoké školy s vyššími nároky u přijímacích zkoušek (tedy školy, kam se v průměru dostane menší počet uchazečů).')

doc.add_heading(f'{h.h2} Rozdíly mezi chlapci a dívkami', 2)

pr17 = df17.groupby('gender')[['prijat', 'w']].apply(nanaverage, weights='w')['prijat']
pr21 = df21.groupby('gender')[['prijat', 'w']].apply(nanaverage, weights='w')['prijat']

tot17 = po17['gender'].value_counts()
zps17 = po17[po17['zapis_bool']]['gender'].value_counts()

tot21 = po21['gender'].value_counts()
zps21 = po21[po21['zapis_bool']]['gender'].value_counts()

doc.add_paragraph(
    f'V roce 2017 dívky uspěly průměrně v {pr17["Dívka"]:.1%} přijímacích řízení, zatímco chlapci v {pr17["Chlapec"]:.1%} případů. Data pro rok 2021 jsou obdobná, dívky uspěly v {pr21["Dívka"]:.1%}, zatímco chlapci v {pr21["Chlapec"]:.1%} přijímacích řízení. Na vysoké školy se následně zapsalo {zps17["Dívka"]:,} dívek ({zps17["Dívka"] / tot17["Dívka"]:.1%} uchazeček) a {zps17["Chlapec"]:,} chlapců ({zps17["Chlapec"] / tot17["Chlapec"]:.1%} uchazečů) v roce 2017, {zps21["Dívka"]:,} dívek ({zps21["Dívka"] / tot21["Dívka"]:.1%} uchazeček) a {zps21["Chlapec"]:,} chlapců ({zps21["Chlapec"] / tot21["Chlapec"]:.1%} uchazečů) v roce 2021.')
doc.add_paragraph('')
doc.add_paragraph(
    'Ačkoli se chlapci a dívky výrazně liší v úspěšnosti u přijímacích řízení, většina tohoto rozdílu je způsobena odlišnými obory, na které se hlásí. Pokud při odhadu pravděpodobnosti přijetí zohledníme konkrétní obory (jde tedy o logistickou regresi s úspěšností jako závislou proměnnou a pohlavím a indikátory pro obory jako nezávislými proměnnými), pak jsou rozdíly v pravděpodobnosti přijetí mnohem menší: chlapci jsou přijati o 2.6 procentního bodu častěji oproti dívkám v roce 2017, o 1.5 procentního bodu častěji v roce 2021. Tento malý rozdíl může být způsobený tím, že dívek se o VŠ uchází výrazně více, jedná se tedy o širší skupinu v rámci celého dovednostního spektra.')

# # porovnání oborů logistickou regresí:
# stata.pdataframe_to_data(df17, force=True)
# # stata.pdataframe_to_data(df21, force=True)
#
# %%stata
# encode aki4, gen(aki4c)
# encode gender, gen(genderc)
# logit prijat genderc i.aki4c [pw=w]
#
# %stata logit prijat genderc [pw=w]
#
# logistic(-0.0580242)
# logistic(-0.102552)


# endregion

# 3. PROFIL UCHAZEČŮ O PEDAGOGICKÉ OBORY
# region
doc.add_heading(f'{h.h1} Profil uchazečů o pedagogické obory', 1)
ped_col = 'ped_obor'

p = doc.add_paragraph('Jako pedagogické obory zařazujeme ty obory, které podle klasifikace oborů ')
add_hyperlink(p, 'https://www.czso.cz/csu/czso/klasifikace-oboru-vzdelani-cz-isced-f-2013', 'CZ-ISCED-F 2013')
p.add_run(' patří do skupiny Vzdělávání a výchova, konkrétně se jedná o tyto obory:')

doc.add_paragraph('')

for _, row in isc21[isc21['aki2'] == '01'].sort_values('aki4').iterrows():
    doc.add_paragraph(row['isced4'], style='Bullet List')

for y in [17, 21]:
    ff = eval(f'po{y}').copy()

    tot_uchazecu = ff.shape[0]
    tot_chlapcu = (ff['gender'] == 'Chlapec').sum()
    tot_divek = (ff['gender'] == 'Dívka').sum()

    ped_uchazecu = ff[f'prihl_{ped_col}_bool'].sum()
    ped_chlapcu = ff[ff['gender'] == 'Chlapec'][f'prihl_{ped_col}_bool'].sum()
    ped_divek = ff[ff['gender'] == 'Dívka'][f'prihl_{ped_col}_bool'].sum()

    doc.add_heading(f'{h.h2} Uchazeči 20{y}', 2)

    doc.add_paragraph(
        f'V roce 20{y} se na pedagogické obory hlásilo {ped_uchazecu:,.0f} ({100 * ped_uchazecu / tot_uchazecu:.3g} %) uchazečů. Z toho bylo {ped_divek:,} dívek ({100 * ped_divek / tot_divek:.3g} % všech dívek) a {ped_chlapcu:,} chlapců ({100 * ped_chlapcu / tot_chlapcu:.3g} % všech chlapců). Z uchazečů na pedagogické obory bylo dívek {100 * ped_divek / ped_uchazecu:.3g} % (oproti podílu dívek {100 * tot_divek / tot_uchazecu:.3g} % mezi všemi uchazeči).')

    doc.add_paragraph(
        'Následující tabulka shrnuje, jaký podíl uchazečů se hlásil na pedagogické obory z jednotlivých krajů (podle lokality středních škol; procenta vyjadřují podíl ze všech uchazečů daného kraje).')
    doc.add_paragraph('')

    foo = pedf_app(ff, 'ss_kraj', ped_col=ped_col)
    table = doc.add_table(rows=1, cols=3, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kraj'
    hdr_cells[1].text = 'Uchazeči o ped. obory'
    hdr_cells[2].text = 'Podíl ped. oborů'

    for _, row in foo.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['ss_kraj']
        row_cells[1].text = f'{row[f"prihl_{ped_col}_bool"]:,}'
        row_cells[2].text = f'{row[f"{ped_col}_rel"]} %'

    doc.add_paragraph('')
    p = doc.add_paragraph('Další tabulka ukazuje podobné srovnání podle typu střední školy.')
    if y == 17:
        p.add_run(
            ' ("Není SŠ" zpravidla označuje uchazeče, které se nehlásí přímo ze střední školy, ale není tomu tak vždy – v některých případech byly takto označeni i aktuální maturanti, naopak u některých starších uchazečů byl typ střední školy vyplněný správně; malý počet uchazečů s jiným typem střední školy byl vynechán.)')
    if y == 21:
        p.add_run(
            ' Celkově se počet absolventů gymnázií hlásících se na vysoké školy změnil jen mírně, oproti tomu počet absolventů středních odborných škol, kteří se hlásí na vysoké školy, znatelně narostl. Nejvíce studentů se na pedagogické obory hlásí ze středních odborných škol, v relativním vyjádření je to pak největší podíl z absolventů čtyřletých gymnázií.')
    doc.add_paragraph('')

    foo = pedf_app(ff, 'ss_typ_g', ped_col=ped_col)
    table = doc.add_table(rows=1, cols=4, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Typ školy'
    hdr_cells[1].text = 'Uchazeči o ped. obory'
    hdr_cells[2].text = 'Uchazeči celkem'
    hdr_cells[3].text = 'Podíl ped. oborů'

    for _, row in foo[foo['ss_typ_g'] != 'Jiné'].iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['ss_typ_g']
        row_cells[1].text = f'{row[f"prihl_{ped_col}_bool"]:,}'
        row_cells[2].text = f'{row.prihl_bool:,}'
        row_cells[3].text = f'{row[f"{ped_col}_rel"]} %'

    doc.add_heading('Portfolia přihlášek uchazečů', 3)

    avg_prihl = ff[~ff[f'prihl_{ped_col}_bool']]['prihl'].mean()
    avg_prihl_uch = ff[ff[f'prihl_{ped_col}_bool']]['prihl'].mean()
    avg_prihl_uch_p = ff[ff[f'prihl_{ped_col}_bool']][f'prihl_{ped_col}'].mean()

    doc.add_paragraph(
        f'Uchazeči o pedagogické obory si podávají více přihlášek oproti ostatním uchazečům. Uchazeč, který se nehlásil na pedagogický obor, si podal průměrně {avg_prihl:.1f} přihlášek, oproti tomu uchazeč o pedagogický obor si podal průměrně {avg_prihl_uch:.1f} přihlášek (z toho průměrně {avg_prihl_uch_p:.1f} přihlášek na pedagogické obory). Následující graf shrnuje četnosti počtu přihlášek podaných pouze uchazeči o pedagogické obory, kde navíc barevný odstín naznačuje, jaký podíl přihlášek si podali na pedagogické, respektive nepedagogické obory (tedy pokud si podali pouze jednu přihlášku, pak se vždy jednalo o přihlášku na pedagogický obor). Přibližně polovina přihlášek těchto uchazečů směřovala na pedagogické obory, zároveň stále platí, že dívky si podávají více přihlášek oproti chlapcům.')
    doc.add_paragraph('')

    ff['count'] = 1
    foo = ff[ff[f'prihl_{ped_col}_bool']].groupby(['gender', 'prihl']).agg(
        {'count': 'sum', f'prihl_{ped_col}': 'mean', f'prihl_ne{ped_col}': 'mean'}).reset_index()
    tot_count = foo['count'].sum()
    foo['ratio'] = foo['count'] / tot_count
    foo[f'prihl_{ped_col}_pct'] = 100 * foo['ratio'] * foo[f'prihl_{ped_col}'] / foo['prihl']
    foo[f'prihl_ne{ped_col}_pct'] = 100 * foo['ratio'] * foo[f'prihl_ne{ped_col}'] / foo['prihl']

    bar = foo.set_index(['gender', 'prihl'])[[f'prihl_{ped_col}_pct', f'prihl_ne{ped_col}_pct']].stack().rename(
        'pct').reset_index()
    bar['label'] = bar.apply(
        lambda row: f'{row.gender}, {"pedagogický" if row.level_2 == f"prihl_{ped_col}_pct" else "nepedagogický"}',
        axis=1)
    bar['prihl'] = bar['prihl'].astype(int)
    bar = bar[bar['prihl'] <= 12]

    hue_order = ['Dívka, pedagogický', 'Dívka, nepedagogický', 'Chlapec, pedagogický', 'Chlapec, nepedagogický']
    hue_color = ['tab:red', 'tab:orange', 'tab:blue', 'tab:cyan']

    barcum = bar.set_index(['prihl', 'label'])['pct'].unstack()[hue_order].cumsum(axis=1).stack().rename(
        'pct').reset_index()

    fig, ax = plt.subplots()
    sns.barplot(data=barcum, x='pct', y='prihl', hue='label', orient='h', dodge=False, hue_order=hue_order[::-1],
                palette=hue_color[::-1], width=0.7, alpha=0.8)
    ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
    ax.set(title='Počet přihlášek uchazečů o pedagogické obory', xlabel='Relativní četnost', ylabel='Počet přihlášek')
    fig.tight_layout()

    figpath = savefig(fig, f'uchazeci_portfolia_cetnost_{y}')
    doc.add_picture(figpath, width=CHART_WIDTH)

    doc.add_paragraph('')
    doc.add_paragraph(
        'Další graf srovnává skupiny oborů, na které se hlásí uchazeči o pedagogické obory a ostatní uchazeči. X-ová osa znázorňuje relativní četnosti přihlášek do jednotlivých oborů (s výjimkou pedagogických oborů, které tvoří u pedagogických uchazečů přibližně polovinu všech jejich přihlášek). Uchazeči o pedagogické obory se relativně častěji hlásí na humanitní a společenskovědní obory, oproti tomu se výrazně méně hlásí na technické obory a informatiku.')
    doc.add_paragraph('')

    df = eval(f'df{y}')

    uchp = ff[ff[f'prihl_{ped_col}_bool']][['id']]
    uchnp = ff[~ff[f'prihl_{ped_col}_bool']][['id']]
    dfp = pd.merge(df, uchp)
    dfp = dfp[dfp['aki2'] != '01']
    dfnp = pd.merge(df, uchnp)

    frp = (100 * dfp['isced2'].value_counts() / dfp.shape[0]).rename(ped_col).reset_index()
    frnp = (100 * dfnp['isced2'].value_counts() / dfnp.shape[0]).rename(f'ne{ped_col}').reset_index()
    fr = pd.merge(frp, frnp, how='outer').fillna(0)
    frr = fr.set_index('index').unstack().rename('pct').reset_index()
    frr['label'] = np.where(frr['level_0'] == 'ped_obor', 'Pedagogičtí', 'Nepedagogičtí')

    fig, ax = plt.subplots()
    sns.barplot(data=frr, x='pct', y='index', hue='label')
    ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
    ax.get_legend().set_title('Uchazeči')
    ax.set(title='Rozdělení přihlášek mezi obory', xlabel='Relativní četnost', ylabel=None)
    fig.tight_layout()

    figpath = savefig(fig, f'uchazeci_portfolia_obory_{y}')
    doc.add_picture(figpath, width=CHART_WIDTH)

    # aplikacni ambice
    isc = eval(f'isc{y}').copy()
    isc['ap_ambice'] = isc['prijat'] / isc['total']

    df = pd.merge(df, isc[['aki4', 'ap_ambice']])

    uchp = ff[ff[f'prihl_{ped_col}_bool']][['id']]
    uchnp = ff[~ff[f'prihl_{ped_col}_bool']][['id']]
    dfp = pd.merge(df, uchp)
    dfp_p = dfp[dfp['aki2'] == '01']
    dfp_np = dfp[dfp['aki2'] != '01']
    dfnp = pd.merge(df, uchnp)

    aa_all = np.average(df['ap_ambice'], weights=df['w'])
    aa_p = np.average(dfp['ap_ambice'], weights=dfp['w'])
    aa_p_p = np.average(dfp_p['ap_ambice'], weights=dfp_p['w'])
    aa_p_np = np.average(dfp_np['ap_ambice'], weights=dfp_np['w'])
    aa_np = np.average(dfnp['ap_ambice'], weights=dfnp['w'])

    doc.add_paragraph('')
    doc.add_paragraph(
        f'Uchazeči o pedagogické obory se také častěji hlásí na obory, kam je náročnější se dostat. Souhrnně je přijato {aa_all:.1%} uchazečů, na pedagogických oborech je to pouze {aa_p_p:.1%} uchazečů. Průměrná míra přijetí na oborech, kam se hlásí nepedagogičtí uchazeči, je {aa_np:.1%}, oproti tomu průměrná míra na oborech, kam se hlásí pedagogičtí uchazeči, je {aa_p_np:.1%} (po vynechání pedagogických oborů). Rozdíl je způsobený především selektivitou uměleckých a některých společenskovědních oborů a také vysokou mírou přijetí na technické obory.')

    doc.add_paragraph('')
    doc.add_paragraph(
        'Portfolia uchazečů je možné analyzovat také pomocí podmíněných pravděpodobností, které udávají sílu vazby mezi dvěma oborovými skupinami: konkrétně hodnota v grafu v řádku I a sloupci J udává pravděpodobnost P[J | I], tedy pravděpodobnost, že se uchazeč hlásící se na obor I přihlásil i na obor J. Názvy oborových skupin jsou v grafu zkrácené, plné názvy širších oborových skupin jsou:')
    doc.add_paragraph('')

    isced2 = isc[['aki2', 'isced2']].drop_duplicates().set_index('aki2').sort_index().values[:, 0]
    for i in isced2:
        doc.add_paragraph(i, style='Bullet List')

    doc.add_paragraph('')
    doc.add_paragraph(
        'Z grafu je patrné, že uchazeči pedagogických oborů se dále hlásí na humanitní, umělecké a společenskovědní obory, do zdravotnictví a sociálních služeb a  případně také na administrativu nebo právo.')
    doc.add_paragraph('')

    aki2_dict = isc[['aki2', 'isced2']].drop_duplicates().set_index('aki2')['isced2'].to_dict()
    aki2_dict_short = {k: v if len(v) < 20 else v[:18] + '...' for k, v in aki2_dict.items()}
    aki2s = df['aki2'].drop_duplicates().dropna().sort_values().values
    apr = df[['id', 'aki2']].drop_duplicates()
    apr['one'] = 1
    apr = apr.pivot(index='id', columns='aki2', values='one')[aki2s].fillna(0)

    conf = pd.DataFrame(index=apr.columns, columns=apr.columns, dtype='float')

    # napočítej podmíněné pravděpodobnosti
    for i in aki2s:
        for j in aki2s:
            # conf[X -> Y] = # X & Y / # X
            conf.loc[i, j] = (apr[i] * apr[j]).sum() / apr[i].sum()

    conf.columns = conf.columns.map(aki2_dict_short)
    conf.index = conf.index.map(aki2_dict_short)

    mask = pd.DataFrame(np.eye(10, dtype='bool'), columns=conf.columns, index=conf.index)

    fig, ax = plt.subplots()
    sns.heatmap(100 * conf, annot=True, mask=mask, cmap='YlOrBr')
    ax.set(xlabel=None, ylabel=None, title='Podmíněné pravděpodobnosti mezi obory: Pr[řádky | sloupce]')
    fig.tight_layout()

    figpath = savefig(fig, f'uchazeci_portfolia_apriori_{y}')
    doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_heading(f'{h.h2} Srovnání vývoje', 2)
doc.add_paragraph(
    'Následující graf srovnává podíl uchazečů o pedagogické obory z jednotlivých krajů v letech 2017 a 2021. Je patrné, že podíl zájemců o pedagogické obory se v roce 2021 výrazně zvýšil téměř ve všech krajích.')
doc.add_paragraph('')

foo17 = pedf_app(po17, 'ss_kraj', ped_col=ped_col)[['ss_kraj', f'{ped_col}_rel']]
foo17['year'] = '2017'
foo21 = pedf_app(po21, 'ss_kraj', ped_col=ped_col)[['ss_kraj', f'{ped_col}_rel']]
foo21['year'] = '2021'
foo = pd.concat([foo21, foo17], ignore_index=True)

# fig, ax = plt.subplots()
# sns.scatterplot(data=foo, y='ss_kraj', x=f'{ped_col}_rel', hue='year', s=60)
# ax.set(xlabel='Podíl uchazečů', ylabel=None, title='Podíl uchazečů o pedagogické obory')
# plt.legend(title='Rok')
# plt.plot(foo21[f'{ped_col}_rel'], foo21.ss_kraj)
# plt.plot(foo17[f'{ped_col}_rel'], foo17.ss_kraj)
# ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
# fig.tight_layout()

foo = foo.sort_values(['year', 'ped_obor_rel'])

foo_u = foo.pivot(index='ss_kraj', columns='year', values='ped_obor_rel').sort_values('2017')
foo_u['max'] = foo_u[['2017', '2021']].max(axis=1)
foo_u['zmena'] = foo_u['2021'] - foo_u['2017']
foo_u = foo_u.reset_index()

fig, ax = plt.subplots()
sns.barplot(data=foo, y='ss_kraj', x=f'{ped_col}_rel', hue='year', hue_order=['2017', '2021'], width=0.7, alpha=0.9)
for idx, row in foo_u.iterrows():
    ax.text(x=row['max'] + 0.2, y=row.name, s=f'{row["zmena"]:+.1f} pb', va='center', alpha=0.8)
xmin, xmax = ax.get_xlim()
ax.set_xlim(xmin, xmax + 3)
ax.set(xlabel='Podíl uchazečů', ylabel=None, title='Podíl uchazečů o pedagogické obory')
plt.legend(title='Rok')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'uchazeci_obory_kraje_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_paragraph('')
doc.add_paragraph(
    'Další graf srovnává vývoj podílu uchazečů mezi lety 2017 a 2021 podle typu střední školy. Je patrné, že nejvíce zájemců o pedagogické obory přibylo mezi absolventy 4-letých gymnázií, případně také z víceletých gymnázií.')
doc.add_paragraph('')

foo17 = pedf_app(po17, 'ss_typ_g', ped_col=ped_col)[['ss_typ_g', f'{ped_col}_rel']]
foo17['year'] = '2017'
foo17 = foo17[foo17['ss_typ_g'] != 'Jiné']
foo21 = pedf_app(po21, 'ss_typ_g', ped_col=ped_col)[['ss_typ_g', f'{ped_col}_rel']]
foo21['year'] = '2021'
foo21 = foo21[foo21['ss_typ_g'] != 'Jiné']
foo = pd.concat([foo21, foo17], ignore_index=True)

foo_u = foo.pivot(index='ss_typ_g', columns='year', values='ped_obor_rel')
foo_u['max'] = foo_u[['2017', '2021']].max(axis=1)
foo_u['zmena'] = foo_u['2021'] - foo_u['2017']
foo_u = foo_u.reset_index()

fig, ax = plt.subplots()
sns.barplot(data=foo, y='ss_typ_g', x=f'{ped_col}_rel', hue='year', hue_order=['2017', '2021'], width=0.7, alpha=0.9)
for idx, row in foo_u.iterrows():
    ax.text(x=row['max'] + 0.2, y=row.name, s=f'{row["zmena"]:+.1f} pb', va='center', alpha=0.8)
xmin, xmax = ax.get_xlim()
ax.set_xlim(xmin, xmax + 3)
ax.set(xlabel='Podíl uchazečů', ylabel=None, title='Podíl uchazečů o pedagogické obory')
plt.legend(title='Rok', loc='lower right')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'uchazeci_obory_typ_ss_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_heading(f'{h.h2} Podíl z přihlášek', 2)

avg_prihl_17 = po17['prihl'].mean()
avg_prihl_ped_17 = po17[f'prihl_{ped_col}'].mean()
avg_prihl_neped_17 = po17[f'prihl_ne{ped_col}'].mean()
avg_prihl_21 = po21['prihl'].mean()
avg_prihl_ped_21 = po21[f'prihl_{ped_col}'].mean()
avg_prihl_neped_21 = po21[f'prihl_ne{ped_col}'].mean()

doc.add_paragraph(
    f'Data ukazují, že na pedagogické obory se v roce 2021 hlásí vyšší podíl uchazečů než v roce 2021. Zároveň si v roce 2021 průměrný uchazeč podává více přihlášek: {avg_prihl_21:.2g} oproti {avg_prihl_17:.2g} přihlášce v roce 2017. Z toho si v roce 2017 podal průměrný uchazeč {avg_prihl_ped_17:.2g} přihlášek na pedagogické obory a {avg_prihl_neped_17:.2g} na ostatní obory. V roce 2021 to bylo {avg_prihl_ped_21:.2g} přihlášek na pedagogické obory (nárůst o {100 * avg_prihl_ped_21 / avg_prihl_ped_17 - 100:.3g} %) a {avg_prihl_neped_21:.2g} na ostatní obory (mírný pokles, o {100 * avg_prihl_neped_21 / avg_prihl_neped_17 - 100:.3g} %).')

doc.add_paragraph(
    f'Pro ověření výsledků ukazují následující tabulky a grafy obdobná srovnání jako předchozí podíly podle krajů a typů střední školy, avšak vyjádřené jako podíl všech přihlášek, nikoli jako podíl uchazečů.')

df17.groupby('ss_kraj')[['ones', 'pedf']].sum().reset_index()

for y in [17, 21]:
    df = eval(f'df{y}')

    tot_prihlasek = df.shape[0]
    ped_prihlasek = df['ped_obor'].sum()

    doc.add_heading(f'20{y}:', 3)

    doc.add_paragraph(
        f'V roce 20{y} bylo na pedagogické obory podáno {ped_prihlasek:,.0f} z celkového počtu {tot_prihlasek:,.0f} přihlášek na vysoké školy od českých uchazečů (podíl pedagogických oborů tedy byl {100 * ped_prihlasek / tot_prihlasek:.3g} %). Následující tabulka shrnuje, jaký byl podíl přihlášek na pedagogické obory ze všech přihlášek podaných z daného kraje.')
    doc.add_paragraph('')

    foo = df.groupby('ss_kraj')[['ones', 'ped_obor']].sum().reset_index()
    foo['ped_obor_rel'] = 100 * foo['ped_obor'] / foo['ones']

    table = doc.add_table(rows=1, cols=3, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kraj'
    hdr_cells[1].text = 'Přihlášky na ped. obory'
    hdr_cells[2].text = 'Podíl ped. oborů'

    for _, row in foo.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['ss_kraj']
        row_cells[1].text = f'{row.ped_obor:,}'
        row_cells[2].text = f'{row.ped_obor_rel:.1f} %'

    doc.add_paragraph('')
    doc.add_paragraph('Další tabulka obdobně srovnává typy středních škol.')
    doc.add_paragraph('')

    foo = df.groupby('ss_typ_g')[['ones', 'ped_obor']].sum().reset_index()
    foo['ped_obor_rel'] = 100 * foo['ped_obor'] / foo['ones']

    table = doc.add_table(rows=1, cols=4, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Typ školy'
    hdr_cells[1].text = 'Přihlášky na ped. obory'
    hdr_cells[2].text = 'Přihlášky celkem'
    hdr_cells[3].text = 'Podíl ped. oborů'

    for _, row in foo[foo['ss_typ_g'] != 'Jiné'].iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['ss_typ_g']
        row_cells[1].text = f'{row.ped_obor:,}'
        row_cells[2].text = f'{row.ones:,}'
        row_cells[3].text = f'{row.ped_obor_rel:.1f} %'

doc.add_heading('Vývoj mezi lety 2017 a 2021', 3)

doc.add_paragraph(
    'Graf srovnávající vývoj podílu přihlášek na pedagogické obory v jednotlivých krajích potvrzuje, že vyšší zájem o pedagogické obory není daný jenom vyšším množstvím přihlášek. Podíl přihlášek na pedagogické obor se mezi lety 2017 a 2021 zvýšil téměř ve všech krajích. Přiřazení do oborů se však mezi lety 2017 a 2021 výrazně proměnilo, výsledky tak mohou být ovlivněny také odlišným kódováním uchazečů.')
doc.add_paragraph('')

foo17 = df17.groupby('ss_kraj')[['ones', 'ped_obor']].sum().reset_index()
foo17['ped_obor_rel'] = 100 * foo17['ped_obor'] / foo17['ones']
foo17['year'] = '2017'
foo21 = df21.groupby('ss_kraj')[['ones', 'ped_obor']].sum().reset_index()
foo21['ped_obor_rel'] = 100 * foo21['ped_obor'] / foo21['ones']
foo21['year'] = '2021'

foo = pd.concat([foo21, foo17], ignore_index=True)

foo['ss_kraj'] = foo['ss_kraj'].astype('str')  # zajímavé - minule to fungovalo bez toho, ne?
foo = foo.sort_values(['year', 'ped_obor_rel']).reset_index(drop=True)

foo_u = foo.pivot(index='ss_kraj', columns='year', values='ped_obor_rel').sort_values('2017')
foo_u['max'] = foo_u[['2017', '2021']].max(axis=1)
foo_u['zmena'] = foo_u['2021'] - foo_u['2017']
foo_u = foo_u.reset_index()

fig, ax = plt.subplots()
sns.barplot(data=foo, y='ss_kraj', x=f'{ped_col}_rel', hue='year', hue_order=['2017', '2021'], width=0.7, alpha=0.9)
for idx, row in foo_u.iterrows():
    ax.text(x=row['max'] + 0.2, y=row.name, s=f'{row["zmena"]:+.1f} pb', va='center', alpha=0.8)
xmin, xmax = ax.get_xlim()
ax.set_xlim(xmin, xmax + 3)
ax.set(xlabel='Podíl přihlášek', ylabel=None, title='Podíl přihlášek na pedagogické obory')
plt.legend(title='Rok')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'prihlasky_obory_kraje_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_paragraph('')
doc.add_paragraph(
    'Další graf srovnává vývoj podílu přihlášek na pedagogické obory mezi lety 2017 a 2021 podle typu střední školy. U absolventů SOŠ se podíl přihlášek na pedagogické obory zvýšil jen mírně, oproti tomu u absolventů gymnázií podíl přihlášek na pedagogické obory výrazně narostl (zejména u čtyřletých gymnázií), podíl se zvýšil také u uchazečů, kteří se nehlásí ze střední školy.')
doc.add_paragraph('')

foo17 = df17.groupby('ss_typ_g')[['ones', 'ped_obor']].sum().reset_index()
foo17['ped_obor_rel'] = 100 * foo17['ped_obor'] / foo17['ones']
foo17['year'] = '2017'
foo17 = foo17[foo17['ss_typ_g'] != 'Jiné']
foo21 = df21.groupby('ss_typ_g')[['ones', 'ped_obor']].sum().reset_index()
foo21['ped_obor_rel'] = 100 * foo21['ped_obor'] / foo21['ones']
foo21['year'] = '2021'
foo21 = foo21[foo21['ss_typ_g'] != 'Jiné']
foo = pd.concat([foo21, foo17], ignore_index=True)
foo['ss_typ_g'] = foo['ss_typ_g'].cat.remove_unused_categories()

foo_u = foo.pivot(index='ss_typ_g', columns='year', values='ped_obor_rel')
foo_u['max'] = foo_u[['2017', '2021']].max(axis=1)
foo_u['zmena'] = foo_u['2021'] - foo_u['2017']
foo_u = foo_u.reset_index()

fig, ax = plt.subplots()
sns.barplot(data=foo, y='ss_typ_g', x=f'{ped_col}_rel', hue='year', hue_order=['2017', '2021'], width=0.7, alpha=0.9)
for idx, row in foo_u.iterrows():
    ax.text(x=row['max'] + 0.2, y=row.name, s=f'{row["zmena"]:+.1f} pb', va='center', alpha=0.8)
xmin, xmax = ax.get_xlim()
ax.set_xlim(xmin, xmax + 3)
ax.set(xlabel='Podíl přihlášek', ylabel=None, title='Podíl přihlášek na pedagogické obory')
plt.legend(title='Rok', loc='lower right')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'prihlasky_obory_typ_ss_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

# endregion

# 4. ÚSPĚŠNOST UCHAZEČŮ O PEDAGOGICKÉ OBORY
# region
doc.add_heading(f'{h.h1} Úspěšnost uchazečů o pedagogické obory', 1)

for y in [17, 21]:
    ff = eval(f'po{y}')
    ped_col = 'ped_obor'

    avg_prijat = ff['prijat_bool'].mean()
    avg_prijat_divka = ff[ff['gender'] == 'Dívka']['prijat_bool'].mean()
    avg_prijat_chlapec = ff[ff['gender'] == 'Chlapec']['prijat_bool'].mean()

    ff = ff[ff[f'prihl_{ped_col}_bool']]
    avg_prijat_ped = ff[f'prijat_{ped_col}_bool'].mean()
    avg_prijat_ped_divka = ff[ff['gender'] == 'Dívka'][f'prijat_{ped_col}_bool'].mean()
    avg_prijat_ped_chlapec = ff[ff['gender'] == 'Chlapec'][f'prijat_{ped_col}_bool'].mean()
    tot_prijat_ped = ff[f'prijat_{ped_col}_bool'].sum()

    doc.add_heading(f'{h.h2} Rok 20{y}', 2)

    doc.add_paragraph(
        f'V roce 20{y} bylo přijato {100 * avg_prijat:.1f} % uchazečů, kteří se hlásili na vysoké školy (u dívek {100 * avg_prijat_divka:.1f} %, u chlapců {100 * avg_prijat_chlapec:.1f} % – rozdíl je pravděpodobně dán odlišným výběrem oborů).')
    doc.add_paragraph('')

    doc.add_paragraph(
        f'Na pedagogické obory bylo přijato {100 * avg_prijat_ped:.1f} % z uchazečů (u dívek {100 * avg_prijat_ped_divka:.1f} %, u chlapců {100 * avg_prijat_ped_chlapec:.1f} %), celkem se jednalo o {tot_prijat_ped:,.0f} osob.')
    doc.add_paragraph('')

    cols = [f'prihl_ne{ped_col}', f'prihl_{ped_col}', f'prijat_ne{ped_col}', f'prijat_{ped_col}', f'zapis_ne{ped_col}',
            f'zapis_{ped_col}']
    foo = (ff[cols] > 0).value_counts().rename('count').reset_index().query(
        f'prihl_ne{ped_col} & prihl_{ped_col} & prijat_ne{ped_col} & prijat_{ped_col}').drop(
        columns=[f'prihl_ne{ped_col}', f'prihl_{ped_col}', f'prijat_ne{ped_col}', f'prijat_{ped_col}']).reset_index(
        drop=True)
    ano_ne = {True: 'Ano', False: 'Ne'}
    foo[f'zapis_{ped_col}'] = foo[f'zapis_{ped_col}'].map(ano_ne)
    foo[f'zapis_ne{ped_col}'] = foo[f'zapis_ne{ped_col}'].map(ano_ne)

    tot = foo['count'].sum()

    doc.add_paragraph(
        f'Celkem {tot:,} uchazečů se hlásilo a zároveň bylo přijato na pedagogický obor i na jiný obor. U těchto uchazečů je možné sledovat, na který z oborů se skutečně zapsali, jak shrnuje následující tabulka.')
    doc.add_paragraph('')

    table = doc.add_table(rows=1, cols=4, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Zápis ped. obor'
    hdr_cells[1].text = 'Zápis jinde'
    hdr_cells[2].text = 'Počet'
    hdr_cells[3].text = 'Podíl'

    for _, row in foo.sort_values([f'zapis_{ped_col}', f'zapis_ne{ped_col}']).iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row[f'zapis_{ped_col}']
        row_cells[1].text = row[f'zapis_ne{ped_col}']
        row_cells[2].text = f'{row["count"]:,}'
        row_cells[3].text = f'{100 * row["count"] / tot:.1f} %'

    if y == 21:
        doc.add_paragraph('')
        doc.add_paragraph(
            'Data naznačují, že přijímací řízení na pedagogických oborech bylo v roce 2021 srovnatelně náročné jako v roce 2017. Také uchazeči, kteří byli zároveň přijatí na pedagogický a nepedagogický obor, si v obou letech vybírali pro zápis mírně častěji nepedagogický obor oproti pedagogickému.')

doc.add_heading(f'{h.h2} Úspěšnost podle krajů', 2)
ped_col = 'ped_obor'

# ff = ff[ff['prihl_pedf_bool']]
# avg_prijat_pedf = ff['prijat_pedf_bool'].mean()

foo17 = po17[po17[f'prihl_{ped_col}_bool']].copy().groupby('ss_kraj')[f'prijat_{ped_col}_bool'].mean().reset_index()
foo17['year'] = '2017'
foo17['prijat'] = 100 * foo17[f'prijat_{ped_col}_bool']
foo21 = po21[po21[f'prihl_{ped_col}_bool']].copy().groupby('ss_kraj')[f'prijat_{ped_col}_bool'].mean().reset_index()
foo21['year'] = '2021'
foo21['prijat'] = 100 * foo21[f'prijat_{ped_col}_bool']

foo = pd.concat([foo21, foo17], ignore_index=True)

fig, ax = plt.subplots()
sns.barplot(data=foo, y='ss_kraj', x='prijat', hue='year', hue_order=['2017', '2021'], width=0.7, alpha=0.9)
ax.set(xlabel='Podíl přijatých', ylabel=None, title='Podíl přijatých na pedagogické obory')
plt.legend(title='Rok')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'prijati_ped_obory_kraje_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_heading(f'{h.h2} Úspěšnost podle typu střední školy', 2)

foo17 = po17[po17[f'prihl_{ped_col}_bool']].copy().groupby('ss_typ_g')[f'prijat_{ped_col}_bool'].mean().reset_index()
foo17['year'] = '2017'
foo17 = foo17[foo17['ss_typ_g'] != 'Jiné']
foo17['prijat'] = 100 * foo17[f'prijat_{ped_col}_bool']
foo21 = po21[po21[f'prihl_{ped_col}_bool']].copy().groupby('ss_typ_g')[f'prijat_{ped_col}_bool'].mean().reset_index()
foo21['year'] = '2021'
foo21 = foo21[foo21['ss_typ_g'] != 'Jiné']
foo21['prijat'] = 100 * foo21[f'prijat_{ped_col}_bool']
foo = pd.concat([foo21, foo17], ignore_index=True)

fig, ax = plt.subplots()
sns.barplot(data=foo, y='ss_typ_g', x='prijat', hue='year', hue_order=['2017', '2021'], width=0.7, alpha=0.9)
ax.set(xlabel='Podíl přijatých', ylabel=None, title='Podíl přijatých na pedagogické obory')
plt.legend(title='Rok')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'prijati_ped_obory_typ_ss_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_heading(f'{h.h2} Náročnost přijetí podle oboru', 2)

foo17 = df17[df17[ped_col]].copy()
foo17 = foo17.groupby('aki4')['prijat'].mean().reset_index()
foo17['year'] = '2017'
foo17['prijat'] = 100 * foo17['prijat']

foo21 = df21[df21[ped_col]].copy()
foo21 = foo21.groupby('aki4')['prijat'].mean().reset_index()
foo21['year'] = '2021'
foo21['prijat'] = 100 * foo21['prijat']

foo = pd.concat([foo21, foo17], ignore_index=True)

isced4short = {
    '0111': 'Pedagogika',
    '0112': 'Předškolní vzdělávání a výchova',
    '0113': 'Bez předmětové specializace',
    '0114': 'S předmětovou specializací',
    '0188': 'Interdisciplinární programy',
}

foo['isced4short'] = foo['aki4'].map(isced4short)

fig, ax = plt.subplots()
sns.barplot(data=foo, y='isced4short', x='prijat', hue='year', hue_order=['2017', '2021'], width=0.7, alpha=0.9)
ax.set(xlabel='Podíl přijatých', ylabel=None, title='Podíl přijatých na pedagogické obory')
plt.legend(title='Rok')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'prijati_ped_obory_obor_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

# endregion

# 5. SROVNÁNÍ PEDAGOGICKÝCH OBORŮ
# region

doc.add_heading(f'{h.h1} Srovnání pedagogických oborů', 1)

p = doc.add_paragraph(
    'V roce 2021 jsou CZ-ISCED-F kódy součástí kódů studijních programů, proto přiřazení uchazečů do jednotlivých oborů je poměrně přímočaré. V roce 2017 byly používany starší kódy studijních programů a oborů, které sice lze převést na ISCED, avšak tento převod není spolehlivý kvůli nejednoznačnostem v zařazení oborů (například obor ')
r = p.add_run('Biologie se zaměřením na vzdělávání')
r.italic = True
p.add_run(
    ' je na některých fakultách pod odborným studijním programem a tedy není zařazen jako pedagogický obor). Z tohoto důvodu jsou uvedeny nejprve výsledky pro rok 2021, následně jako orientační také výsledky pro rok 2017.')
doc.add_paragraph('')

# pro popis grafu
isced4short = {
    '0111': 'Pedagogika',
    '0112': 'Předškolní vzdělávání a výchova',
    '0113': 'Bez předmětové specializace',
    '0114': 'S předmětovou specializací',
    '0188': 'Interdisciplinární programy',
}

sorter = {
    21: [3, 0, 4, 2, 1],
    17: [3, 0, 2, 1],
}

for y in [21, 17]:

    doc.add_heading(f'{h.h2} Srovnání oborů v roce 20{y}', 2)

    # procentuální zastoupení jednotlivých oborů
    df = eval(f'df{y}')
    dfpo = df[df['ped_obor']].copy()
    dfpo['isced4short'] = dfpo['aki4'].map(isced4short)

    tot = dfpo.shape[0]
    foo = (100 * dfpo['isced4short'].value_counts().sort_index().iloc[sorter[y]] / tot).reset_index()

    colors = sns.color_palette(n_colors=2)

    fig, ax = plt.subplots(figsize=(12, 4))
    sns.barplot(data=foo, y='index', x='isced4short', ax=ax, width=0.6, color=colors[0])
    for i, row in foo.iterrows():
        plt.text(row['isced4short'] + 0.5, i, f'{row.isced4short:.1f} %', va='center')

    ax.set(xlabel='Podíl uchazečů', ylabel='', title='Podíl uchazečů v jednotlivých pedagogických oborech')
    ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
    fig.tight_layout()

    figpath = savefig(fig, f'ped_obory_podil_{y}')

    p = doc.add_paragraph(
        f'V roce 20{y} se na pedagogické obory (dle klasifikace CZ-ISCED-F) hlásilo {tot:,} z celkového počtu {df.shape[0]:,} uchazečů ({100 * tot / df.shape[0]:.1f} %). Následující graf ukazuje podíl jednotlivých pedagogických oborů.')

    if y == 17:
        p.add_run(' V roce 2017 nebyl žádný z bakalářských studijních programů zařazený jako ')
        r = p.add_run('Příprava učitelů s předmětovou specializací')
        r.italic = True
        p.add_run(', proto tato kategorie není v grafech zobrazena.')

    doc.add_paragraph('')
    doc.add_picture(figpath, width=CHART_WIDTH)
    doc.add_paragraph('')

    doc.add_heading('Podíl dívek a chlapců', 3)
    doc.add_paragraph(
        f'Jednotlivé obory se výrazně liší v zastoupení dívek a chlapců, jak ukazují další dva grafy – první srovnává počty přihlášených chlapců a dívek, druhý potom počty uchazečů skutečně zapsaných do daných oborů.')

    foo = pd.crosstab(dfpo['isced4short'], dfpo['gender'], margins=True).reset_index().iloc[sorter[y]]

    fig, ax = plt.subplots(figsize=(12, 4))
    sns.barplot(data=foo, y='isced4short', x='All', label='Chlapci', ax=ax, color=colors[0], width=0.6)
    sns.barplot(data=foo, y='isced4short', x='Dívka', label='Dívky', ax=ax, color=colors[1], width=0.6)
    ax.legend()
    ax.set(xlabel='Počet uchazečů', ylabel='', title='Zastoupení dívek a chlapců v jednotlivých pedagogických oborech')
    fig.tight_layout()

    figpath = savefig(fig, f'ped_obory_gender_prijati_{y}')
    doc.add_paragraph('')
    doc.add_picture(figpath, width=CHART_WIDTH)

    foo = pd.crosstab(dfpo[dfpo['zaps_zapsal']]['isced4short'], dfpo[dfpo['zaps_zapsal']]['gender'],
                      margins=True).reset_index().iloc[sorter[y]]

    fig, ax = plt.subplots(figsize=(12, 4))
    sns.barplot(data=foo, y='isced4short', x='All', label='Chlapci', ax=ax, color=colors[0], width=0.6)
    sns.barplot(data=foo, y='isced4short', x='Dívka', label='Dívky', ax=ax, color=colors[1], width=0.6)
    ax.legend()
    ax.set(xlabel='Počet zapsaných', ylabel='', title='Počty zapsaných na pedagogické obory podle gender')
    fig.tight_layout()

    figpath = savefig(fig, f'ped_obory_gender_zapsani_{y}')
    doc.add_paragraph('')
    doc.add_picture(figpath, width=CHART_WIDTH)
    doc.add_paragraph('')

    doc.add_heading('Podíl přijatých a zapsaných uchazečů', 3)

    doc.add_paragraph(
        'Jednotlivé pedagogické obory se liší také ve své selektivitě, tedy v tom, kolik uchazečů je na jednotlivé obory přijato. Graf srovnává tuto selektivitu jednotlivých pedagogických oborů, další graf následně vyjadřuje, kolik z těchto přijatých uchazečů se na daný obor opravdu zapsalo.')

    foo_tot = dfpo['isced4short'].value_counts().sort_index().iloc[sorter[y]]
    foo_prij = dfpo[dfpo['prijat']]['isced4short'].value_counts().sort_index().iloc[sorter[y]]
    foo_zaps = dfpo[dfpo['zaps_zapsal']]['isced4short'].value_counts().sort_index().iloc[sorter[y]]

    foo = (100 * foo_prij / foo_tot).reset_index()

    fig, ax = plt.subplots(figsize=(12, 4))
    sns.barplot(data=foo, y='index', x='isced4short', ax=ax, width=0.6, color=colors[0])
    for i, row in foo.iterrows():
        plt.text(row['isced4short'] + 0.5, i, f'{row.isced4short:.1f} %', va='center')

    ax.set(xlabel='Podíl přijatých', ylabel='', title='Podíl přijatých v jednotlivých pedagogických oborech')
    ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
    fig.tight_layout()

    figpath = savefig(fig, f'ped_obory_podil_prijatych_{y}')
    doc.add_paragraph('')
    doc.add_picture(figpath, width=CHART_WIDTH)
    doc.add_paragraph('')

    foo = (100 * foo_zaps / foo_prij).reset_index()

    fig, ax = plt.subplots(figsize=(12, 4))
    sns.barplot(data=foo, y='index', x='isced4short', ax=ax, width=0.6, color=colors[0])
    for i, row in foo.iterrows():
        plt.text(row['isced4short'] + 0.5, i, f'{row.isced4short:.1f} %', va='center')

    ax.set(xlabel='Podíl zapsaných', ylabel='', title='Podíl zapsaných (z počtu přijatých) v jednotlivých oborech')
    ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
    fig.tight_layout()

    figpath = savefig(fig, f'ped_obory_podil_zapsanych_{y}')
    doc.add_paragraph('')
    doc.add_picture(figpath, width=CHART_WIDTH)
    doc.add_paragraph('')

doc.add_heading(f'{h.h2} Shrnutí', 2)
p = doc.add_paragraph(
    'Z grafů je jasně patrné odlišné kódování oborů v letech 2017 a 2021. Zároveň je však zřejmé velice nízké zastoupení chlapců v některých pedagogických oborech, podíl chlapců na oborech ')
r = p.add_run('Příprava učitelů bez předmětové specializace')
r.italic = True
p.add_run(' a ')
r = p.add_run('Pedagogika')
r.italic = True
p.add_run(' se pohybuje okolo 5 %, na oboru ')
r = p.add_run('Příprava učitelů pro předškolní vzdělávání a výchovu')
r.italic = True
p.add_run(
    ' je podíl chlapců přibližně 1 %. Nerovnosti mezi obory jsou jak v počtech uchazečů, tak v počtech skutečně zapsaných. Na nejvíce nevyvážený obor ')
r = p.add_run('Příprava učitelů pro předškolní vzdělávání a výchovu')
r.italic = True
p.add_run(
    ' se v roce 2017 hlásili dva chlapci a nezapsal se ani jeden, v roce 2021 se hlásilo 19 chlapců a zapsali se tři (obecně byl však počet uchazečů o tento obor v roce 2021 výrazně vyšší, což může souviset také s odlišným kódováním).')

doc.add_paragraph('')
p = doc.add_paragraph(
    'Jednotlivé pedagogické obory se liší také ve své selektivitě. Nejnižší podíl přijatých uchazečů, přibližně 20 % vykazuje obor ')
r = p.add_run('Příprava učitelů pro předškolní vzdělávání a výchovu')
r.italic = True
p.add_run(' (zároveň tento obor má také nejnižší podíl skutečně zapsaných přijatých uchazečů). Také na obor ')
r = p.add_run('Pedagogika')
r.italic = True
p.add_run(' je mírně obtížnější se dostat ve srovnání s ostatními pedagogickými obory.')

# endregion

# 6. REVEALED PREFERENCES
# region
doc.add_heading(f'{h.h1} Revealed Preferences', 1)

for y in [17, 21]:
    ff = eval(f'po{y}')
    ped_col = 'ped_obor'

    doc.add_heading(f'{h.h2} Projevené preference uchazečů v roce 20{y}', 2)

    doc.add_paragraph(
        'Pokud byl uchazeč přijatý na dva různé obory, lze sledovat, kam se nakonec skutečně zapsal. Porovnáním jednotlivých dvojic oborů tak můžeme usuzovat na skutečné preference uchazečů ohledně studijních oborů.')
    doc.add_paragraph('')

    if y == 17:
        doc.add_paragraph(
            'Následující graf uvažuje všechny uchazeče, kteří byli přijati na dva různé obory X a Y a následně se do jednoho z těchto oborů zapsali a do druhého nikoli. Hodnota v řádku X a sloupci Y udává podíl uchazečů (v procentech), kteří si vybrali pro zápis obor X; obdobně hodnota v řádku Y a sloupci X udává podíl uchazečů, kteří se zapsali na obor Y (tato preference mezi obory je kódována i barvou, kde modrá značí preferenci pro daný řádek, oproti tomu červená barva značí preferenci jiných oborů). Velikost bublin znázorňuje celkové množství uchazečů, kteří byli na danou kombinaci oborů přijati (velké bubliny jsou významnější, malé příliš ne).')
        doc.add_paragraph('')
        doc.add_paragraph(
            'Z grafu je tak například patrné, že uchazeči často upřednostňují oborovou skupinu "Obchod, administrativa a právo" (pravděpodobně kvůli právu), mnoho z nich dává přednost také zdravotním oborům před přírodovědnými nebo technickými. Uchazeči pedagogických oborů také dávají přednost skupině "Obchod, administrativa a právo", naopak před přírodními vědami volí spíše pedagogické obory.')
        doc.add_paragraph('')
    else:
        doc.add_paragraph(
            'Uchazeči v roce 2021 preferují zdravotnictví před jinými obory (vyšší preference oproti roku 2017 může být důsledkem epidemie COVID-19), naopak mají výrazně nižší zájem o služby. Uchazeči pedagogických oborů preferují zdravotní obory, pokud se na ně také dostanou, naopak dávají přednost pedagogickým oborům před službami nebo technickými obory. V některých ohledech se revealed preferences v roce 2021 liší od revealed preferences v roce 2017, příčinou rozdílů může být jak epidemie COVID-19, tak mírně odlišné kódování oborů v těchto letech.')
        doc.add_paragraph('')

    df = eval(f'df{y}')
    isc = eval(f'isc{y}').copy()

    foo = df[df['prijat']].copy()
    akivar = 'aki2'
    foo_id = foo[['id', akivar]].drop_duplicates()['id'].value_counts()
    foo_id_mult = foo_id[foo_id > 1].index

    foo = foo[foo['id'].isin(foo_id_mult)].copy()
    bar = foo.groupby(['id', akivar])['zaps_zapsal'].sum().reset_index()
    bar['pref'] = np.where(bar['zaps_zapsal'] > 0.5, 1, -1)

    aki2s = isc['aki2'].drop_duplicates().dropna().sort_values().values
    barpr = bar.pivot(index='id', columns=akivar, values='pref').fillna(0).astype(int)
    barpr['min'] = barpr.min(axis=1)
    barpr['max'] = barpr.max(axis=1)
    barpr = barpr[(barpr['min'] == -1) & (barpr['max'] == 1)].drop(columns=['min', 'max'])


    def comp_prefs(data, akis):
        assert data.shape[1] == len(akis)
        prefs = np.full([len(akis), len(akis)], 0)
        tots = np.full([len(akis), len(akis)], 0)

        def inner(x):
            idx = np.outer(x > 0, x < 0)
            prefs[idx] += 1
            tots[idx] += 1
            tots[idx.T] += 1

        np.apply_along_axis(inner, 1, data.values)
        return prefs, tots


    prefs, tots = comp_prefs(barpr, aki2s)
    prefs = pd.DataFrame(prefs, index=aki2s, columns=aki2s)
    tots = pd.DataFrame(tots, index=aki2s, columns=aki2s)

    aki2_dict = isc[['aki2', 'isced2']].drop_duplicates().set_index('aki2')['isced2'].to_dict()
    aki2_dict_short = {k: v if len(v) < 20 else v[:18] + '...' for k, v in aki2_dict.items()}

    mask = pd.DataFrame(np.eye(len(aki2s), dtype='bool'), columns=conf.columns, index=conf.index)

    rp = pd.merge(tots.stack().rename('tot').reset_index(), prefs.stack().rename('pref').reset_index())
    rp['ppct'] = 100 * rp['pref'] / rp['tot']
    rp['x'] = rp['level_1'].map(aki2_dict_short)
    rp['y'] = rp['level_0'].map(aki2_dict_short)
    rp = rp.dropna(subset='ppct')
    rp['sqrt_tot'] = np.sqrt(rp['tot'])

    isced2s = [aki2_dict_short[x] for x in aki2s]
    rp['x'] = pd.Categorical(rp['x'], categories=isced2s, ordered=True)
    rp['y'] = pd.Categorical(rp['y'], categories=isced2s, ordered=True)

    fig, ax = plt.subplots()
    sns.scatterplot(data=rp, x='x', y='y', hue='ppct', palette='RdYlBu', size='sqrt_tot', sizes=(5, 3000))
    for _, row in rp.iterrows():
        color, alpha = ('black', 0.8) if 25 < row['ppct'] < 75 else ('white', 1)
        ax.text(x=row['x'], y=row['y'], s=f'{row["ppct"]:.0f}', va='center', ha='center', alpha=alpha, color=color)
    ax.set(xlabel=None, ylabel=None, title='Preference při zápisu do širších oborových skupin')
    plt.xticks(rotation=30, ha='right')
    plt.legend([], [], frameon=False)
    fig.tight_layout()

    figpath = savefig(fig, f'uchazeci_revealed_preferences_{y}')
    doc.add_picture(figpath, width=CHART_WIDTH)

    if y == 17:
        doc.add_paragraph('')
        doc.add_paragraph(
            f'Širší oborové skupiny podle ISCED-F vymezují deset kategorií. Analogicky lze napočítat revealed preferences i na užších oborových skupinách (36 kategorií) nebo na oborech (83 kategorií), avšak graf se potom stává velice nepřehledným. Oba grafy zde také přikládám, znázorněny jsou pouze kombinace oborů, do kterých bylo přijato alespoň deset studentů (a grafy jsou čitelné jen při výrazném zvětšení dokumentu).')
        doc.add_paragraph('')
        doc.add_paragraph(
            f'Výsledky pro užší oborové skupiny potvrzují, že nejvíce studenti preferují právo a zdravotní obory{"" if y == 17 else ", případně také umělecké obory v roce 2021 (které však přijímají jen malé množství studentů)"}. Výsledky pro jednotlivé obory nabízí detailnější vhled do rozhodování studentů mezi blízkými obory. Výsledky poukazují:')
        doc.add_paragraph('')
        doc.add_paragraph(
            'Konkurence mezi pedaogickými obory a ostatními obory se týká především uchazečů o obor "Příprava učitelů s předmětovou specializací", v menší míře také uchazečů o obor "Pedagogika".',
            style='Bullet List')
        doc.add_paragraph(
            '"Humánní medicína" je jeden z nejvíce preferovaných oborů, pokud se je však zároveň student přijatý i na obor "Stomatologie", vybírá si zpravidla stomatologii.',
            style='Bullet List')
        doc.add_paragraph(
            'Pokud je student přijatý na architekturu i na stavebnictví, vybírá si téměř vždy architekturu pro zápis.',
            style='Bullet List')
        doc.add_paragraph(
            'Ze společenských věd jsou nejvíce preferované psychologie a žurnalistika; pokud se však student dostane na oba tyto obory, volí téměř vždy psychologii.',
            style='Bullet List')
        doc.add_paragraph(
            'V roce 2017 si studenti přijatí na právo i na medicínu vybírali častěji právo, v roce 2021 si častěji volili medicínu pro zápis.',
            style='Bullet List')
        doc.add_paragraph('')

    foo = df[df['prijat']].copy()
    akivar = 'aki3'
    foo_id = foo[['id', akivar]].drop_duplicates()['id'].value_counts()
    foo_id_mult = foo_id[foo_id > 1].index

    foo = foo[foo['id'].isin(foo_id_mult)].copy()
    bar = foo.groupby(['id', akivar])['zaps_zapsal'].sum().reset_index()
    bar['pref'] = np.where(bar['zaps_zapsal'] > 0.5, 1, -1)

    akis = isc[akivar].drop_duplicates().dropna().sort_values().values
    barpr = bar.pivot(index='id', columns=akivar, values='pref').fillna(0).astype(int)
    barpr['min'] = barpr.min(axis=1)
    barpr['max'] = barpr.max(axis=1)
    barpr = barpr[(barpr['min'] == -1) & (barpr['max'] == 1)].drop(columns=['min', 'max'])

    prefs, tots = comp_prefs(barpr, akis)

    prefs = pd.DataFrame(prefs, index=akis, columns=akis)
    tots = pd.DataFrame(tots, index=akis, columns=akis)

    aki_dict = isc[[akivar, 'isced' + akivar[-1]]].drop_duplicates().set_index(akivar)['isced' + akivar[-1]].to_dict()
    aki_dict_short = {k: v + ' [' + k + ']' if len(v) < 20 else v[:18] + '... [' + k + ']' for k, v in aki_dict.items()}

    rp = pd.merge(tots.stack().rename('tot').reset_index(), prefs.stack().rename('pref').reset_index())
    rp['ppct'] = 100 * rp['pref'] / rp['tot']
    rp['ppct_fltr'] = np.where(rp['tot'] < 10, np.nan, rp['ppct'])
    rp['x'] = rp['level_1'].map(aki_dict_short)
    rp['y'] = rp['level_0'].map(aki_dict_short)
    rp = rp.dropna(subset='ppct')
    rp['sqrt_tot'] = np.sqrt(rp['tot'])

    isceds = [aki_dict_short[x] for x in akis]
    rp['x'] = pd.Categorical(rp['x'], categories=isceds, ordered=True)
    rp['y'] = pd.Categorical(rp['y'], categories=isceds, ordered=True)

    fig, ax = plt.subplots(figsize=(30, 18))
    sns.scatterplot(data=rp, x='x', y='y', hue='ppct_fltr', palette='RdYlBu', size='sqrt_tot', sizes=(5, 3000))
    for _, row in rp.iterrows():
        if not pd.isna(row['ppct_fltr']):
            color, alpha = ('black', 0.8) if 25 < row['ppct'] < 75 else ('white', 1)
            ax.text(x=row['x'], y=row['y'], s=f'{row["ppct"]:.0f}', va='center', ha='center', alpha=alpha, color=color)
    ax.set(xlabel=None, ylabel=None, title='Preference při zápisu do užších oborových skupin')
    plt.xticks(rotation=30, ha='right')
    plt.legend([], [], frameon=False)
    fig.tight_layout()

    figpath = savefig(fig, f'uchazeci_revealed_preferences_isced3_{y}')
    doc.add_picture(figpath, width=CHART_WIDTH)

    foo = df[df['prijat']].copy()
    akivar = 'aki4'
    foo_id = foo[['id', akivar]].drop_duplicates()['id'].value_counts()
    foo_id_mult = foo_id[foo_id > 1].index

    foo = foo[foo['id'].isin(foo_id_mult)].copy()
    bar = foo.groupby(['id', akivar])['zaps_zapsal'].sum().reset_index()
    bar['pref'] = np.where(bar['zaps_zapsal'] > 0.5, 1, -1)

    akis = isc[akivar].drop_duplicates().dropna().sort_values().values

    barpr = bar.pivot(index='id', columns=akivar, values='pref').fillna(0).astype(int)
    barpr['min'] = barpr.min(axis=1)
    barpr['max'] = barpr.max(axis=1)
    barpr = barpr[(barpr['min'] == -1) & (barpr['max'] == 1)].drop(columns=['min', 'max'])

    akis = barpr.columns.values
    prefs, tots = comp_prefs(barpr, akis)

    prefs = pd.DataFrame(prefs, index=akis, columns=akis)
    tots = pd.DataFrame(tots, index=akis, columns=akis)

    valid = tots.index[tots.max() >= 10]
    prefs = prefs[valid].loc[valid]
    tots = tots[valid].loc[valid]
    akis = tots.columns.values

    aki_dict = isc[[akivar, 'isced' + akivar[-1]]].drop_duplicates().set_index(akivar)['isced' + akivar[-1]].to_dict()
    aki_dict_short = {k: v + ' [' + k + ']' if len(v) < 20 else v[:18] + '... [' + k + ']' for k, v in aki_dict.items()}

    rp = pd.merge(tots.stack().rename('tot').reset_index(), prefs.stack().rename('pref').reset_index())
    rp['ppct'] = 100 * rp['pref'] / rp['tot']
    rp['ppct_fltr'] = np.where(rp['tot'] < 10, np.nan, rp['ppct'])
    rp['x'] = rp['level_1'].map(aki_dict_short)
    rp['y'] = rp['level_0'].map(aki_dict_short)
    rp = rp.dropna(subset='ppct')
    rp['sqrt_tot'] = np.sqrt(rp['tot'])
    isceds = [aki_dict_short[x] for x in akis]
    rp['x'] = pd.Categorical(rp['x'], categories=isceds, ordered=True)
    rp['y'] = pd.Categorical(rp['y'], categories=isceds, ordered=True)

    fig, ax = plt.subplots(figsize=(30, 18))
    sns.scatterplot(data=rp, x='x', y='y', hue='ppct_fltr', palette='RdYlBu', size='sqrt_tot', sizes=(5, 3000))
    for _, row in rp.iterrows():
        if not pd.isna(row['ppct_fltr']):
            color, alpha = ('black', 0.8) if 25 < row['ppct'] < 75 else ('white', 1)
            ax.text(x=row['x'], y=row['y'], s=f'{row["ppct"]:.0f}', va='center', ha='center', alpha=alpha, color=color)
    ax.set(xlabel=None, ylabel=None, title='Preference při zápisu do jednotlivých oborů')
    plt.xticks(rotation=30, ha='right')
    plt.legend([], [], frameon=False)
    fig.tight_layout()

    figpath = savefig(fig, f'uchazeci_revealed_preferences_isced4_{y}')
    doc.add_picture(figpath, width=CHART_WIDTH)

# endregion

# 7. ANALÝZA STŘEDNÍCH ŠKOL
# region
doc.add_heading(f'{h.h1} Analýza středních škol', 1)

doc.add_paragraph(
    'Uchazeče o vysoké školy je možné analyzovat také optikou středních škol, ze kterých se hlásí. Data z databáze uchazeč za rok 2021 jsou napojena na údaje o počtu absolventů jednotlivých škol. Přibližně u sta škol je v databázi Uchazeč více žáků maturujících v daném roce než kolik bylo absolventů: malé rozdíly jsou pravděpodobně způsobené žáky, kteří se sice hlásili na vysoké školy, avšak nepovedlo se jim dokončit studium střední školy v daném roce. U malého množství škol jsou však diskrepance větší, mohou být způsobeny nespolehlivými údaji o absolventech některých škol (příkladem jsou školy s RED IZO 600007472, 600005666, 600005836).')
doc.add_paragraph('')

# zaci = pd.read_excel("../data/uchazec/absolventi_ss.xlsx")
ab_raw = pd.read_excel("../data/uchazec/absolventi_ss.xlsx", sheet_name="List3")
ab = ab_raw.groupby('red_izo')[['absolventi_celkem', 'absolventi_divky']].sum().reset_index().rename(
    columns={"red_izo": "ss_red_izo"})

ped_col = 'ped_obor'
vars_to_sum = ['prihl_bool', 'prijat_bool', 'zapis_bool', f'prihl_{ped_col}_bool', f'prijat_{ped_col}_bool',
               f'zapis_{ped_col}_bool']
foo = po21.groupby(['gender', 'ss_red_izo'])[vars_to_sum].sum()

foo = foo.unstack(level=0).fillna(0).reset_index()
sfx = {'': '', 'Chlapec': '_chlapec', 'Dívka': '_divka'}
foo.columns = [x + sfx[y] for x, y in foo.columns]

foo = pd.merge(foo, ab, how='left')
foo = foo.rename(columns={'absolventi_divky': 'absolventi_divka'})
foo['absolventi_chlapec'] = foo['absolventi_celkem'] - foo['absolventi_divka']

for x in vars_to_sum:
    foo[f'{x}_celkem'] = foo[f'{x}_divka'] + foo[f'{x}_chlapec']

isc21['aki4_podil_prijat'] = isc21['prijat'] / isc21['total']
aki4 = isc21[['aki4', 'aki4_podil_prijat']]
df21_ = pd.merge(df21, aki4, how='left')

ss_ambice = df21_.groupby('ss_red_izo')[['aki4_podil_prijat', 'w']].apply(nanaverage, weights='w').reset_index()
ss_ambice2 = df21_[~df21_['ped_obor']].groupby('ss_red_izo')[['aki4_podil_prijat', 'w']].apply(nanaverage,
                                                                                               weights='w').reset_index().rename(
    columns={'aki4_podil_prijat': 'aki4_podil_prijat_no_ped'})
ss21 = pd.merge(foo, ss_ambice)
ss21 = pd.merge(ss21, ss_ambice2, how='left')

ss_kraj = df21[['ss_red_izo', 'ss_kraj']].drop_duplicates().dropna()
ss21 = pd.merge(ss21, ss_kraj)

foo = po21[['ss_red_izo', 'ss_typ']].drop_duplicates().dropna()
foo = foo[~foo['ss_typ'].isin(['Není SŠ', 'Jiné'])]
bar = foo['ss_red_izo'].value_counts()
multiple = bar[bar > 1]
foo = foo[~foo.ss_red_izo.isin(multiple.index)]

ss21 = pd.merge(ss21, foo, how='left')
ss21['ss_typ'] = ss21['ss_typ'].fillna('Jiné')

# now I should be able to compute everything
ss21['podil_uchazecu'] = ss21['prihl_bool_celkem'] / ss21['absolventi_celkem']
ss21['podil_ped_uchazecu'] = ss21[f'prihl_{ped_col}_bool_celkem'] / ss21['absolventi_celkem']
ss21['podil_ped_uchazecu_z_uchazecu'] = ss21[f'prihl_{ped_col}_bool_celkem'] / ss21['prihl_bool_celkem']
ss21['podil_ped_uchazecu_z_uchazecu_disc'] = pd.cut(ss21['podil_ped_uchazecu_z_uchazecu'],
                                                    bins=[-0.001, 0.25, 0.5, 1.001])
ss21['podil_divek'] = ss21['absolventi_divka'] / ss21['absolventi_celkem']
ss21['podil_divek_disc'] = pd.cut(ss21['podil_divek'], bins=[-0.001, 0.3, 0.7, 1.001])

ss21['podil_uchazecu_pct'] = 100 * ss21['podil_uchazecu']
ss21['podil_ped_uchazecu_pct'] = 100 * ss21['podil_ped_uchazecu']
ss21['podil_ped_uchazecu_z_uchazecu_pct'] = 100 * ss21['podil_ped_uchazecu_z_uchazecu']
ss21['aki4_podil_prijat_pct'] = 100 * ss21['aki4_podil_prijat']
ss21['aki4_podil_prijat_no_ped_pct'] = 100 * ss21['aki4_podil_prijat_no_ped']

# pokud bych to chtěl zobrazit pomoci .show
to_drop = ['podil_ped_uchazecu_z_uchazecu_disc', 'podil_divek_disc']

# weird = ss21[ss21['podil_uchazecu'] > 1]
# weird.drop(columns=to_drop).show('weird')
# import warnings
# warnings.simplefilter(action='ignore', category=FutureWarning)

# priklady problematickych RED IZO skol:
# - 600007472: 1 absolvent, 0 zaku, ale v datech mame 16 uchazecu
# - 600005666: 1 absolvent, 171 zaku (to je trochu podezrele), mame 15 uchazecu
# - 600005836: 7 absolventu, 141 zaku, 43 uchazecu

# toto jsou nejkriklavejsi pripady, celkem je temer sto skol, kde mame vyssi pocet uchazecu nez absolventu. ackoli u vetsiny z nich jsou ty rozdily male, o 1-2 zaky

# zaci[zaci.red_izo == 600007472]
# zaci[zaci.red_izo == 600005666]
# ab_raw[ab_raw.red_izo == 600005666].T
# df21[df21.ss_red_izo == 600005666].show()
# ab_raw[ab_raw.p_izo == 10622110]
# ab_raw[ab_raw.red_izo == 600005836]
# zaci[zaci.red_izo == 600005836]

doc.add_paragraph(
    'Následující graf srovnává pro jednotlivé střední školy podíl absolventů, kteří se uchází o VŠ, a jejich aplikační ambice. Aplikační ambice jsou zde vyjádřené jako průměrná míra přijetí na oborech, kam se absolventi dané SŠ hlásí (nižší hodnota tedy znamená vyšší ambice absolventů). Velikost bublin odpovídá velikosti dané SŠ (resp. počtu absolventů v roce 2021), barva značí typ dané střední školy (kategorie "Jiné" zahrnuje jak školy, které nebyly zařazeny do kategorií SOŠ nebo Gymnázium, tak školy, které kombinují obě tyto kategorie dohromady). Z grafu je patrné, že na gymnáziích se o studium na VŠ uchází téměř všichni absolventi, oproti tomu zájem o studium na VŠ se mezi jednotlivými SOŠ velmi liší. Absolventi gymnázií jsou také ambicióznější v oborech, které si pro své studium vybírají. Na některých školách je počet uchazečů vyšší než počet absolventů: částečně toto lze vysvětlit tím, že některým uchazečům se nepodařilo dokončit střední školu, částečně to může být způsobeno i nepřesnými daty o absolventech. Střední školy, kde počet uchazečů převyšoval počet absolventů o více než o 20 %, jsou z grafu vynechány.')
doc.add_paragraph('')

fig, ax = plt.subplots()
sns.scatterplot(data=ss21[ss21['podil_uchazecu'] <= 1.2], x='podil_uchazecu_pct', y='aki4_podil_prijat_pct',
                hue='ss_typ', size='absolventi_celkem', sizes=(5, 200), alpha=0.8, ec='black')
handles, labels = ax.get_legend_handles_labels()
ax.legend(handles[1:4], labels[1:4], title='Typ střední školy')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.yaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.set(xlabel='Podíl absolventů ucházejících se o VŠ', ylabel='Aplikační ambice',
       title='Podíl uchazečů a jejich ambice podle typů škol')
fig.tight_layout()

figpath = savefig(fig, f'ss_podil_uchazecu_ambice')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_paragraph('')
doc.add_paragraph(
    'Další graf ukazuje, že školy se v zájmu o studium na VŠ také zásadně liší podle převažujícího genderu mezi studenty. Zatímco na většině gymnázií a některých středních školách je zastoupení chlapců a dívek relativně vyrovnané, graf ukazuje, že u středních odborných škol panují velké rozdíly: střední školy s převahou chlapců vykazují výrazně nižší aplikační ambice svých uchazečů (technické obory nejsou příliš selektivní), zároveň je u nich častěji nižší celkový zájem o studium na VŠ. Oproti tomu žáci ze škol s výraznou převahou dívek se častěji hlásí na VŠ a jsou ambicióznější ve svých přihláškách (často se jedna o zdravotnické nebo pedagogické střední školy).')
doc.add_paragraph('')

fig, ax = plt.subplots()
pal = ['royalblue', 'goldenrod', 'firebrick']
sns.scatterplot(data=ss21[ss21['podil_uchazecu'] <= 1.2], x='podil_uchazecu_pct', y='aki4_podil_prijat_pct',
                hue='podil_divek_disc', size='absolventi_celkem', sizes=(5, 200), alpha=0.8, ec='black', palette=pal)
handles, _ = ax.get_legend_handles_labels()
ax.legend(handles[1:4], [r'<30 %', r'30 % – 70 %', r'>70 %'], title='Podíl dívek')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.yaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.set(xlabel='Podíl absolventů ucházejících se o VŠ', ylabel='Aplikační ambice',
       title='Podíl uchazečů a jejich ambice podle zastoupení dívek')
fig.tight_layout()

figpath = savefig(fig, f'ss_podil_uchazecu_ambice_divky')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_paragraph('')
doc.add_paragraph(
    'Uchazeči o pedagogické obory pochází především ze škol, které mají relativně vysoký podíl zájemců o studium na VŠ mezi absolventy a vyšší aplikační ambice (což je dáno také tím, že pedagogické obory jsou poměrně selektivní). Na mnoha gymnáziích se uchází o pedagogický obor alespoň čtvrtina studentů, u některých středních odborných škol je to i více než polovina studentů (převážně pedagogické střední školy).')
doc.add_paragraph('')

fig, ax = plt.subplots()
pal = ['#fdae61', '#abd9e9', '#4575b4']
sns.scatterplot(data=ss21[ss21['podil_uchazecu'] <= 1.2], x='podil_uchazecu_pct', y='aki4_podil_prijat_pct',
                hue='podil_ped_uchazecu_z_uchazecu_disc', size='absolventi_celkem', sizes=(5, 200), alpha=1, ec='black',
                palette=pal)
handles, _ = ax.get_legend_handles_labels()
ax.legend(handles[1:4], [r'<25 %', r'25 % – 50 %', r'>50 %'], title='Podíl ped. uchazečů')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.yaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.set(xlabel='Podíl absolventů ucházejících se o VŠ', ylabel='Aplikační ambice',
       title='Podíl uchazečů, jejich ambice a podíl zájemců o pedagogické obory')
fig.tight_layout()

figpath = savefig(fig, f'ss_podil_uchazecu_ambice_ped')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_paragraph('')
doc.add_paragraph(
    'Rozdíly mezi středními školami ukazuje i následující graf, který dává do vztahu aplikační ambice studentů dané střední školy (opět vyjádřené jako průměrná míra přijetí na oborech, kam se studenti hlásí) a podíl uchazečů o pedagogické obory (z počtu uchazečů o studium na VŠ). Zatímco gymnázia tvoří střední část grafu (ambiciózní studenti ve svých přihláškách, střední množství zájemců o pedagogické obory), střední odborné školy se zde jasně rozpadají do tří rozdílných skupin:')
doc.add_paragraph(
    'Techničtěji zaměřené střední školy tvoří pravou část grafu: technické vysokoškolské obory zpravidla nejsou selektivní, o pedagogické obory na těchto školách není zájem.',
    style='Bullet List')
doc.add_paragraph(
    'Pedagogicky zaměřené střední školy tvoří horní část grafu: studenti zpravidla pokračují v pedagogickém zaměření i na VŠ, které jsou poměrně selektivní. Další graf srovnává, jak by se aplikační ambice SŠ změnily, kdyby při jejich výpočtu nebyly zahrnuty pedagogické obory.',
    style='Bullet List')
doc.add_paragraph(
    'Levou spodní část grafu tvoří střední školy, jejichž studenti se hlásí na velice selektivní obory a zároveň nemají zájem o učitelství. Jedná se o střední zdravotnické školy a umělecky zaměřené školy.')
doc.add_paragraph('')

fig, ax = plt.subplots()
sns.scatterplot(data=ss21, x='aki4_podil_prijat_pct', y='podil_ped_uchazecu_z_uchazecu_pct', hue='ss_typ',
                size='absolventi_celkem', sizes=(5, 200), alpha=0.8, ec='black')
handles, labels = ax.get_legend_handles_labels()
ax.legend(handles[1:4], labels[1:4], title='Typ střední školy')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.yaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.set(xlabel='Aplikační ambice', ylabel='Podíl uchazečů o pedagogické obory',
       title='Podíl uchazečů o pedagogické obory a aplikační ambice')
fig.tight_layout()

figpath = savefig(fig, f'ss_ambice_podil_ped_uchazecu_1')
doc.add_picture(figpath, width=CHART_WIDTH)

fig, ax = plt.subplots()
sns.scatterplot(data=ss21, x='aki4_podil_prijat_no_ped_pct', y='podil_ped_uchazecu_z_uchazecu_pct', hue='ss_typ',
                size='absolventi_celkem', sizes=(5, 200), alpha=0.8, ec='black')
handles, labels = ax.get_legend_handles_labels()
ax.legend(handles[1:4], labels[1:4], title='Typ střední školy')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.yaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.set(xlabel='Aplikační ambice (bez pedagogických oborů)', ylabel='Podíl uchazečů o pedagogické obory',
       title='Podíl uchazečů o pedagogické obory a aplikační ambice')
fig.tight_layout()

figpath = savefig(fig, f'ss_ambice_podil_ped_uchazecu_2')
doc.add_picture(figpath, width=CHART_WIDTH)

# endregion

# 8. PROFIL UCHAZEČŮ O PEDAGOGICKÉ FAKULTY
# region
doc.add_heading(f'{h.h1} Profil uchazečů o pedagogické fakulty', 1)

doc.add_paragraph('Uchazeči se mohli hlásit na některou z osmi pedagogických fakult v České republice:')
doc.add_paragraph('')

for _, row in pedf_vs.iterrows():
    doc.add_paragraph(row['vs_nazev'], style='Bullet List')

for y in [17, 21]:
    ff = eval(f'pf{y}')

    tot_uchazecu = ff.shape[0]
    tot_chlapcu = (ff['gender'] == 'Chlapec').sum()
    tot_divek = (ff['gender'] == 'Dívka').sum()

    pedf_uchazecu = ff['prihl_pedf_bool'].sum()
    pedf_chlapcu = ff[ff['gender'] == 'Chlapec']['prihl_pedf_bool'].sum()
    pedf_divek = ff[ff['gender'] == 'Dívka']['prihl_pedf_bool'].sum()

    doc.add_heading(f'{h.h2} Uchazeči 20{y}', 2)

    doc.add_paragraph(
        f'V roce 20{y} se na pedagogické fakulty hlásilo {pedf_uchazecu:,.0f} uchazečů ({100 * pedf_uchazecu / tot_uchazecu:.3g} %). Z toho bylo {pedf_divek:,} dívek ({100 * pedf_divek / tot_divek:.3g} % všech dívek) a {pedf_chlapcu:,} chlapců ({100 * pedf_chlapcu / tot_chlapcu:.3g} % všech chlapců). Následující tabulka shrnuje, jaký podíl uchazečů se hlásil na pedagogické fakulty z jednotlivých krajů (procenta vyjadřují podíl ze všech uchazečů daného kraje).')
    doc.add_paragraph('')

    foo = pedf_app(ff, 'ss_kraj')
    table = doc.add_table(rows=1, cols=3, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kraj'
    hdr_cells[1].text = 'Uchazeči o PedF'
    hdr_cells[2].text = 'Podíl PedF'

    for _, row in foo.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['ss_kraj']
        row_cells[1].text = f'{row.prihl_pedf_bool:,}'
        row_cells[2].text = f'{row.pedf_rel} %'

    doc.add_paragraph('')
    p = doc.add_paragraph('Další tabulka ukazuje podobné srovnání podle typu střední školy.')
    if y == 17:
        p.add_run(
            ' ("Není SŠ" zpravidla označuje uchazeče, které se nehlásí přímo ze střední školy, ale není tomu tak vždy – v některých případech byly takto označeni i aktuální maturanti, naopak u některých starších uchazečů byl typ střední školy vyplněný správně; malý počet uchazečů s jiným typem střední školy byl vynechán.)')
    if y == 21:
        p.add_run(
            ' Celkově se počet absolventů gymnázií hlásících se na vysoké školy změnil jen mírně, oproti tomu počet absolventů středních odborných škol, kteří se hlásí na vysoké školy, znatelně narostl. Nejvíce studentů se na pedagogické fakulty hlásí ze středních odborných škol, v relativním vyjádření je to pak největší podíl z absolventů čtyřletých gymnázií.')
    doc.add_paragraph('')

    foo = pedf_app(ff, 'ss_typ_g')
    table = doc.add_table(rows=1, cols=4, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Typ školy'
    hdr_cells[1].text = 'Uchazeči o PedF'
    hdr_cells[2].text = 'Uchazeči celkem'
    hdr_cells[3].text = 'Podíl PedF'

    for _, row in foo[foo['ss_typ_g'] != 'Jiné'].iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['ss_typ_g']
        row_cells[1].text = f'{row.prihl_pedf_bool:,}'
        row_cells[2].text = f'{row.prihl_bool:,}'
        row_cells[3].text = f'{row.pedf_rel} %'

doc.add_heading(f'{h.h2} Srovnání vývoje', 2)
doc.add_paragraph(
    'Následující graf srovnává podíl uchazečů o pedagogické fakulty z jednotlivých krajů v letech 2017 a 2021. Je patrné, že podíl zájemců o pedagogické fakulty se v roce 2021 výrazně zvýšil téměř ve všech krajích.')
doc.add_paragraph('')

ped_col = 'pedf'
foo17 = pedf_app(pf17, 'ss_kraj')[['ss_kraj', 'pedf_rel']]
foo17['year'] = '2017'
foo21 = pedf_app(pf21, 'ss_kraj')[['ss_kraj', 'pedf_rel']]
foo21['year'] = '2021'
foo = pd.concat([foo21, foo17], ignore_index=True)

foo = foo.sort_values(['year', 'pedf_rel'])

foo_u = foo.pivot(index='ss_kraj', columns='year', values='pedf_rel').sort_values('2017')
foo_u['max'] = foo_u[['2017', '2021']].max(axis=1)
foo_u['zmena'] = foo_u['2021'] - foo_u['2017']
foo_u = foo_u.reset_index()

fig, ax = plt.subplots()
sns.barplot(data=foo, y='ss_kraj', x=f'{ped_col}_rel', hue='year', hue_order=['2017', '2021'], width=0.7, alpha=0.9)
for idx, row in foo_u.iterrows():
    ax.text(x=row['max'] + 0.2, y=row.name, s=f'{row["zmena"]:+.1f} pb', va='center', alpha=0.8)
xmin, xmax = ax.get_xlim()
ax.set_xlim(xmin, xmax + 3)
ax.set(xlabel='Podíl uchazečů', ylabel=None, title='Podíl uchazečů o pedagogické fakulty')
plt.legend(title='Rok')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'uchazeci_kraje_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_paragraph('')
doc.add_paragraph(
    'Další graf srovnává vývoj podílu uchazečů mezi lety 2017 a 2021 podle typu střední školy. Je patrné, že nejvíce zájemců o pedagogické fakulty přibylo mezi absolventy 4-letých gymnázií, případně také z víceletých gymnázií.')
doc.add_paragraph('')

foo17 = pedf_app(pf17, 'ss_typ_g')[['ss_typ_g', 'pedf_rel']]
foo17['year'] = '2017'
foo17 = foo17[foo17['ss_typ_g'] != 'Jiné']
foo21 = pedf_app(pf21, 'ss_typ_g')[['ss_typ_g', 'pedf_rel']]
foo21['year'] = '2021'
foo21 = foo21[foo21['ss_typ_g'] != 'Jiné']
foo = pd.concat([foo21, foo17], ignore_index=True)

foo_u = foo.pivot(index='ss_typ_g', columns='year', values=f'{ped_col}_rel')
foo_u['max'] = foo_u[['2017', '2021']].max(axis=1)
foo_u['zmena'] = foo_u['2021'] - foo_u['2017']
foo_u = foo_u.reset_index()

fig, ax = plt.subplots()
sns.barplot(data=foo, y='ss_typ_g', x=f'{ped_col}_rel', hue='year', hue_order=['2017', '2021'], width=0.7, alpha=0.9)
for idx, row in foo_u.iterrows():
    ax.text(x=row['max'] + 0.2, y=row.name, s=f'{row["zmena"]:+.1f} pb', va='center', alpha=0.8)
xmin, xmax = ax.get_xlim()
ax.set_xlim(xmin, xmax + 3)
ax.set(xlabel='Podíl uchazečů', ylabel=None, title='Podíl uchazečů o pedagogické fakulty')
plt.legend(title='Rok', loc='lower right')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'uchazeci_typ_ss_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_heading(f'{h.h2} Podíl z přihlášek', 2)

avg_prihl_17 = pf17['prihl'].mean()
avg_prihl_pedf_17 = pf17['prihl_pedf'].mean()
avg_prihl_nepedf_17 = pf17['prihl_nepedf'].mean()
avg_prihl_21 = pf21['prihl'].mean()
avg_prihl_pedf_21 = pf21['prihl_pedf'].mean()
avg_prihl_nepedf_21 = pf21['prihl_nepedf'].mean()

doc.add_paragraph(
    f'Data ukazují, že na pedagogické fakulty se hlásí vyšší podíl uchazečů. Zároveň si v roce 2021 průměrný uchazeč podává více přihlášek: {avg_prihl_21:.2g} oproti {avg_prihl_17:.2g} přihlášce v roce 2017. Z toho si v roce 2017 podal průměrný uchazeč {avg_prihl_pedf_17:.2g} přihlášek na pedagogické fakulty a {avg_prihl_nepedf_17:.2g} na ostatní fakulty. V roce 2021 to bylo {avg_prihl_pedf_21:.2g} přihlášek na pedagogické fakulty (nárůst o {100 * avg_prihl_pedf_21 / avg_prihl_pedf_17 - 100:.3g} %) a {avg_prihl_nepedf_21:.2g} na ostatní fakulty (nárůst o {100 * avg_prihl_nepedf_21 / avg_prihl_nepedf_17 - 100:.3g} %).')

doc.add_paragraph(
    f'Pro ověření výsledků ukazují následující tabulky a grafy obdobná srovnání jako předchozí podíly podle krajů a typů střední školy, avšak vyjádřené jako podíl všech přihlášek, nikoli jako podíl uchazečů.')

df17.groupby('ss_kraj')[['ones', 'pedf']].sum().reset_index()

for y in [17, 21]:
    df = eval(f'df{y}')

    tot_prihlasek = df.shape[0]
    pedf_prihlasek = df['pedf'].sum()

    doc.add_heading(f'20{y}:', 3)

    doc.add_paragraph(
        f'V roce 20{y} bylo na pedagogické fakulty podáno {pedf_prihlasek:,.0f} z celkového počtu {tot_prihlasek:,.0f} přihlášek na vysoké školy od českých uchazečů (podíl pedagogických fakult tedy byl {100 * pedf_prihlasek / tot_prihlasek:.3g} %). Následující tabulka shrnuje, jaký byl podíl přihlášek na pedagogické fakulty ze všech přihlášek podaných z daného kraje.')
    doc.add_paragraph('')

    foo = df.groupby('ss_kraj')[['ones', 'pedf']].sum().reset_index()
    foo['pedf_rel'] = 100 * foo['pedf'] / foo['ones']

    table = doc.add_table(rows=1, cols=3, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kraj'
    hdr_cells[1].text = 'Přihlášky na PedF'
    hdr_cells[2].text = 'Podíl PedF'

    for _, row in foo.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['ss_kraj']
        row_cells[1].text = f'{row.pedf:,}'
        row_cells[2].text = f'{row.pedf_rel:.1f} %'

    doc.add_paragraph('')
    doc.add_paragraph('Další tabulka obdobně srovnává typy středních škol.')
    doc.add_paragraph('')

    foo = df.groupby('ss_typ_g')[['ones', 'pedf']].sum().reset_index()
    foo['pedf_rel'] = 100 * foo['pedf'] / foo['ones']

    table = doc.add_table(rows=1, cols=4, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Typ školy'
    hdr_cells[1].text = 'Přihlášky na PedF'
    hdr_cells[2].text = 'Přihlášky celkem'
    hdr_cells[3].text = 'Podíl PedF'

    for _, row in foo[foo['ss_typ_g'] != 'Jiné'].iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['ss_typ_g']
        row_cells[1].text = f'{row.pedf:,}'
        row_cells[2].text = f'{row.ones:,}'
        row_cells[3].text = f'{row.pedf_rel:.1f} %'

doc.add_heading('Vývoj mezi lety 2017 a 2021', 3)

doc.add_paragraph(
    'Graf srovnávající vývoj podílu přihlášek na pedagogické fakulty v jednotlivých krajích potvrzuje, že vyšší zájem o pedagogické fakulty není daný jenom vyšším množstvím přihlášek. Podíl přihlášek na pedagogické fakulty se mezi lety 2017 a 2021 zvýšil téměř ve všech krajích.')
doc.add_paragraph('')

foo17 = df17.groupby('ss_kraj')[['ones', 'pedf']].sum().reset_index()
foo17['pedf_rel'] = 100 * foo17['pedf'] / foo17['ones']
foo17['year'] = '2017'
foo21 = df21.groupby('ss_kraj')[['ones', 'pedf']].sum().reset_index()
foo21['pedf_rel'] = 100 * foo21['pedf'] / foo21['ones']
foo21['year'] = '2021'

foo = pd.concat([foo21, foo17], ignore_index=True)

foo['ss_kraj'] = foo['ss_kraj'].astype('str')  # zajímavé - minule to fungovalo bez toho, ne?
foo = foo.sort_values(['year', f'{ped_col}_rel']).reset_index(drop=True)

foo_u = foo.pivot(index='ss_kraj', columns='year', values=f'{ped_col}_rel').sort_values('2017')
foo_u['max'] = foo_u[['2017', '2021']].max(axis=1)
foo_u['zmena'] = foo_u['2021'] - foo_u['2017']
foo_u = foo_u.reset_index()

fig, ax = plt.subplots()
sns.barplot(data=foo, y='ss_kraj', x=f'{ped_col}_rel', hue='year', hue_order=['2017', '2021'], width=0.7, alpha=0.9)
for idx, row in foo_u.iterrows():
    ax.text(x=row['max'] + 0.2, y=row.name, s=f'{row["zmena"]:+.1f} pb', va='center', alpha=0.8)
xmin, xmax = ax.get_xlim()
ax.set_xlim(xmin, xmax + 3)
ax.set(xlabel='Podíl přihlášek', ylabel=None, title='Podíl přihlášek na pedagogické fakulty')
plt.legend(title='Rok')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'prihlasky_kraje_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_paragraph('')
doc.add_paragraph(
    'Další graf srovnává vývoj podílu přihlášek na pedagogické fakulty mezi lety 2017 a 2021 podle typu střední školy. U absolventů SOŠ se podíl přihlášek na pedagogické fakulty prakticky nezměnil, oproti tomu u absolventů gymnázií podíl přihlášek na pedagogické fakulty narostl (zejména u čtyřletých gymnázií), podíl se zvýšil také u uchazečů, kteří se nehlásí ze střední školy.')
doc.add_paragraph('')

foo17 = df17.groupby('ss_typ_g')[['ones', 'pedf']].sum().reset_index()
foo17['pedf_rel'] = 100 * foo17['pedf'] / foo17['ones']
foo17['year'] = '2017'
foo17 = foo17[foo17['ss_typ_g'] != 'Jiné']
foo21 = df21.groupby('ss_typ_g')[['ones', 'pedf']].sum().reset_index()
foo21['pedf_rel'] = 100 * foo21['pedf'] / foo21['ones']
foo21['year'] = '2021'
foo21 = foo21[foo21['ss_typ_g'] != 'Jiné']
foo = pd.concat([foo21, foo17], ignore_index=True)

foo['ss_typ_g'] = foo['ss_typ_g'].cat.remove_unused_categories()

foo_u = foo.pivot(index='ss_typ_g', columns='year', values=f'{ped_col}_rel')
foo_u['max'] = foo_u[['2017', '2021']].max(axis=1)
foo_u['zmena'] = foo_u['2021'] - foo_u['2017']
foo_u = foo_u.reset_index()

fig, ax = plt.subplots()
sns.barplot(data=foo, y='ss_typ_g', x=f'{ped_col}_rel', hue='year', hue_order=['2017', '2021'], width=0.7, alpha=0.9)
for idx, row in foo_u.iterrows():
    ax.text(x=row['max'] + 0.2, y=row.name, s=f'{row["zmena"]:+.1f} pb', va='center', alpha=0.8)
xmin, xmax = ax.get_xlim()
ax.set_xlim(xmin, xmax + 3)
ax.set(xlabel='Podíl přihlášek', ylabel=None, title='Podíl přihlášek na pedagogické fakulty')
plt.legend(title='Rok', loc='lower right')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()
figpath = savefig(fig, f'prihlasky_typ_ss_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

# endregion

# 9. ÚSPĚŠNOST UCHAZEČŮ O PEDAGOGICKÉ FAKULTY
# region
doc.add_heading(f'{h.h1} Úspěšnost uchazečů o pedagogické fakulty', 1)

for y in [17, 21]:
    ff = eval(f'pf{y}')

    avg_prijat = ff['prijat_bool'].mean()
    avg_prijat_divka = ff[ff['gender'] == 'Dívka']['prijat_bool'].mean()
    avg_prijat_chlapec = ff[ff['gender'] == 'Chlapec']['prijat_bool'].mean()

    ff = ff[ff['prihl_pedf_bool']]
    avg_prijat_pedf = ff['prijat_pedf_bool'].mean()
    avg_prijat_pedf_divka = ff[ff['gender'] == 'Dívka']['prijat_pedf_bool'].mean()
    avg_prijat_pedf_chlapec = ff[ff['gender'] == 'Chlapec']['prijat_pedf_bool'].mean()
    tot_prijat_pedf = ff['prijat_pedf_bool'].sum()

    doc.add_heading(f'{h.h2} Rok 20{y}', 2)

    doc.add_paragraph(
        f'V roce 20{y} bylo přijato {100 * avg_prijat:.1f} % uchazečů, kteří se hlásili na vysoké školy (u dívek {100 * avg_prijat_divka:.1f} %, u chlapců {100 * avg_prijat_chlapec:.1f} % – rozdíl je pravděpodobně dán odlišným výběrem oborů).')
    doc.add_paragraph('')

    doc.add_paragraph(
        f'Na pedagogické fakulty bylo přijato {100 * avg_prijat_pedf:.1f} % z uchazečů (u dívek {100 * avg_prijat_pedf_divka:.1f} %, u chlapců {100 * avg_prijat_pedf_chlapec:.1f} %), celkem se jednalo o {tot_prijat_pedf:,.0f} osob.')
    doc.add_paragraph('')

    cols = ['prihl_nepedf', 'prihl_pedf', 'prijat_nepedf', 'prijat_pedf', 'zapis_nepedf', 'zapis_pedf']
    foo = (ff[cols] > 0).value_counts().rename('count').reset_index().query(
        'prihl_nepedf & prihl_pedf & prijat_nepedf & prijat_pedf').drop(
        columns=['prihl_nepedf', 'prihl_pedf', 'prijat_nepedf', 'prijat_pedf']).reset_index(drop=True)
    ano_ne = {True: 'Ano', False: 'Ne'}
    foo['zapis_pedf'] = foo['zapis_pedf'].map(ano_ne)
    foo['zapis_nepedf'] = foo['zapis_nepedf'].map(ano_ne)

    tot = foo['count'].sum()

    doc.add_paragraph(
        f'Celkem {tot:,} uchazečů se hlásilo a zároveň bylo přijato na pedagogickou fakultu i na jinou fakultu. U těchto uchazečů je možné sledovat, na kterou z fakult se skutečně zapsali, jak shrnuje následující tabulka.')
    doc.add_paragraph('')

    table = doc.add_table(rows=1, cols=4, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Zápis PedF'
    hdr_cells[1].text = 'Zápis jinde'
    hdr_cells[2].text = 'Počet'
    hdr_cells[3].text = 'Podíl'

    for _, row in foo.sort_values(['zapis_pedf', 'zapis_nepedf']).iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['zapis_pedf']
        row_cells[1].text = row['zapis_nepedf']
        row_cells[2].text = f'{row["count"]:,}'
        row_cells[3].text = f'{100 * row["count"] / tot:.1f} %'

doc.add_paragraph('')
doc.add_paragraph(
    'Data naznačují, že přijímací řízení na pedagogické fakulty v roce 2021 bylo náročnější než v roce 2017. Zároveň uchazeči, kteří si při zápisu vybírali mezi pedagogickou a nepedagogickou fakultou, dávali v roce 2021 častěji přednost pedagogické fakultě, oproti tomu v roce 2017 dávali spíše přednost nepedagogické fakultě.')

doc.add_heading(f'{h.h2} Úspěšnost podle krajů', 2)

# ff = ff[ff['prihl_pedf_bool']]
# avg_prijat_pedf = ff['prijat_pedf_bool'].mean()

foo17 = pf17[pf17['prihl_pedf_bool']].copy().groupby('ss_kraj')['prijat_pedf_bool'].mean().reset_index()
foo17['year'] = '2017'
foo17['prijat'] = 100 * foo17['prijat_pedf_bool']
foo21 = pf21[pf21['prihl_pedf_bool']].copy().groupby('ss_kraj')['prijat_pedf_bool'].mean().reset_index()
foo21['year'] = '2021'
foo21['prijat'] = 100 * foo21['prijat_pedf_bool']

foo = pd.concat([foo21, foo17], ignore_index=True)

fig, ax = plt.subplots()
sns.barplot(data=foo, y='ss_kraj', x='prijat', hue='year', hue_order=['2017', '2021'], width=0.7, alpha=0.9)
ax.set(xlabel='Podíl přijatých', ylabel=None, title='Podíl přijatých na pedagogické fakulty')
plt.legend(title='Rok')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'prijati_pedf_kraje_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_heading(f'{h.h2} Úspěšnost podle typu střední školy', 2)

foo17 = pf17[pf17['prihl_pedf_bool']].copy().groupby('ss_typ_g')['prijat_pedf_bool'].mean().reset_index()
foo17['year'] = '2017'
foo17 = foo17[foo17['ss_typ_g'] != 'Jiné']
foo17['prijat'] = 100 * foo17['prijat_pedf_bool']
foo21 = pf21[pf21['prihl_pedf_bool']].copy().groupby('ss_typ_g')['prijat_pedf_bool'].mean().reset_index()
foo21['year'] = '2021'
foo21 = foo21[foo21['ss_typ_g'] != 'Jiné']
foo21['prijat'] = 100 * foo21['prijat_pedf_bool']
foo = pd.concat([foo21, foo17], ignore_index=True)

fig, ax = plt.subplots()
sns.barplot(data=foo, y='ss_typ_g', x='prijat', hue='year', hue_order=['2017', '2021'], width=0.7, alpha=0.9)
ax.set(xlabel='Podíl přijatých', ylabel=None, title='Podíl přijatých na pedagogické fakulty')
plt.legend(title='Rok')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'prijati_pedf_typ_ss_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_heading(f'{h.h2} Náročnost přijetí podle fakult', 2)

foo17 = df17[df17['pedf']].copy()
foo17['vs_nazev'] = foo17['vs_nazev'].cat.remove_unused_categories()
foo17 = foo17.groupby('vs_nazev')['prijat'].mean().reset_index()
foo17['year'] = '2017'
foo17['prijat'] = 100 * foo17['prijat']

foo21 = df21[df21['pedf']].copy()
foo21['vs_nazev'] = foo21['vs_nazev'].cat.remove_unused_categories()
foo21 = foo21.groupby('vs_nazev')['prijat'].mean().reset_index()
foo21['year'] = '2021'
foo21['prijat'] = 100 * foo21['prijat']

foo = pd.concat([foo21, foo17], ignore_index=True)

fig, ax = plt.subplots()
sns.barplot(data=foo, y='vs_nazev', x='prijat', hue='year', hue_order=['2017', '2021'], width=0.7, alpha=0.9)
ax.set(xlabel='Podíl přijatých', ylabel=None, title='Podíl přijatých na pedagogické fakulty')
plt.legend(title='Rok')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'prijati_pedf_vs_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

# endregion

doc.show()

# FINALIZACE

# genderové rozdíly mezi uchazeči
# region
# pro popis grafu
isced4short = {
    '0111': 'Pedagogika',
    '0112': 'Předškolní vzdělávání a výchova',
    '0113': 'Bez předmětové specializace',
    '0114': 'S předmětovou specializací',
    '0188': 'Interdisciplinární programy',
}

sorter = {
    21: [3, 0, 4, 2, 1],
    17: [3, 0, 2, 1],
}

y = 21

for y in [21, 17]:

    doc.add_heading(f'{h.h2} Srovnání oborů v roce 20{y}', 2)

    # procentuální zastoupení jednotlivých oborů
    df = eval(f'df{y}')
    dfpo = df[df['ped_obor']].copy()
    dfpo['isced4short'] = dfpo['aki4'].map(isced4short)

    tot = dfpo.shape[0]
    foo = (100 * dfpo['isced4short'].value_counts().sort_index().iloc[sorter[y]] / tot).reset_index()

    colors = sns.color_palette(n_colors=2)

    fig, ax = plt.subplots(figsize=(12, 4))
    sns.barplot(data=foo, y='index', x='isced4short', ax=ax, width=0.6, color=colors[0])
    for i, row in foo.iterrows():
        plt.text(row['isced4short'] + 0.5, i, f'{row.isced4short:.1f} %', va='center')

    ax.set(xlabel='Podíl uchazečů', ylabel='', title='Podíl uchazečů v jednotlivých pedagogických oborech')
    ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
    fig.tight_layout()

    fig.show()

    foo = pd.crosstab(dfpo['isced4short'], dfpo['gender'], margins=True).reset_index().iloc[sorter[y]]

    fig, ax = plt.subplots(figsize=(12, 4))
    sns.barplot(data=foo, y='isced4short', x='All', label='Chlapci', ax=ax, color=colors[0], width=0.6)
    sns.barplot(data=foo, y='isced4short', x='Dívka', label='Dívky', ax=ax, color=colors[1], width=0.6)
    ax.legend()
    ax.set(xlabel='Počet přihlášek', ylabel='')
    ax.get_xaxis().set_major_formatter(mpl.ticker.FuncFormatter(lambda x, p: format(int(x), ',')))
    fig.suptitle('Zastoupení dívek a chlapců mezi uchazeči o pedagogické obory')
    fig.tight_layout()
    fig.show()

    isc21.show()
    foo['Podíl'] = foo['Chlapec'] / foo['All']
    foo

    1099 / foo['Chlapec'].sum()

    foo = pd.crosstab(dfpo[dfpo['zaps_zapsal']]['isced4short'], dfpo[dfpo['zaps_zapsal']]['gender'],
                      margins=True).reset_index().iloc[sorter[y]]

    fig, ax = plt.subplots(figsize=(12, 4))
    sns.barplot(data=foo, y='isced4short', x='All', label='Chlapci', ax=ax, color=colors[0], width=0.6)
    sns.barplot(data=foo, y='isced4short', x='Dívka', label='Dívky', ax=ax, color=colors[1], width=0.6)
    ax.legend()
    ax.set(xlabel='Počet zapsaných', ylabel='')
    ax.get_xaxis().set_major_formatter(mpl.ticker.FuncFormatter(lambda x, p: format(int(x), ',')))
    fig.suptitle('Zastoupení chlapců a dívek mezi zapsanými studenty na pedagogické obory')
    fig.tight_layout()
    fig.show()

#     figpath = savefig(fig, f'ped_obory_gender_zapsani_{y}')
#     doc.add_paragraph('')
#     doc.add_picture(figpath, width=CHART_WIDTH)
#     doc.add_paragraph('')
#
#     doc.add_heading('Podíl přijatých a zapsaných uchazečů', 3)
#
#     doc.add_paragraph(
#         'Jednotlivé pedagogické obory se liší také ve své selektivitě, tedy v tom, kolik uchazečů je na jednotlivé obory přijato. Graf srovnává tuto selektivitu jednotlivých pedagogických oborů, další graf následně vyjadřuje, kolik z těchto přijatých uchazečů se na daný obor opravdu zapsalo.')
#
#     foo_tot = dfpo['isced4short'].value_counts().sort_index().iloc[sorter[y]]
#     foo_prij = dfpo[dfpo['prijat']]['isced4short'].value_counts().sort_index().iloc[sorter[y]]
#     foo_zaps = dfpo[dfpo['zaps_zapsal']]['isced4short'].value_counts().sort_index().iloc[sorter[y]]
#
#     foo = (100 * foo_prij / foo_tot).reset_index()
#
#     fig, ax = plt.subplots(figsize=(12, 4))
#     sns.barplot(data=foo, y='index', x='isced4short', ax=ax, width=0.6, color=colors[0])
#     for i, row in foo.iterrows():
#         plt.text(row['isced4short'] + 0.5, i, f'{row.isced4short:.1f} %', va='center')
#
#     ax.set(xlabel='Podíl přijatých', ylabel='', title='Podíl přijatých v jednotlivých pedagogických oborech')
#     ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
#     fig.tight_layout()
#
#     figpath = savefig(fig, f'ped_obory_podil_prijatych_{y}')
#     doc.add_paragraph('')
#     doc.add_picture(figpath, width=CHART_WIDTH)
#     doc.add_paragraph('')
#
#     foo = (100 * foo_zaps / foo_prij).reset_index()
#
#     fig, ax = plt.subplots(figsize=(12, 4))
#     sns.barplot(data=foo, y='index', x='isced4short', ax=ax, width=0.6, color=colors[0])
#     for i, row in foo.iterrows():
#         plt.text(row['isced4short'] + 0.5, i, f'{row.isced4short:.1f} %', va='center')
#
#     ax.set(xlabel='Podíl zapsaných', ylabel='', title='Podíl zapsaných (z počtu přijatých) v jednotlivých oborech')
#     ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
#     fig.tight_layout()
#
#     figpath = savefig(fig, f'ped_obory_podil_zapsanych_{y}')
#     doc.add_paragraph('')
#     doc.add_picture(figpath, width=CHART_WIDTH)
#     doc.add_paragraph('')
#
# doc.add_heading(f'{h.h2} Shrnutí', 2)
# p = doc.add_paragraph(
#     'Z grafů je jasně patrné odlišné kódování oborů v letech 2017 a 2021. Zároveň je však zřejmé velice nízké zastoupení chlapců v některých pedagogických oborech, podíl chlapců na oborech ')
# r = p.add_run('Příprava učitelů bez předmětové specializace')
# r.italic = True
# p.add_run(' a ')
# r = p.add_run('Pedagogika')
# r.italic = True
# p.add_run(' se pohybuje okolo 5 %, na oboru ')
# r = p.add_run('Příprava učitelů pro předškolní vzdělávání a výchovu')
# r.italic = True
# p.add_run(
#     ' je podíl chlapců přibližně 1 %. Nerovnosti mezi obory jsou jak v počtech uchazečů, tak v počtech skutečně zapsaných. Na nejvíce nevyvážený obor ')
# r = p.add_run('Příprava učitelů pro předškolní vzdělávání a výchovu')
# r.italic = True
# p.add_run(
#     ' se v roce 2017 hlásili dva chlapci a nezapsal se ani jeden, v roce 2021 se hlásilo 19 chlapců a zapsali se tři (obecně byl však počet uchazečů o tento obor v roce 2021 výrazně vyšší, což může souviset také s odlišným kódováním).')
#
# doc.add_paragraph('')
# p = doc.add_paragraph(
#     'Jednotlivé pedagogické obory se liší také ve své selektivitě. Nejnižší podíl přijatých uchazečů, přibližně 20 % vykazuje obor ')
# r = p.add_run('Příprava učitelů pro předškolní vzdělávání a výchovu')
# r.italic = True
# p.add_run(' (zároveň tento obor má také nejnižší podíl skutečně zapsaných přijatých uchazečů). Také na obor ')
# r = p.add_run('Pedagogika')
# r.italic = True
# p.add_run(' je mírně obtížnější se dostat ve srovnání s ostatními pedagogickými obory.')
# endregion


# revealed preferences - sloupec pro 2017 a pro 2021
# region

plt.rcParams['figure.figsize'] = 16, 9
font_size()

figs = []
rps = {}

for y in [17, 21]:
    ff = eval(f'po{y}')
    ped_col = 'ped_obor'

    df = eval(f'df{y}')
    isc = eval(f'isc{y}').copy()

    foo = df[df['prijat']].copy()
    akivar = 'aki2'
    foo_id = foo[['id', akivar]].drop_duplicates()['id'].value_counts()
    foo_id_mult = foo_id[foo_id > 1].index

    foo = foo[foo['id'].isin(foo_id_mult)].copy()
    bar = foo.groupby(['id', akivar])['zaps_zapsal'].sum().reset_index()
    bar['pref'] = np.where(bar['zaps_zapsal'] > 0.5, 1, -1)

    aki2s = isc['aki2'].drop_duplicates().dropna().sort_values().values
    barpr = bar.pivot(index='id', columns=akivar, values='pref').fillna(0).astype(int)
    barpr['min'] = barpr.min(axis=1)
    barpr['max'] = barpr.max(axis=1)
    barpr = barpr[(barpr['min'] == -1) & (barpr['max'] == 1)].drop(columns=['min', 'max'])


    def comp_prefs(data, akis):
        assert data.shape[1] == len(akis)
        prefs = np.full([len(akis), len(akis)], 0)
        tots = np.full([len(akis), len(akis)], 0)

        def inner(x):
            idx = np.outer(x > 0, x < 0)
            prefs[idx] += 1
            tots[idx] += 1
            tots[idx.T] += 1

        np.apply_along_axis(inner, 1, data.values)
        return prefs, tots


    prefs, tots = comp_prefs(barpr, aki2s)
    prefs = pd.DataFrame(prefs, index=aki2s, columns=aki2s)
    tots = pd.DataFrame(tots, index=aki2s, columns=aki2s)

    aki2_dict = isc[['aki2', 'isced2']].drop_duplicates().set_index('aki2')['isced2'].to_dict()
    aki2_dict_short = {k: v if len(v) < 20 else v[:18] + '...' for k, v in aki2_dict.items()}

    rp = pd.merge(tots.stack().rename('tot').reset_index(), prefs.stack().rename('pref').reset_index())
    rp['ppct'] = 100 * rp['pref'] / rp['tot']
    rp['x'] = rp['level_1'].map(aki2_dict_short)
    rp['y'] = rp['level_0'].map(aki2_dict_short)
    rp = rp.dropna(subset='ppct')
    rp['sqrt_tot'] = np.sqrt(rp['tot'])

    isced2s = [aki2_dict_short[x] for x in aki2s]
    rp['x'] = pd.Categorical(rp['x'], categories=isced2s, ordered=True)
    rp['y'] = pd.Categorical(rp['y'], categories=isced2s, ordered=True)
    rps[y] = rp
    fig, ax = plt.subplots()
    sns.scatterplot(data=rp, x='x', y='y', hue='ppct', palette='RdYlBu', size='sqrt_tot', sizes=(5, 3000))
    for _, row in rp.iterrows():
        color, alpha = ('black', 0.8) if 25 < row['ppct'] < 75 else ('white', 1)
        ax.text(x=row['x'], y=row['y'], s=f'{row["ppct"]:.0f}', va='center', ha='center', alpha=alpha, color=color)
    ax.set(xlabel=None, ylabel=None, title='Preference při zápisu do širších oborových skupin')
    plt.xticks(rotation=30, ha='right')
    plt.legend([], [], frameon=False)
    fig.tight_layout()
    figs.append(fig)

rt.Leaf(figs, title='Revealed preferences', num_cols=1).show()

rp17 = rps[17]
rp21 = rps[21]

cols = ['level_1', 'ppct', 'sqrt_tot']
rp17 = rp17[rp17['level_0'] == '01'][cols]
rp17['year'] = 'Rok 2017'
rp21 = rp21[rp21['level_0'] == '01'][cols]
rp21['year'] = 'Rok 2021'
rp = pd.concat([rp17, rp21])
rp['label'] = rp['level_1'].map(aki2_dict)
rp['label'] = pd.Categorical(rp['label'], categories=[aki2_dict[x] for x in aki2s], ordered=True)
rp['label_short'] = rp['level_1'].map(aki2_dict_short)
rp['label_short'] = pd.Categorical(rp['label_short'], categories=[aki2_dict_short[x] for x in aki2s], ordered=True)
rp['year'] = pd.Categorical(rp['year'])

font_size_word()

plt.rcParams['figure.figsize'] = 10, 7
fig, ax = plt.subplots()
sns.scatterplot(data=rp, x='year', y='label', hue='ppct', palette='RdYlBu', size='sqrt_tot', sizes=(500, 2000))
for _, row in rp.iterrows():
    color, alpha = ('black', 0.8) if 24 < row['ppct'] < 76 else ('white', 1)
    ax.text(x=row['year'], y=row['label'], s=f'{row["ppct"]:.0f}', va='center', ha='center', alpha=alpha, color=color)
ax.set(xlabel=None, ylabel=None)
fig.suptitle('Podíl zapsaných do pedagogických oborů')
ax.set(xlim=(-1, 2))
# plt.xticks(rotation=30, ha='right')
plt.legend([], [], frameon=False)
fig.tight_layout()

fig.show()

fig2, ax2 = plt.subplots()
bool(ax2.get_title())

rp

rps[17].columns

rp17 = rps[17]
rp21 = rps[21]

cols = ['ppct', 'sqrt_tot', 'pref', 'tot']
rp17 = rp17[rp17['level_0'] == '01'][['level_1'] + cols].rename(columns={c: f'{c}_17' for c in cols})
rp21 = rp21[rp21['level_0'] == '01'][['level_1'] + cols].rename(columns={c: f'{c}_21' for c in cols})

rp = pd.merge(rp17, rp21)
rp['label'] = rp['level_1'].map(aki2_dict)
# rp['label'] = pd.Categorical(rp['label'], categories=[aki2_dict[x] for x in aki2s], ordered=True)

# rp['label_short'] = rp['level_1'].map(aki2_dict_short)
# rp['label_short'] = pd.Categorical(rp['label_short'], categories=[aki2_dict_short[x] for x in aki2s], ordered=True)
# rp['year'] = pd.Categorical(rp['year'])

doc = Document('template.docx')
doc._body.clear_content()

h = uchazec.HNumbering()

section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(25.4)
section.right_margin = Mm(25.4)
section.top_margin = Mm(25.4)
section.bottom_margin = Mm(25.4)

doc.add_heading('IDEA: Uchazeč', 0)

# doc.add_paragraph(f'Celkem {tot:,} uchazečů se hlásilo a zároveň bylo přijato na pedagogický obor i na jiný obor. U těchto uchazečů je možné sledovat, na který z oborů se skutečně zapsali, jak shrnuje následující tabulka.')
# doc.add_paragraph('')

table = doc.add_table(rows=1, cols=3, style='Plain Table 1')
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Oborová skupina'
hdr_cells[1].text = 'Rok 2017'
hdr_cells[2].text = 'Rok 2021'

for _, row in rp.sort_values(['label']).iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row['label']
    row_cells[1].text = f'{row.ppct_17:.3g} % ({row.pref_17} / {row.tot_17})'
    row_cells[2].text = f'{row.ppct_21:.3g} % ({row.pref_21} / {row.tot_21})'

doc.show()

if y == 21:
    doc.add_paragraph('')
    doc.add_paragraph(
        'Data naznačují, že přijímací řízení na pedagogických oborech bylo v roce 2021 srovnatelně náročné jako v roce 2017. Také uchazeči, kteří byli zároveň přijatí na pedagogický a nepedagogický obor, si v obou letech vybírali pro zápis mírně častěji nepedagogický obor oproti pedagogickému.')

doc.add_heading(f'{h.h2} Úspěšnost podle krajů', 2)

rp.columns
foo = rp.sum()
foo['pref_17'] / foo['tot_17']
foo['pref_21'] / foo['tot_21']

# endregion


# TODO: ruzne dalsi komentare od Dana
# popis portfolii prihlasek
# efekt vzdalenosti?
# kam se zak zapsal, kdyz byl prijat na vice nez jednu skolu


# ABSOLVENTI
os.getcwd()

zaci = pd.read_excel("../data/uchazec/absolventi_ss.xlsx")
ab_raw = pd.read_excel("../data/uchazec/absolventi_ss.xlsx", sheet_name="List3")
ab = ab_raw.groupby('red_izo')[['absolventi_celkem', 'absolventi_divky']].sum().reset_index().rename(
    columns={"red_izo": "ss_red_izo"})

# df21 = pd.merge(df21, ab, how='left')

ped_col = 'ped_obor'
vars_to_sum = ['prihl_bool', 'prijat_bool', 'zapis_bool', f'prihl_{ped_col}_bool', f'prijat_{ped_col}_bool',
               f'zapis_{ped_col}_bool']
foo = po21.groupby(['gender', 'ss_red_izo'])[vars_to_sum].sum()

foo = foo.unstack(level=0).fillna(0).reset_index()
sfx = {'': '', 'Chlapec': '_chlapec', 'Dívka': '_divka'}
foo.columns = [x + sfx[y] for x, y in foo.columns]

foo = pd.merge(foo, ab, how='left')
foo = foo.rename(columns={'absolventi_divky': 'absolventi_divka'})
foo['absolventi_chlapec'] = foo['absolventi_celkem'] - foo['absolventi_divka']

for x in vars_to_sum:
    foo[f'{x}_celkem'] = foo[f'{x}_divka'] + foo[f'{x}_chlapec']

isc21['aki4_podil_prijat'] = isc21['prijat'] / isc21['total']
aki4 = isc21[['aki4', 'aki4_podil_prijat']]
df21 = pd.merge(df21, aki4, how='left')

ss_ambice = df21.groupby('ss_red_izo')[['aki4_podil_prijat', 'w']].apply(nanaverage, weights='w').reset_index()
ss_ambice2 = df21[~df21['ped_obor']].groupby('ss_red_izo')[['aki4_podil_prijat', 'w']].apply(nanaverage,
                                                                                             weights='w').reset_index().rename(
    columns={'aki4_podil_prijat': 'aki4_podil_prijat_no_ped'})
ss21 = pd.merge(foo, ss_ambice)
ss21 = pd.merge(ss21, ss_ambice2, how='left')

ss_kraj = df21[['ss_red_izo', 'ss_kraj']].drop_duplicates().dropna()
ss21 = pd.merge(ss21, ss_kraj)

foo = po21[['ss_red_izo', 'ss_typ']].drop_duplicates().dropna()
foo = foo[~foo['ss_typ'].isin(['Není SŠ', 'Jiné'])]
bar = foo['ss_red_izo'].value_counts()
multiple = bar[bar > 1]
foo = foo[~foo.ss_red_izo.isin(multiple.index)]

ss21 = pd.merge(ss21, foo, how='left')
ss21['ss_typ'] = ss21['ss_typ'].fillna('Jiné')

# now I should be able to compute everything


bar = po21[['ss_red_izo', 'ss_kraj', 'ss_typ', 'ss_typ_g']].drop_duplicates().dropna()
bar[bar['ss_red_izo'] == 600014321]

ss21['podil_uchazecu'] = ss21['prihl_bool_celkem'] / ss21['absolventi_celkem']
ss21['podil_ped_uchazecu'] = ss21[f'prihl_{ped_col}_bool_celkem'] / ss21['absolventi_celkem']
ss21['podil_ped_uchazecu_z_uchazecu'] = ss21[f'prihl_{ped_col}_bool_celkem'] / ss21['prihl_bool_celkem']
# ss21['ppuzu_disc'] = pd.cut(ss21['podil_ped_uchazecu_z_uchazecu'], bins=4)
ss21['ppuzu_disc'] = pd.cut(ss21['podil_ped_uchazecu_z_uchazecu'], bins=[-0.001, 0.25, 0.5, 1.001])
ss21['podil_divek'] = ss21['absolventi_divka'] / ss21['absolventi_celkem']
ss21['podil_divek_disc'] = pd.cut(ss21['podil_divek'], bins=[-0.001, 0.3, 0.7, 1.001])

to_drop = ['ppuzu_disc', 'podil_divek_disc']

weird = ss21[ss21['podil_uchazecu'] > 1]
weird.drop(columns=to_drop).show('weird')

# import warnings
# warnings.simplefilter(action='ignore', category=FutureWarning)

# priklady problematickych RED IZO skol:
# - 600007472: 1 absolvent, 0 zaku, ale v datech mame 16 uchazecu
# - 600005666: 1 absolvent, 171 zaku (to je trochu podezrele), mame 15 uchazecu
# - 600005836: 7 absolventu, 141 zaku, 43 uchazecu

# toto jsou nejkriklavejsi pripady, celkem je temer sto skol, kde mame vyssi pocet uchazecu nez absolventu. ackoli u vetsiny z nich jsou ty rozdily male, o 1-2 zaky

# zaci[zaci.red_izo == 600007472]
# zaci[zaci.red_izo == 600005666]
# ab_raw[ab_raw.red_izo == 600005666].T
# df21[df21.ss_red_izo == 600005666].show()
# ab_raw[ab_raw.p_izo == 10622110]
# ab_raw[ab_raw.red_izo == 600005836]
# zaci[zaci.red_izo == 600005836]

sns.scatterplot(data=ss21[ss21['podil_uchazecu'] <= 1.2], x='podil_uchazecu', y='aki4_podil_prijat',
                hue='podil_ped_uchazecu_z_uchazecu').show()
sns.scatterplot(data=ss21[ss21['podil_uchazecu'] <= 1.2], x='podil_uchazecu', y='aki4_podil_prijat', hue='ppuzu_disc',
                size='absolventi_celkem', sizes=(5, 200), alpha=0.8).show()
sns.scatterplot(data=ss21[ss21['podil_uchazecu'] <= 1.2], x='podil_uchazecu', y='aki4_podil_prijat',
                hue='podil_divek_disc', size='absolventi_celkem', sizes=(5, 200), alpha=0.8).show()
sns.scatterplot(data=ss21[ss21['podil_uchazecu'] <= 1.2], x='podil_uchazecu', y='aki4_podil_prijat', hue='ss_kraj',
                size='absolventi_celkem', sizes=(5, 200), alpha=0.8).show()
sns.scatterplot(data=ss21[ss21['podil_uchazecu'] <= 1.2], x='podil_uchazecu', y='aki4_podil_prijat', hue='ss_kraj',
                size='absolventi_celkem', sizes=(5, 200), alpha=0.8).show()
sns.scatterplot(data=ss21[ss21['podil_uchazecu'] <= 1.2], x='podil_uchazecu', y='aki4_podil_prijat', hue='ss_typ',
                size='absolventi_celkem', sizes=(5, 200), alpha=0.8).show()

sns.scatterplot(data=ss21, x='aki4_podil_prijat', y='podil_ped_uchazecu_z_uchazecu', hue='ss_typ',
                size='absolventi_celkem', sizes=(5, 200), alpha=0.8).show()
sns.scatterplot(data=ss21, x='aki4_podil_prijat_no_ped', y='podil_ped_uchazecu_z_uchazecu', hue='ss_typ',
                size='absolventi_celkem', sizes=(5, 200), alpha=0.8).show()

ss21.groupby('ss_kraj')[
    ['absolventi_celkem', 'podil_divek', 'podil_uchazecu', 'podil_ped_uchazecu', 'aki4_podil_prijat']].apply(nanaverage,
                                                                                                             weights='absolventi_celkem').reset_index()

fig, ax = plt.subplots()
sns.lineplot(x=range(10), y=range(10), marker='o')
fig.show()

ss21['absolventi_celkem'] = ss21['absolventi_celkem'].fillna(0)

ss21.groupby('ss_kraj')[
    ['absolventi_celkem', 'podil_divek', 'podil_uchazecu', 'podil_ped_uchazecu', 'aki4_podil_prijat']].apply(nanaverage,
                                                                                                             weights='absolventi_celkem').reset_index().show()
ss21.groupby('ss_kraj')[
    ['absolventi_celkem', 'podil_divek', 'podil_uchazecu', 'podil_ped_uchazecu', 'aki4_podil_prijat']].apply(nanaverage,
                                                                                                             weights='absolventi_celkem').reset_index()

ss21.groupby('ss_kraj')[['absolventi_celkem', 'aki4_podil_prijat']].apply(nanaverage,
                                                                          weights='absolventi_celkem').reset_index()
ss21.groupby('ss_kraj')[['absolventi_celkem', 'aki4_podil_prijat']].average()

for k, v in ss21.groupby('ss_kraj')[['absolventi_celkem', 'aki4_podil_prijat']].groups.items():
    break

k
v
x = ss21.iloc[v]
weights = 'absolventi_celkem'

w = x[weights]
x = x.drop(columns=[weights])
mask = np.isnan(x)
xm = np.ma.masked_array(x, mask=mask)
if len(x.shape) == 1:
    return np.ma.average(xm, weights=w)
else:
    res = np.ma.average(xm, weights=w, axis=0)
    return pd.Series(res, x.columns) if isinstance(x, pd.DataFrame) else res

w

po21

df21.shape
np.sum(df21.ss_red_izo == 600020070)
df21.show()

ab.p_izo.dtype
df21.ss_izo.dtype

pd.isna(df21.ss_izo).sum()

redizo = 600020070

redizo in ab.red_izo.dtype

ab[ab.red_izo == redizo]

df_izos = df21[["ss_red_izo"]].value_counts().rename('count').reset_index().rename(columns={"ss_red_izo": "red_izo"})
foo = pd.merge(ab[["red_izo"]].drop_duplicates(), df_izos, how="outer", indicator=True)

foo.groupby("_merge")["count"].sum()
foo.show(num_rows=1500)
foo["_merge"].value_counts()

df_izos = df21[["ss_red_izo"]].drop_duplicates().rename(columns={"ss_red_izo": "p_izo"})
foo = pd.merge(ab[["p_izo"]].drop_duplicates(), df_izos, how="outer", indicator=True)
foo.show()

foo["_merge"].value_counts()

df21[df21['ss_red_izo'] == 691001197]['ss_izo'].drop_duplicates()

sys_root = 'D:\\' if sys.platform == 'win32' else '/mnt/d'
data_root = os.path.join(sys_root, 'projects', 'idea', 'data')

li = []
for y in range(1999, 2023):
    pdf = pd.read_csv(f'{data_root}/uchazec/stredni-skoly/m{y}.csv', encoding='cp1250')
    pdf['year'] = y
    li.append(pdf)
ss = pd.concat(li, axis=0, ignore_index=True).drop_duplicates().reset_index(drop=True)
ss = ss.sort_values('year', ascending=False).drop_duplicates('IZO').sort_index()
ss['OBOR'] = ss['OBOR'].str.strip()

sss = ss[['IZO', 'VUSC']].drop_duplicates().reset_index(drop=True)
sss.shape

ssss = ss[['IZO', 'VUSC', 'RED_IZO']].drop_duplicates().reset_index(drop=True)
ssss.shape

# year can be 17 or 21
year = 21
path = os.path.join(data_root, 'uchazec', f'0022MUCH{year}P')
df = pd.read_csv(f'{path}.csv', encoding='cp1250', low_memory=False)
df.columns

df_izos = df21[["ss_izo"]].value_counts().rename('count').reset_index().rename(columns={"ss_izo": "p_izo"})
foo = pd.merge(ab[["p_izo"]].drop_duplicates(), df_izos, how="outer", indicator=True)

foo.groupby("_merge")["count"].sum()
foo.show(num_rows=1500)

# KAM SE ZAPSAL
# region
ped_col = 'ped_obor'
y = 21
ff = eval(f'po{y}').copy()
df = eval(f'df{y}').copy()
isc = eval(f'isc{y}').copy()

# ...
foo = df[df['prijat']].copy()
akivar = 'aki4'
foo_id = foo[['id', akivar]].drop_duplicates()['id'].value_counts()
foo_id_mult = foo_id[foo_id > 1].index

foo = foo[foo['id'].isin(foo_id_mult)].copy()
bar = foo.groupby(['id', akivar])['zaps_zapsal'].sum().reset_index()
bar['pref'] = np.where(bar['zaps_zapsal'] > 0.5, 1, -1)

akis = isc[akivar].drop_duplicates().dropna().sort_values().values

barpr = bar.pivot(index='id', columns=akivar, values='pref').fillna(0).astype(int)
barpr['min'] = barpr.min(axis=1)
barpr['max'] = barpr.max(axis=1)
barpr = barpr[(barpr['min'] == -1) & (barpr['max'] == 1)].drop(columns=['min', 'max'])


def comp_prefs(data, akis):
    assert data.shape[1] == len(akis)
    prefs = np.full([len(akis), len(akis)], 0)
    tots = np.full([len(akis), len(akis)], 0)

    def inner(x):
        idx = np.outer(x > 0, x < 0)
        prefs[idx] += 1
        tots[idx] += 1
        tots[idx.T] += 1

    np.apply_along_axis(inner, 1, data.values)
    return prefs, tots


start = datetime.now()
akis = barpr.columns.values
prefs, tots = comp_prefs(barpr, akis)
print(f'{datetime.now() - start} s')

prefs = pd.DataFrame(prefs, index=akis, columns=akis)
tots = pd.DataFrame(tots, index=akis, columns=akis)

valid = tots.index[tots.max() >= 10]
prefs = prefs[valid].loc[valid]
tots = tots[valid].loc[valid]

akis = tots.columns.values

aki_dict = isc[[akivar, 'isced' + akivar[-1]]].drop_duplicates().set_index(akivar)['isced' + akivar[-1]].to_dict()
aki_dict_short = {k: v + ' [' + k + ']' if len(v) < 20 else v[:18] + '... [' + k + ']' for k, v in aki_dict.items()}

len(akis)
len(isc['aki4'].drop_duplicates().dropna().sort_values().values)

rp = pd.merge(tots.stack().rename('tot').reset_index(), prefs.stack().rename('pref').reset_index())
rp['ppct'] = 100 * rp['pref'] / rp['tot']
rp['ppct_fltr'] = np.where(rp['tot'] < 10, np.nan, rp['ppct'])
rp['x'] = rp['level_1'].map(aki_dict_short)
rp['y'] = rp['level_0'].map(aki_dict_short)
rp = rp.dropna(subset='ppct')
rp['sqrt_tot'] = np.sqrt(rp['tot'])

isceds = [aki_dict_short[x] for x in akis]

rp['x'] = pd.Categorical(rp['x'], categories=isceds, ordered=True)
rp['y'] = pd.Categorical(rp['y'], categories=isceds, ordered=True)

fig, ax = plt.subplots(figsize=(30, 18))
sns.scatterplot(data=rp, x='x', y='y', hue='ppct_fltr', palette='RdYlBu', size='sqrt_tot', sizes=(5, 3000))
for _, row in rp.iterrows():
    if not pd.isna(row['ppct_fltr']):
        color, alpha = ('black', 0.8) if 25 < row['ppct'] < 75 else ('white', 1)
        ax.text(x=row['x'], y=row['y'], s=f'{row["ppct"]:.0f}', va='center', ha='center', alpha=alpha, color=color)
ax.set(xlabel=None, ylabel=None, title='Revealed Preferences při zápisu do oborů')
plt.xticks(rotation=30, ha='right')
plt.legend([], [], frameon=False)
fig.tight_layout()
fig.show()

tots

# APRIORI
ped_col = 'ped_obor'
y = 21
ff = eval(f'po{y}').copy()
df = eval(f'df{y}').copy()
isc = eval(f'isc{y}').copy()

df['isced2'].cat.categories
df['isced2'] = pd.Categorical(df['isced2'])
apr = df[['id', 'isced2']].drop_duplicates()
apr['one'] = 1
apr.pivot(index='id', columns='isced2', values='one')[df['isced2'].cat.categories].fillna(0)

aki2_dict = isc[['aki2', 'isced2']].drop_duplicates().set_index('aki2')['isced2'].to_dict()
aki2_dict_short = {k: v if len(v) < 20 else v[:18] + '...' for k, v in aki2_dict.items()}

aki2s = df['aki2'].drop_duplicates().dropna().sort_values().values
apr = df[['id', 'aki2']].drop_duplicates()
apr['one'] = 1
apr = apr.pivot(index='id', columns='aki2', values='one')[aki2s].fillna(0)

conf = pd.DataFrame(index=apr.columns, columns=apr.columns, dtype='float')

(apr['01'] * apr['09']).sum() / apr['01'].sum()

for i in aki2s:
    for j in aki2s:
        # conf[X -> Y] = # X & Y / # X
        conf.loc[i, j] = (apr[i] * apr[j]).sum() / apr[i].sum()
        # if i == j:
        #     conf.loc[i, j] = np.nan

conf.columns = conf.columns.map(aki2_dict_short)
conf.index = conf.index.map(aki2_dict_short)

mask = pd.DataFrame(np.eye(10, dtype='bool'), columns=conf.columns, index=conf.index)

fig, ax = plt.subplots()
sns.heatmap(100 * conf, annot=True, mask=mask, cmap='YlOrBr')
ax.set(xlabel=None, ylabel=None, title='Apriori: Asociace mezi jednotlivými obory [řádky --> sloupce]')
fig.tight_layout()
fig.show()

isced2 = isc[['aki2', 'isced2']].drop_duplicates().set_index('aki2').sort_index().values[:, 0]

# kam se hlasi ve slozeni oboru?
isc21.show('isc21')

ped_col = 'ped_obor'
y = 21
ff = eval(f'po{y}').copy()
df = eval(f'df{y}').copy()
isc = eval(f'isc{y}').copy()

# aplikacni ambice jednotlivych oboru
isc['ap_ambice'] = isc['prijat'] / isc['total']

df = pd.merge(df, isc[['aki4', 'ap_ambice']])

uchp = ff[ff[f'prihl_{ped_col}_bool']][['id']]
uchnp = ff[~ff[f'prihl_{ped_col}_bool']][['id']]
dfp = pd.merge(df, uchp)
dfp_p = dfp[dfp['aki2'] == '01']
dfp_np = dfp[dfp['aki2'] != '01']
dfnp = pd.merge(df, uchnp)

np.average(df['ap_ambice'], weights=df['w'])
np.average(dfp['ap_ambice'], weights=dfp['w'])
np.average(dfp_p['ap_ambice'], weights=dfp_p['w'])
np.average(dfp_np['ap_ambice'], weights=dfp_np['w'])
np.average(dfnp['ap_ambice'], weights=dfnp['w'])

df.show('df')

np.round(100 * df['isced2'].value_counts() / df.shape[0], 1)

frp = (100 * dfp['isced2'].value_counts() / dfp.shape[0]).rename(ped_col).reset_index()
frnp = (100 * dfnp['isced2'].value_counts() / dfnp.shape[0]).rename(f'ne{ped_col}').reset_index()
fr = pd.merge(frp, frnp, how='outer').fillna(0)
frr = fr.set_index('index').unstack().rename('pct').reset_index()
frr['label'] = np.where(frr['level_0'] == 'ped_obor', 'Pedagogičtí', 'Nepedagogičtí')

fig, ax = plt.subplots()
sns.barplot(data=frr, x='pct', y='index', hue='label')
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.get_legend().set_title('Uchazeči')
ax.set(title='Rozdělení přihlášek mezi obory', xlabel='Relativní četnost', ylabel=None)
fig.tight_layout()
fig.show()

df.crosstab('isced2')
df.cross_tab('isced2')

df

po21.show('po21')

ped_col = 'ped_obor'
tot_prihlasek = eval(f'df{y}').shape[0]
tot_uchazecu = ff.shape[0]
ff_chlapci = ff[ff['gender'] == 'Chlapec']
ff_divky = ff[ff['gender'] == 'Dívka']
tot_chlapcu = ff_chlapci.shape[0]
tot_divek = ff_divky.shape[0]

# celkové statistiky
avg_prihlasek = ff['prihl'].mean()
max_prihlasek = ff['prihl'].max()
med_prihlasek = ff['prihl'].median()

# chlapci
avg_prihlasek_chlapci = ff_chlapci['prihl'].mean()
max_prihlasek_chlapci = ff_chlapci['prihl'].max()
med_prihlasek_chlapci = ff_chlapci['prihl'].median()

# dívky
avg_prihlasek_divky = ff_divky['prihl'].mean()
max_prihlasek_divky = ff_divky['prihl'].max()
med_prihlasek_divky = ff_divky['prihl'].median()

doc.add_heading(f'{h.h2} Přihlášky 20{y}', 2)
doc.add_paragraph(
    f'V roce 20{y} bylo podáno {tot_prihlasek:,} přihlášek na vysoké školy od {tot_uchazecu:,} českých uchazečů, z toho se hlásilo {tot_chlapcu:,} chlapců ({100 * tot_chlapcu / tot_uchazecu:.3g} %) a {tot_divek:,} dívek ({100 * tot_divek / tot_uchazecu:.3g} %). Průměrný uchazeč si podal {avg_prihlasek:.1f} přihlášek (chlapci {avg_prihlasek_chlapci:.1f}, dívky {avg_prihlasek_divky:.1f}), mediánový uchazeč {med_prihlasek:.0f} (chlapci {med_prihlasek_chlapci:.0f}, dívky {med_prihlasek_divky:.0f}), maximální počet podaných přihlášek byl {max_prihlasek:.0f} (chlapci {max_prihlasek_chlapci:.0f}, dívky {max_prihlasek_divky:.0f}). Četnost počtu přihlášek jednotlivých uchazečů ukazuje následující graf.')
doc.add_paragraph('')

# fig, ax = plt.subplots()
# ax = sns.histplot(data=ff, y='prihl', discrete=True, binwidth=0.8, shrink=0.8, ec=None)
# ax.set(xlabel='Četnost', ylabel='Počet přihlášek', title='Počet přihlášek jednotlivých uchazečů')
# ax.set_ylim(0, 13)
# ax.invert_yaxis()
# fig.tight_layout()

# v procentech, rozdělení chlapci a dívky
gender_order = ['Chlapec', 'Dívka']
gender_colors = ['tab:blue', 'tab:red']

fig, ax = plt.subplots()
ax = sns.histplot(data=ff[ff[f'prihl_{ped_col}_bool']], y='prihl', discrete=True, binwidth=0.8, shrink=0.8, ec=None,
                  hue='gender', stat='percent', multiple='stack', hue_order=gender_order, palette=gender_colors)
ax.set(xlabel='Četnost', ylabel='Počet přihlášek', title='Počet přihlášek jednotlivých uchazečů')
ax.set_ylim(0, 13)
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.get_legend().set_title('Gender')
ax.invert_yaxis()
fig.tight_layout()

fig.show()

ff[ff[f'prihl_{ped_col}_bool']][['gender', f'prihl_{ped_col}', f'prihl_ne{ped_col}']]

ff[ff[f'prihl_{ped_col}_bool']].groupby('prihl')[f'prihl_{ped_col}', f'prihl_ne{ped_col}'].mean()
ff[ff[f'prihl_{ped_col}_bool']].groupby('prihl')[[f'prihl_{ped_col}', f'prihl_ne{ped_col}', 'prihl']].agg(
    {f'prihl_{ped_col}': 'mean', f'prihl_ne{ped_col}': 'mean', 'prihl': ' count'})

ff['divka'] = ff['gender'] == 'Dívka'
ff['chlapec'] = ff['gender'] == 'Chlapec'
ff['count'] = 1

# foo = ff[ff[f'prihl_{ped_col}_bool']].groupby('prihl').agg({f'prihl_{ped_col}': 'sum', f'prihl_ne{ped_col}': 'sum', 'id': 'count', 'divka': 'sum'}).reset_index()
foo = ff[ff[f'prihl_{ped_col}_bool']].groupby('prihl').['count', 'divka', 'chlapec', f'prihl_{ped_col}': 'sum',
      f'prihl_ne{ped_col}': 'sum', 'id': 'count', 'divka': 'sum'}).reset_index()

foo[f'{ped_col}_ratio'] = foo[f'prihl_{ped_col}'] / foo['prihl']
foo = foo.rename(columns={'id': 'count'})
foo['divka_count'] = foo['divka'] * foo['count']
foo['chlapec_count'] = foo['count'] - foo['divka_count']
foo[f'{ped_col}_count'] = foo[f'{ped_col}_ratio'] * foo['count']

ff[ff[f'prihl_{ped_col}_bool']][['prihl', f'prihl_{ped_col}']].mean()

foo = ff[ff[f'prihl_{ped_col}_bool']].groupby(['gender', 'prihl']).agg(
    {'count': 'sum', f'prihl_{ped_col}': 'mean', f'prihl_ne{ped_col}': 'mean'}).reset_index()
foo.show()

tot_count = foo['count'].sum()
foo['ratio'] = foo['count'] / tot_count
foo[f'prihl_{ped_col}_pct'] = 100 * foo['ratio'] * foo[f'prihl_{ped_col}'] / foo['prihl']
foo[f'prihl_ne{ped_col}_pct'] = 100 * foo['ratio'] * foo[f'prihl_ne{ped_col}'] / foo['prihl']

bar = foo.set_index(['gender', 'prihl'])[[f'prihl_{ped_col}_pct', f'prihl_ne{ped_col}_pct']].stack().rename(
    'pct').reset_index()

bar['label'] = bar.apply(
    lambda row: f'{row.gender}, {"pedagogický" if row.level_2 == f"prihl_{ped_col}_pct" else "nepedagogický"}', axis=1)
bar['prihl'] = bar['prihl'].astype(int)
bar = bar[bar['prihl'] <= 12]

hue_order = ['Dívka, pedagogický', 'Dívka, nepedagogický', 'Chlapec, pedagogický', 'Chlapec, nepedagogický']
hue_color = ['tab:red', 'tab:orange', 'tab:blue', 'tab:cyan']

barcum = bar.set_index(['prihl', 'label'])['pct'].unstack()[hue_order].cumsum(axis=1).stack().rename(
    'pct').reset_index()

fig, ax = plt.subplots()
sns.barplot(data=barcum, x='pct', y='prihl', hue='label', orient='h', dodge=False, hue_order=hue_order[::-1],
            palette=hue_color[::-1], width=0.7, alpha=0.8)
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.show()

# endregion
# OLDER STUFF
# region


df21.show('df21')
df21['typ_st'].value_counts()
df21[df21['typ_st'] == 'Magisterský'].show('df21_magistersky')

df21[df21['typ_st'] == 'Magisterský']['fak_nazev'].cat.remove_unused_categories().value_counts()

df21['typ_st'].value_counts()
df17['typ_st'].value_counts()

# TODO: diskrepance v datech pro jednoho uchazeče

other_keys = ['gender', 'rmat', 'ss_kraj', 'ss_typ', 'ss_gym_delka', 'ss_typ_g']
df_ids = df.groupby('id')[other_keys].first().reset_index()

# 21
df = df21.copy()

df.shape[0]  # 144 118
foo = df[['id'] + other_keys].drop_duplicates()
df[['id']].drop_duplicates()

multiple_ids = foo['id'].value_counts()
dups = foo[foo['id'].isin(multiple_ids[multiple_ids > 1].index)]

dups.show('dups21', num_rows=2000)

dups[dups['id'] == '0059QQ1542']

fr = dups[dups['id'] == '0006QQ6740']

dups[dups['id'] == '0257QQ9509']

for idx, fr in dups.groupby('id'):
    break

idx
fr


def fix_duplicates(fr):
    if fr.shape[0] == 1:
        return fr
    else:
        res = pd.Series(dtype='object')
        res['id'] = fr['id'].iloc[0]

        # droppers!
        drop_duplicates = lambda x: x.drop_duplicates()
        drop_nan = lambda x: x.dropna()
        drop_jine = lambda x: x[x != 'Jiné'].copy()
        drop_neni_ss = lambda x: x[x != 'Není SŠ'].copy()
        drop_gym = lambda x: x[x != 'Gymnázium'].copy()
        drop_0 = lambda x: x[x != 0].copy()
        drop_6 = lambda x: x[x != 6].copy()
        drop_8 = lambda x: x[x != 8].copy()
        drop_gym6 = lambda x: x[x != 'Gym 6-leté'].copy()
        drop_gym8 = lambda x: x[x != 'Gym 8-leté'].copy()

        droppers = {
            'gender': [drop_duplicates, drop_nan],
            'rmat': [drop_duplicates, drop_nan],
            'ss_kraj': [drop_duplicates, drop_nan],
            'ss_typ': [drop_duplicates, drop_nan, drop_jine, drop_neni_ss, drop_gym],
            'ss_gym_delka': [drop_duplicates, drop_nan, drop_0, drop_6, drop_8],
            'ss_typ_g': [drop_duplicates, drop_nan, drop_jine, drop_neni_ss, drop_gym6, drop_gym8],
        }

        for k, v in droppers.items():
            x = fr[k]
            for d in v:
                x = d(x)
                if len(x) == 1:
                    res[k] = x.iloc[0]
                    break
            if len(x) > 1:
                print(
                    f'ERROR IN DEDUPLICATION FOR {res["id"]}: {k} has values [{", ".join(x.astype("str"))}], using the first one.')
                res[k] = x.iloc[0]

        return res


df21['ss_typ_g'].value_counts()

logger('')
foo = df21[['id'] + other_keys].drop_duplicates().groupby('id').apply(fix_duplicates)
logger('')

logger('')
df_ids_dups = df21[['id'] + other_keys].drop_duplicates()
multiple_ids = df_ids_dups['id'].value_counts()
dups = df_ids_dups[df_ids_dups['id'].isin(multiple_ids[multiple_ids > 1].index)]
no_dups = df_ids_dups[df_ids_dups['id'].isin(multiple_ids[multiple_ids == 1].index)]
fixed_dups = dups.groupby('id').apply(fix_duplicates)
logger('')

no_dups
fixed_dups

df_ids = pd.concat([no_dups, fixed_dups])

df_ids['ss_typ'].value_counts()

pf21['ss_typ'].value_counts()
pf21.shape[0]

dups

dups.shape
foo = dups.groupby('id').apply(fix_duplicates)

x = fr['ss_gym_delka']
x
drop_duplicates(x)
drop_nan(x)
len(x)

", ".join(x)

x.astype('str')

fix_duplicates(fr)

for x in ['gender', 'rmat', 'ss_kraj']:
    xcl = fr[x].dropna().drop_duplicates()
    if len(xcl) == 1:
        res[x] = xcl[0]
    else:
        print(f'ERROR IN DEDUPLICATION: {x} has values [{", ".join(xcl)}], using the first one.')
        res[x] = xcl[0]

for x in ['ss_typ']

xcl = fr['ss_typ'].drop_duplicates()
if len(xcl) == 1:
    res['ss_typ'] = xcl[0]
else:
    xcl[xcl != 'Jiné']
    if len(xcl) == 1:
        res['ss_typ'] = xcl[0]
    else:
        xcl[xcl != 'Není SŠ']
        if len(xcl) == 1:
            res['ss_typ'] = xcl[0]
        else:
            print(f'ERROR IN DEDUPLICATION: ss_typ has values [{", ".join(xcl)}], using the first one.')
            res['ss_typ'] = xcl[0]


def get_per_app(df, ped_col='pedf'):
    print(f'Dataset podle uchazečů - připravuji')
    df['prijat'] = df['vypr'].str.startswith('Přijat')
    df['zaps_neprijat'] = df['zaps'] == 'Uchazeč nebyl přijat'
    df['zaps_zapsal'] = df['zaps'] == 'Uchazeč se zapsal'
    df['zaps_nezapsal'] = df['zaps'] == 'Uchazeč se nezapsal'

    other_keys = ['gender', 'rmat', 'ss_kraj', 'ss_typ', 'ss_gym_delka', 'ss_typ_g']
    df_ids = df.groupby('id')[other_keys].first().reset_index()

    variables = ['ones', 'prijat', 'zaps_zapsal']
    cols = [f'prihl_ne{ped_col}', f'prihl_{ped_col}', f'prijat_ne{ped_col}', f'prijat_{ped_col}', f'zapis_ne{ped_col}',
            f'zapis_{ped_col}']
    foo = df.groupby(['id', ped_col])[variables].sum().unstack(ped_col)
    foo.columns = cols
    foo = foo.fillna(0).reset_index()

    ff = pd.merge(df_ids, foo)

    ff['prihl'] = ff[f'prihl_{ped_col}'] + ff[f'prihl_ne{ped_col}']
    ff['prijat'] = ff[f'prijat_{ped_col}'] + ff[f'prijat_ne{ped_col}']
    ff['zapis'] = ff[f'zapis_{ped_col}'] + ff[f'zapis_ne{ped_col}']

    for c in cols:
        ff[f'{c}_bool'] = ff[c] > 0

    ff['prihl_bool'] = ff[f'prihl_{ped_col}_bool'] | ff[f'prihl_ne{ped_col}_bool']
    ff['prijat_bool'] = ff[f'prijat_{ped_col}_bool'] | ff[f'prijat_ne{ped_col}_bool']
    ff['zapis_bool'] = ff[f'zapis_{ped_col}_bool'] | ff[f'zapis_ne{ped_col}_bool']

    print(f'Dataset podle uchazečů - hotovo')
    return ff


# TO DO?
# - obdoba krajů pro typy střední školy
# - úspěšnost uchazečů
# - ???

# endregion
# region

isc21.show('isc21')

df21[['aki2', 'ped_obor']].value_counts().unstack()

df21.groupby('aki4')
pd.crosstab(df21['aki2'], df21['ped_obor'], margins=True)
pd.crosstab(df21['aki2'], df21['ped_obor'], margins=True, normalize='index')

dfpo21 = df21[df21['ped_obor']].copy()
dfpo17 = df17[df17['ped_obor']].copy()
dfpo17['isced4short'] = dfpo17['aki4'].map(isced4short)
dfpo21.shape

pd.crosstab(dfpo21['isced4short'], dfpo21['gender'], margins=True).reset_index().iloc[[3, 0, 4, 2, 1]]
pd.crosstab(dfpo21[dfpo21['zaps_zapsal']]['isced4short'], dfpo21[dfpo21['zaps_zapsal']]['gender'],
            margins=True).reset_index().iloc[[3, 0, 4, 2, 1]]

pd.crosstab(dfpo17[dfpo17['zaps_zapsal']]['isced4short'], dfpo17[dfpo17['zaps_zapsal']]['gender'],
            margins=True).reset_index().iloc[[3, 0, 4, 2, 1]]
pd.crosstab(dfpo17['isced4short'], dfpo17['gender'], margins=True).reset_index().iloc[[3, 0, 4, 2, 1]]

# podily prijatych a nasledne zapsanych
foo21_tot = pd.crosstab(dfpo21['isced4short'], dfpo21['gender'], margins=True).iloc[[3, 0, 4, 2, 1]]
foo21_zaps = \
pd.crosstab(dfpo21[dfpo21['zaps_zapsal']]['isced4short'], dfpo21[dfpo21['zaps_zapsal']]['gender'], margins=True).iloc[
    [3, 0, 4, 2, 1]]
foo21_prij = \
pd.crosstab(dfpo21[dfpo21['prijat']]['isced4short'], dfpo21[dfpo21['prijat']]['gender'], margins=True).iloc[
    [3, 0, 4, 2, 1]]

foo21_prij / foo21_tot
foo21_zaps / foo21_prij

y = 17
foo_tot = dfpo['isced4short'].value_counts().sort_index().iloc[sorter[y]]
foo_prij = dfpo[dfpo['prijat']]['isced4short'].value_counts().sort_index().iloc[sorter[y]]
foo_zaps = dfpo[dfpo['zaps_zapsal']]['isced4short'].value_counts().sort_index().iloc[sorter[y]]

foo_prij / foo_tot

pd.crosstab(dfpo21['isced4'], dfpo21['gender'], margins=True, normalize='index').iloc[[3, 2, 4, 1, 0]]

pd.crosstab(dfpo21['isced4'], dfpo21['gender'], margins=True).iloc[[3, 2, 4, 1, 0, 5]]
pd.crosstab(dfpo21['isced4'], dfpo21['gender'], margins=True).iloc[[3, 2, 4, 1, 0, 5]]
np.round(
    100 * pd.crosstab(dfpo21['isced4'], dfpo21['gender'], margins=True, normalize='index').iloc[[3, 2, 4, 1, 0, 5]], 1)

dfpo21[['isced4', 'gender']].value_counts()

isced4short = {
    '0111': 'Pedagogika',
    '0112': 'Předškolní vzdělávání a výchova',
    '0113': 'Bez předmětové specializace',
    '0114': 'S předmětovou specializací',
    '0188': 'Interdisciplinární programy',
}

dfpo21['isced4short'] = dfpo21['aki4'].map(isced4short)

pd.crosstab(df21['aki2'], df21['ped_obor'], margins=True, normalize='index')

foo = dfpo21[['isced4short', 'gender']].value_counts().rename('count').reset_index()

pd.crosstab(df21['isced4short'], df21['ped_obor'], margins=True, normalize='index')

sns.barplot(data=foo, y='isced4short', x='count', hue='gender').show()

dfpo21['isced4short']

foo = pd.crosstab(dfpo21['isced4short'], dfpo21['gender'], margins=True).reset_index().iloc[[3, 0, 4, 2, 1]]

plt.rcParams['figure.figsize']

fig, ax = plt.subplots(figsize=(12, 4))
colors = sns.color_palette(n_colors=2)
plt.subplots_adjust(left=0.3, bottom=0.16)
sns.barplot(data=foo, y='isced4short', x='All', label='Chlapci', ax=ax, color=colors[0], width=0.6)
sns.barplot(data=foo, y='isced4short', x='Dívka', label='Dívky', ax=ax, color=colors[1], width=0.6)
ax.legend()
ax.set(xlabel='Počet uchazečů', ylabel='')
ax.show()

dfpo21.columns

plt.close('all')

foo = pd.crosstab(dfpo21[dfpo21['zaps_zapsal']]['isced4short'], dfpo21[dfpo21['zaps_zapsal']]['gender'],
                  margins=True).reset_index().iloc[[3, 0, 4, 2, 1]]

fig, ax = plt.subplots(figsize=(12, 4))
colors = sns.color_palette(n_colors=2)
plt.subplots_adjust(left=0.3, bottom=0.16)
sns.barplot(data=foo, y='isced4short', x='All', label='Chlapci', ax=ax, color=colors[0], width=0.6)
sns.barplot(data=foo, y='isced4short', x='Dívka', label='Dívky', ax=ax, color=colors[1], width=0.6)
ax.legend()
ax.set(xlabel='Počet zapsaných', ylabel='')
ax.show()

ped_obory = isc21[isc21['aki2'] == '01'].sort_values('aki4')
ped_obory

df21.show('df21')
df21.shape
tot21 = df.shape[0]

ppo21 = df21[df21['ped_obor']]
totpo21 = ppo21.shape[0]

ppo21['isced4'].value_counts()
df17[df17['ped_obor']]['isced4'].value_counts()

isc17.show('isc17')
isc21.show('isc21')

df17[df17['ped_obor']].show('df17_ped_obor')

raw17, _ = uchazec.filter_data(pd.read_parquet('temp/uchazec/uch17.parquet'))

raw17.show('raw17', num_rows=10_000)

df17.shape
df17[df17['ped_obor']].show('df17_ped_obor')

np.sum(df17['program'].str[1:5] == '7504')

raw17 = pd.read_parquet('temp/uchazec/uch17.parquet')
np.sum(raw17['program'].str[1:5] == '7504')

np.sum(pd.isna(raw17['program']))
np.sum(raw17['program'] == '')
np.sum(raw17['program'].str.len != 5)
raw17[raw17['program'].str.len != 5]

raw17['program'][raw17['program'].str[1] == '7'].value_counts()

raw17[raw17['program'] == 'B7507'].show()
df17[df17['program'] == 'B7507'].show('df17_B7507')

raw17[raw17['obor1'].str.startswith('7504')].show('raw17_7504', num_rows=5_000)

###  NÁVRHY OD DANA  ###
# region
# 1.	Muzeme fakulty rozdelit podle zakladnich oborovych skupin (mediciny, humanitni, prirodovedni, technicke, humanitni, spolecenske]
# 2.	Muzeme fakulty rozdelit podle miry prevysu poptavky na zaklade oboru (nejvyssi v pravech a medicine a psychologii, mensi v prirodnich vedach humanitnich a spolecenskych, nejmensi v technickych)
# 3.	Muzeme rozdelit stredni skoly podle aplikacnich ambici jejich uchazcu hlasit se na VS s velkym previsem
# 4.	Muzeme rozdelit stredni skoly podle uspesnosti jejich uchazecu na vysoce poptavanych VS
# 5.	Muzeme rozdelit stredni skoly podle podilu uchazecu, kteri se nedostali na zadnou skolu
# 6.	Muzeme rozdelit stredni skoly podle podilu holek mezi uchazeci
# 7.	Muzeme se podivat, kam se zapsali uchazeci o pedf pokud byli prijati zaroven na jinou skolu (podle previsu ci oborove skupiny)

# ad 1:
# - podívej se, co vlastně všechno máme za fakulty, jestli to zvládneme nějak snadno manuálně přiřadit
# - nešla by použít informace ze studijních programů? jejich začátek tuším byl obor?

# - 2021: studijní programy jsou pravděpodobně konstruovány pomocí ISCED-F-2013
# - 2017: existuje převodník studijních programů SP <--> ISCED
# - a zdá se, že ISCED-F-2013 je konstruován hierarchicky, tedy umožňuje snadné sdružování do vyšších jednotek

# vytvoř něco jako ZAPS dataset na kompletních datech
# ok, isc17 a isc21


# ÚSPĚŠNOST CHLAPCŮ A DÍVEK
# region

# plain: chlapci 63 %, dívky 47 %
df21.groupby('gender')[['prijat', 'w']].apply(nanaverage, weights='w')

df21[df21['pedf']].groupby('gender')[['vypr_flag', 'w']].apply(nanaverage, weights='w')
df21.shape
df21[df21['pedf_isc']].groupby('gender')[['vypr_flag', 'w']].apply(nanaverage, weights='w')

# správně bych ale měl porovnávat přímo obory
stata.pdataframe_to_data(df17, force=True)

% % stata
encode
aki4, gen(aki4c)
encode
gender, gen(genderc)
logit
prijat
genderc
i.aki4c[pw = w]

% stata
logit
prijat
genderc[pw = w]

% stata
logit
vypr_flag
genderc
i.aki4c[pw = w]

% stata
logit
vypr_flag
genderc
i.aki4c[pw = w]

logistic(-0.0580242)
logistic(-0.102552)

po21['gender'].value_counts()
po21.columns

po21[po21['zapis_bool']]['gender'].value_counts()

logistic(1.138966)
logistic(1.138966 - .6199063)

# koeficient na gender: -.0660265
logistic(-.0660265)  # 0.4835 -> tzn. dívky se dostávají v průměru o necelá dvě procenta méně často

ff21['gender'].value_counts()
# Ale hlásí se jich o třetinu více než chlapců...
# Dívka      28835
# Chlapec    21640

ff21[ff21['prijat_bool']]['gender'].value_counts()
# U přijatých je rozdíl menší, ale pořád výrazný:
# Dívka      23879
# Chlapec    19479

# ROZDÍLY V PŘIJETÍ PODLE TYPŮ ŠKOL
df21.show()

df21.groupby('ss_typ')[['vypr_flag', 'w']].apply(nanaverage, weights='w')
#            vypr_flag
# Gymnázium   0.507897
# Jiné        0.441986
# Není SŠ     0.661740
# SOŠ         0.557848

ff21['ss_typ'].value_counts()
# SOŠ          27998
# Gymnázium    20633
# Není SŠ       1571
# Jiné           273

% % stata
encode
ss_typ, gen(ss_typc)
logit
vypr_flag
i.ss_typc
i.aki4c[pw = w]

% stata
logit
vypr_flag
i.ss_typc
i.aki4c[pw = w], baselevels

logistic(-.1424101)  # SOŠ mají pravděpodobnost přijetí cca o 3.5 % nižší než gymnázia
# "Není SŠ" vyšší, "Jiné" zřetelně nižší -> ale měli bychom mít jen maturanty, tedy není jisté, co to vlastně znamená...

# Zajímavé exploratorní cvičení, není jisté, kam s tím dál...

# endregion

# APLIKAČNÍ AMBICE
font_size()
plt.rcParams['figure.figsize'] = 10, 6

isc21.show('isc21')
df21.show('df21')

isc21['podil_prijat'] = isc21['prijat'] / isc21['total']

df21.shape
df21 = pd.merge(df21, isc21[['aki4', 'podil_prijat']])  # ztratil jsem 7 studentů, u kterých neidentifikujeme obor

df21.groupby('ss_typ')[['podil_prijat', 'w']].apply(nanaverage, weights='w')
#            podil_prijat
# Gymnázium      0.455536
# Jiné           0.496687
# Není SŠ        0.525337
# SOŠ            0.526474

df21.groupby('ss_kraj')[['podil_prijat', 'w']].apply(nanaverage, weights='w')
# Rozdíly zanedbatelné...
#
# Hlavní město Praha        0.503435
# Jihomoravský kraj         0.487093
# Jihočeský kraj            0.507915
# Karlovarský kraj          0.488719
# Kraj Vysočina             0.496976
# Královéhradecký kraj      0.495167
# Liberecký kraj            0.490569
# Moravskoslezský kraj      0.492409
# Olomoucký kraj            0.494605
# Pardubický kraj           0.498985
# Plzeňský kraj             0.499008
# Středočeský kraj          0.501238
# Zlínský kraj              0.499091
# Ústecký kraj              0.500719

ss21 = df21.groupby('ss_izo')[['podil_prijat', 'w']].apply(nanaverage, weights='w').reset_index()
ssnopf21 = df21[~df21['pedf']].groupby('ss_izo')[['podil_prijat', 'w']].apply(nanaverage, weights='w').reset_index()
ssnopfi21 = df21[~df21['pedf_isc']].groupby('ss_izo')[['podil_prijat', 'w']].apply(nanaverage,
                                                                                   weights='w').reset_index()
ssnopf21 = ssnopf21.rename(columns={'podil_prijat': 'podil_pr_mimo_pedf'})
ssnopfi21 = ssnopfi21.rename(columns={'podil_prijat': 'podil_pr_mimo_pedfi'})
ss21.show()

sstyp21 = df21[['ss_izo', 'ss_typ']].value_counts().reset_index()
sstyp21 = sstyp21.sort_values(0, ascending=False).groupby('ss_izo')['ss_typ'].first().reset_index()

ss.groupby(['ss_izo', 'id'])

sns.histplot(ss21['podil_prijat']).show()

sspedf21 = df21.groupby(['id', 'ss_izo'])['pedf'].any().reset_index()
sspedfi21 = df21.groupby(['id', 'ss_izo'])['pedf_isc'].any().reset_index()
sszaci21 = sspedf21['ss_izo'].value_counts().rename('ss_zaci').reset_index().rename(columns={'index': 'ss_izo'})

ss21 = pd.merge(ss21, sspedf21.groupby('ss_izo')['pedf'].mean().reset_index())
ss21 = pd.merge(ss21, sszaci21)
ss21 = pd.merge(ss21, ssnopf21)
ss21 = pd.merge(ss21, sstyp21)

ss21 = pd.merge(ss21, sspedfi21.groupby('ss_izo')['pedf_isc'].mean().reset_index())
ss21 = pd.merge(ss21, ssnopfi21)

sns.histplot(ss21['ss_zaci']).show()

plt.figure()
ax1 = sns.scatterplot(data=ss21, x='podil_prijat', y='pedf', size='ss_zaci', sizes=(2, 300), alpha=0.4, hue='ss_typ')
ax1.set_title('Aplikační ambice SŠ a podíl PedF uchazečů')

plt.figure()
ax2 = sns.scatterplot(data=ss21, x='podil_pr_mimo_pedf', y='pedf', size='ss_zaci', sizes=(2, 300), alpha=0.4,
                      hue='ss_typ')
ax2.set_title('Aplikační ambice SŠ (bez PedF) a podíl PedF uchazečů')

Chart([ax1, ax2]).show()

sns.scatterplot(data=ss21[ss21['ss_typ'] != 'Gymnázium'], x='podil_pr_mimo_pedf', y='pedf', size='ss_zaci',
                sizes=(2, 300), alpha=0.4, hue='ss_typ').show()

df21['pedf_isc'] = df21['aki2'] == '01'
df21[['pedf', 'pedf_isc']].value_counts().unstack()

sns.histplot(ss21['ss_zaci']).show()

plt.figure()
ax1 = sns.scatterplot(data=ss21, x='podil_prijat', y='pedf_isc', size='ss_zaci', sizes=(2, 300), alpha=0.4,
                      hue='ss_typ')
ax1.set_title('Aplikační ambice SŠ a podíl Ped-ISCED uchazečů')

plt.figure()
ax2 = sns.scatterplot(data=ss21, x='podil_pr_mimo_pedfi', y='pedf_isc', size='ss_zaci', sizes=(2, 300), alpha=0.4,
                      hue='ss_typ')
ax2.set_title('Aplikační ambice SŠ (bez Ped-ISCED) a podíl Ped-ISCED uchazečů')

Chart([ax1, ax2]).show()

df21[df21['aki2'] == '01'].show()

isc21

li = []
for y in range(1999, 2023):
    pdf = pd.read_csv(f'{uchazec.data_root}/uchazec/stredni-skoly/m{y}.csv', encoding='cp1250')
    pdf['year'] = y
    li.append(pdf)
ss = pd.concat(li, axis=0, ignore_index=True).drop_duplicates().reset_index(drop=True)
ss = ss.sort_values('year', ascending=False).drop_duplicates('IZO').sort_index()
ss['OBOR'] = ss['OBOR'].str.strip()
ss.show('ss')

ss = ss[['IZO', 'ZAR_PLN', 'MISTO']].rename(columns={'IZO': 'ss_izo'}).drop_duplicates()
ss21 = pd.merge(ss21, ss, how='left')

ss21[(ss21['ss_zaci'] > 50) & (ss21['pedf'] > 0.6)].show()

ss21[(ss21['ss_zaci'] > 50) & (ss21['pedf'] < 0.4) & (ss21['ss_typ'] == 'SOŠ') & (
            ss21['podil_pr_mimo_pedf'] < 0.45)].show()

pd.isna(df21['aki4']).sum()

% stata
list
aki4 in 1 / 10

% % stata
encode
aki4, gen(aki4c)
encode
gender, gen(genderc)
tab
genderc

% stata
logit
vypr_flag
genderc
i.aki4c[pw = w]
er = stata.get_ereturn()

sigmoid(er['e(b)'][0, -1])
sigmoid(er['e(b)'][0, 0] + er['e(b)'][0, -1])

er

er['e(b)'].shape

er['e(b)'][0, 0]
er['e(b)'][0, -1]

rr = stata.get_return()
type(rr['r(PT_rtitles)'])

sr = stata.get_sreturn()

% stata
help
xtlogit

x = np.linspace(-10, 10, 1000)


def sigmoid(x):
    return 1 / (1 + np.exp(-x))


sns.lineplot(x=x, y=sigmoid(x)).show()

sigmoid(0.5190593)  # 0.627
sigmoid(0.5190593 - 0.6199063)  # 0.475

sigmoid(-0.676048)  # 0.337
sigmoid(-0.676048 - 0.0660265)  # 0.323

stata.pdataframe_to_data(df21)
stata.run('encode aki4, gen(aki4c)')
stata.run('encode gender, gen(genderc)')
stata.run('list aki4 aki4c in 1/10')
stata.run('logit vypr_flag genderc i.aki4c [pw=w]')

config.status()

import sys

sys.path.append('C:\\Program Files\\Stata17\\utilities')
from pystata import config

config.init('mp')
config.status()
from pystata import stata

s = stata.run

s('list aki4 aki4c in 1/10')
s('help logit')

df21['gender'].dtype

stata.run('d gender')

df21.show('df21')
ff21.show('ff21')
isc21.show('isc21')

df21.to_stata(f'{uchazec.data_root}/uchazec/df21.dta', write_index=False, version=118,
              variable_labels=uchazec.variable_labels)

font_size()

figs17 = uchazec.plot_isced(isc17)
figs21 = uchazec.plot_isced(isc21)

Selector([Selector(figs21, title='Rok 2021'), Selector(figs17, title='Rok 2017')],
         title='Podíl přijatých a zapsaných podle let a oborů ISCED').show()

# ok, this is pretty simple and seems to work well
# next steps?
# -> novy cistsi soubor?

isc21['selektivita'] = isc21['prijat'] / isc21['total']

df21 = pd.merge(df21, isc21[['aki4', 'selektivita']])

df21.shape

rep21

df21['selektivita'].mean()
np.average(df21['selektivita'], weights=df21['w'])

df21.groupby('gender')[['selektivita', 'vypr_flag']].mean()

df21.groupby('pedf')[['selektivita']].mean()

np.sum(df21['pedf'])

np.mean(df21['ss_izo'].value_counts())

df21[df21.pedf].show()

# teď už bych měl mít komplet ISCED-F v datasetech
df21.show('df21')

i2_zaps = df21[['isced2', 'zaps']].value_counts().unstack()
i2_counts = df21['isced2'].value_counts()

i2_zaps = i2_zaps.sort_index()
i2_counts = i2_counts.sort_index()
i2_zaps['total'] = i2_counts

i2_zaps = df21[['isced2', 'zaps']].value_counts().unstack()
i2_zaps['total'] = i2_zaps.sum(axis=1)
i2_zaps.T / i2_zaps['total']
np.round(100 * i2_zaps / i2_zaps.sum(axis=1)[:, np.newaxis], 1)

zaps = df21[['isced2', 'zaps']].value_counts().unstack()
zaps['total'] = zaps.sum(axis=1)
zaps['prijati'] = zaps['Uchazeč se nezapsal'] + zaps['Uchazeč se zapsal']
zaps['prijati_pct'] = np.round(100 * zaps['prijati'] / zaps['total'], 1)
zaps['zapsani_pct'] = np.round(100 * zaps['Uchazeč se zapsal'] / zaps['prijati'], 1)
zaps.show('zaps21_2')

zaps = df17[['isced2', 'zaps']].value_counts().unstack()
zaps['total'] = zaps.sum(axis=1)
zaps['prijati'] = zaps['Uchazeč se nezapsal'] + zaps['Uchazeč se zapsal']
zaps['prijati_pct'] = np.round(100 * zaps['prijati'] / zaps['total'], 1)
zaps['zapsani_pct'] = np.round(100 * zaps['Uchazeč se zapsal'] / zaps['prijati'], 1)
zaps.show('zaps17_2')

zaps = df21[['isced3', 'aki3', 'isced2', 'aki2', 'zaps']].value_counts().unstack()
zaps['total'] = zaps.sum(axis=1)
zaps['prijati'] = zaps['Uchazeč se nezapsal'] + zaps['Uchazeč se zapsal']
zaps['prijati_pct'] = np.round(100 * zaps['prijati'] / zaps['total'], 1)
zaps['zapsani_pct'] = np.round(100 * zaps['Uchazeč se zapsal'] / zaps['prijati'], 1)
zaps = zaps.drop(columns=['Uchazeč se nezapsal', 'Uchazeč se zapsal', 'Uchazeč nebyl přijat'])
zaps.show('zaps21_3')

zaps = df17[['isced4', 'aki4', 'isced3', 'aki3', 'isced2', 'aki2', 'zaps']].value_counts().unstack()
zaps['total'] = zaps.sum(axis=1)
zaps['prijati'] = zaps['Uchazeč se nezapsal'] + zaps['Uchazeč se zapsal']
zaps['prijati_pct'] = np.round(100 * zaps['prijati'] / zaps['total'], 1)
zaps['zapsani_pct'] = np.round(100 * zaps['Uchazeč se zapsal'] / zaps['prijati'], 1)
zaps = zaps.drop(columns=['Uchazeč se nezapsal', 'Uchazeč se zapsal', 'Uchazeč nebyl přijat'])
# zaps.show('zaps17_4')
zaps = zaps.reset_index()

zaps = df21[['isced4', 'aki4', 'isced3', 'aki3', 'isced2', 'aki2', 'zaps']].value_counts().unstack()
zaps['total'] = zaps.sum(axis=1)
zaps['prijati'] = zaps['Uchazeč se nezapsal'] + zaps['Uchazeč se zapsal']
zaps['prijati_pct'] = np.round(100 * zaps['prijati'] / zaps['total'], 1)
zaps['zapsani_pct'] = np.round(100 * zaps['Uchazeč se zapsal'] / zaps['prijati'], 1)
zaps = zaps.drop(columns=['Uchazeč se nezapsal', 'Uchazeč se zapsal', 'Uchazeč nebyl přijat'])
# zaps.show('zaps21_4')
zaps = zaps.reset_index()

i2s = zaps['isced2'].drop_duplicates().sort_values().values
colors = sns.color_palette(n_colors=10)
colors = {i: c for i, c in zip(i2s, colors)}
norm = plt.Normalize(zaps['total'].min(), zaps['total'].max())

font_size()

len(isc21['isced2'].unique())

xmin, xmax = 0, 110
ymin, ymax = 32, 108

figs21 = []

fig, ax = plt.subplots(figsize=(16, 9))

for i2, idf in zaps.groupby('isced2'):
    sns.scatterplot(x='prijati_pct', y='zapsani_pct', data=idf, label=i2, color=colors[i2], size='total',
                    sizes=(10, 800), size_norm=norm, legend=False, ec=None)
    # labels
    for _, row in idf.iterrows():
        plt.text(row['prijati_pct'], row['zapsani_pct'] + 0.2 + np.sqrt(row['total']) / 80, row['isced4'],
                 color=colors[i2], ha='center', va='bottom')

# xmin, xmax = ax.get_xlim()
# ymin, ymax = ax.get_ylim()
ax.set(xlim=(xmin, xmax), ylim=(ymin, ymax))

ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.yaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.set(xlabel='Podíl přijatých', ylabel='Podíl zapsaných', title='Podíl přijatých a zapsaných podle oborů, 2021')
plt.legend()

figs21.append(Chart(fig, title='Všechny obory'))

for i2, idf in zaps.groupby('isced2'):
    fig, ax = plt.subplots(figsize=(16, 9))
    sns.scatterplot(x='prijati_pct', y='zapsani_pct', data=idf, label=i2, color=colors[i2], size='total',
                    sizes=(10, 800), size_norm=norm, legend=False, ec=None)
    # labels
    for _, row in idf.iterrows():
        plt.text(row['prijati_pct'], row['zapsani_pct'] + 0.2 + np.sqrt(row['total']) / 80, row['isced4'],
                 color=colors[i2], ha='center', va='bottom')

    ax.set(xlim=(xmin, xmax), ylim=(ymin, ymax))

    ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
    ax.yaxis.set_major_formatter(mpl.ticker.PercentFormatter())
    ax.set(xlabel='Podíl přijatých', ylabel='Podíl zapsaných', title=i2)
    plt.legend()

    figs21.append(Chart(fig, title=i2))

sel21 = Selector(figs21, title='Rok 2021')

zaps = df17[['isced4', 'aki4', 'isced3', 'aki3', 'isced2', 'aki2', 'zaps']].value_counts().unstack()
zaps['total'] = zaps.sum(axis=1)
zaps['prijati'] = zaps['Uchazeč se nezapsal'] + zaps['Uchazeč se zapsal']
zaps['prijati_pct'] = np.round(100 * zaps['prijati'] / zaps['total'], 1)
zaps['zapsani_pct'] = np.round(100 * zaps['Uchazeč se zapsal'] / zaps['prijati'], 1)
zaps = zaps.drop(columns=['Uchazeč se nezapsal', 'Uchazeč se zapsal', 'Uchazeč nebyl přijat'])
# zaps.show('zaps17_4')
zaps = zaps.reset_index()

figs17 = []

fig, ax = plt.subplots(figsize=(16, 9))

for i2, idf in zaps.groupby('isced2'):
    sns.scatterplot(x='prijati_pct', y='zapsani_pct', data=idf, label=i2, color=colors[i2], size='total',
                    sizes=(10, 800), size_norm=norm, legend=False, ec=None)
    # labels
    for _, row in idf.iterrows():
        plt.text(row['prijati_pct'], row['zapsani_pct'] + 0.2 + np.sqrt(row['total']) / 80, row['isced4'],
                 color=colors[i2], ha='center', va='bottom')

ax.set(xlim=(xmin, xmax), ylim=(ymin, ymax))

ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.yaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.set(xlabel='Podíl přijatých', ylabel='Podíl zapsaných', title='Podíl přijatých a zapsaných podle oborů, 2021')
plt.legend()

figs17.append(Chart(fig, title='Všechny obory'))

for i2, idf in zaps.groupby('isced2'):
    fig, ax = plt.subplots(figsize=(16, 9))
    sns.scatterplot(x='prijati_pct', y='zapsani_pct', data=idf, label=i2, color=colors[i2], size='total',
                    sizes=(10, 800), size_norm=norm, legend=False, ec=None)
    # labels
    for _, row in idf.iterrows():
        plt.text(row['prijati_pct'], row['zapsani_pct'] + 0.2 + np.sqrt(row['total']) / 80, row['isced4'],
                 color=colors[i2], ha='center', va='bottom')

    ax.set(xlim=(xmin, xmax), ylim=(ymin, ymax))

    ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
    ax.yaxis.set_major_formatter(mpl.ticker.PercentFormatter())
    ax.set(xlabel='Podíl přijatých', ylabel='Podíl zapsaných', title=i2)
    plt.legend()

    figs17.append(Chart(fig, title=i2))

sel17 = Selector(figs17, title='Rok 2017')

Selector([sel21, sel17], title='Podíl přijatých a zapsaných podle let a oborů ISCED').show()

ax.show()

i2

# příprava kategorií
df21.head(n=1000).show('df21')
df17.head(n=1000).show('df17')
df21[['fak_nazev', 'vs_nazev']].drop_duplicates().reset_index(drop=True).show()

# akko, akvo, program, regpro?
# program
sys_root = 'D:\\' if sys.platform == 'win32' else '/mnt/d'
data_root = os.path.join(sys_root, 'projects', 'idea', 'data')

# převod SP <--> ISCED
sp_isced = pd.read_excel(f'{data_root}/uchazec/Prevodnik_ISCED.xlsx')
sp_isced['isced'] = sp_isced['isced'].astype('str').str.pad(width=4, side='left', fillchar='0')
sp_isced.show('sp_isced', num_rows=5000)
sp_isced_dict = sp_isced.set_index('kód SP')['isced'].to_dict()

# ISCED-F
aki = {}
aki_xml = {
    2: 'http://stistko.uiv.cz/katalog/textdata/C131443AKI2.xml',
    3: 'http://stistko.uiv.cz/katalog/textdata/C13149AKI3.xml',
    4: 'http://stistko.uiv.cz/katalog/textdata/C13156AKI4.xml'
}

for i in range(2, 5):
    foo = pd.read_xml(aki_xml[i], encoding='cp1250', xpath='./veta')
    foo['KOD'] = foo['KOD'].astype('str').str.pad(width=i, side='left', fillchar='0')
    foo = foo[['KOD', 'TXT']].rename(columns={'KOD': f'aki{i}', 'TXT': f'isced{i}'})
    aki[i] = foo

df21['aki4'] = np.where(df21['program'].str.len() == 5, df21['program'].map(sp_isced_dict), df21['program'].str[1:5])
df21['aki3'] = df21['aki4'].str[:3]
df21['aki2'] = df21['aki4'].str[:2]

df21 = pd.merge(df21, aki[4], how='left')
df21 = pd.merge(df21, aki[3], how='left')
df21 = pd.merge(df21, aki[2], how='left')

df17['aki4'] = np.where(df17['program'].str.len() == 5, df17['program'].map(sp_isced_dict), df17['program'].str[1:5])
df17['aki3'] = df17['aki4'].str[:3]
df17['aki2'] = df17['aki4'].str[:2]

df17 = pd.merge(df17, aki[4], how='left')
df17 = pd.merge(df17, aki[3], how='left')
df17 = pd.merge(df17, aki[2], how='left')

aki[2].show('aki2')
aki[3].show('aki3')
aki[4].show('aki4')

df21[['vs_nazev', 'fak_nazev', 'program', 'aki2', 'aki3', 'aki4', 'isced2', 'isced3', 'isced4']].head(1000).show(
    'df21_isced')

df17[['vs_nazev', 'fak_nazev', 'program', 'aki2', 'aki3', 'aki4', 'isced2', 'isced3', 'isced4']].head(1000).show(
    'df17_isced')

df21 = df21.drop(columns=['aki2', 'aki3', 'aki4', 'isced2', 'isced3', 'isced4'])

aki2 = pd.read_xml(aki2_xml, encoding='cp1250', xpath='./veta')
aki3 = pd.read_xml(aki3_xml, encoding='cp1250', xpath='./veta')
aki4 = pd.read_xml(aki4_xml, encoding='cp1250', xpath='./veta')
aki2['KOD'] = aki2['KOD'].astype('str').str.pad(width=2, side='left', fillchar='0')
aki3['KOD'] = aki3['KOD'].astype('str').str.pad(width=3, side='left', fillchar='0')
aki4['KOD'] = aki4['KOD'].astype('str').str.pad(width=4, side='left', fillchar='0')

aki2.show('aki2')
aki3.show('aki3')
aki4.show('aki4')

df21['aki4'] = df21['program'].str[1:5]

df17['program'].str.len().value_counts()
df17.shape

np.sum(sp_isced['kód SP'].str.len() != 5)

program = pd.read_csv(f'{data_root}/uchazec/uch3-03-6/cisel/program.csv', encoding='cp1250')
program.show('program')

akko_xml = 'http://stistko.uiv.cz/katalog/textdata/C95743AKKO.xml'
akko = pd.read_xml(akko_xml, encoding='cp1250', xpath='./veta')
akko.show('akko')

akko[akko['KOD'].str.len() == 1].show('akko1')
akko[akko['KOD'].str.len() == 2].show('akko2')

akvo_xml = 'http://stistko.uiv.cz/katalog/textdata/C95619AKVO.xml'
akvo = pd.read_xml(akvo_xml, encoding='cp1250', xpath='./veta')
akvo.show('akvo')

regpro = pd.read_csv(f'{data_root}/uchazec/uch3-03-6/cisel/regpro.csv', encoding='cp1250')
regpro.show('regpro')

df17['obor1']

akko1 = akko[akko['KOD'].str.len() == 1][['KOD', 'TXT']].rename(columns={'KOD': 'obor11', 'TXT': 'akko1'})
akko1['akko1'] = akko1['akko1'].str.lower()

foo = df17[['vs_nazev', 'fak_nazev', 'obor1']].value_counts().rename('prihlasky').reset_index()
foo['obor11'] = foo['obor1'].str[0]
foo = pd.merge(foo, akko1)
foo['pocet'] = 1
foo.head()

foo
np.sum(foo.groupby(['vs_nazev', 'fak_nazev', 'akko1'])['pocet'].sum() > 0)

foo.groupby(['vs_nazev', 'fak_nazev', 'akko1']).sum()
bar = foo.groupby(['vs_nazev', 'fak_nazev', 'akko1'])[['prihlasky', 'pocet']].sum().reset_index()
bar = bar[bar['pocet'] > 0].copy()
bar['minus_pocet'] = -bar['pocet']
bar['minus_prihlasky'] = -bar['prihlasky']
bar.sort_values(['vs_nazev', 'fak_nazev', 'minus_pocet', 'minus_prihlasky']).show('bar')

# endregion

###  SOME DEBUGGING  ###
# region

# můžu nahrát ten kompletní číselník?
# http://stistko.uiv.cz/katalog/ciselnika.asp -> MCPR
url_old = 'http://stistko.uiv.cz/katalog/ciselnik11x.asp?idc=MCPR&ciselnik=V%FDsledek+p%F8ij%EDmac%EDho+%F8%EDzen%ED&aap=on&poznamka='
html = urlopen(url_old).read()
cleanr = re.compile('<.*?>')
lines = [l.strip() for l in re.sub(cleanr, '', html.decode('windows-1250')).split('\r\n') if l.count(';') > 4]
text_data = StringIO('\n'.join(lines))
rg = pd.read_csv(text_data, sep=';', index_col=False)
rg.columns = [c.upper() for c in rg.columns]
rg = rg.groupby('KOD').last().reset_index()

# proč je tak často chybějící výsledek přijímacího řízení v datech z roku 2017
# ok, asi změna v kódování v číselníku

year = 17
sys_root = 'D:\\' if sys.platform == 'win32' else '/mnt/d'
data_root = os.path.join(sys_root, 'projects', 'idea', 'data')
path = os.path.join(data_root, 'uchazec', f'0022MUCH{year}P')

foo = pd.read_csv(f'{path}.csv', encoding='cp1250', low_memory=False)
foo['VYPR']
foo.head(1000).show()

np.round(100 * foo['VYPR'].value_counts() / foo.shape[0], 2)

np.round(100 * foo['ZAPS'].value_counts() / foo.shape[0], 2)

url = 'http://stistko.uiv.cz/katalog/textdata/C21925MCPR.xml'

rg = pd.read_xml(url, encoding='cp1250', xpath='./veta')
rg

rg = pd.merge(rg, pd.DataFrame({'KOD': range(10, 20)}), how='outer')
rg['TXT'] = rg['TXT'].fillna('Přijat - pravděpodobně, chybí v číselníku')
rg = pd.merge(rg, pd.DataFrame({'KOD': range(20, 30)}), how='outer')
rg['TXT'] = rg['TXT'].fillna('Nepřijat - pravděpodobně, chybí v číselníku')
rg = pd.merge(rg, pd.DataFrame({'KOD': range(30, 40)}), how='outer')
rg['TXT'] = rg['TXT'].fillna('Přijímací řízení zrušeno - pravděpodobně, chybí v číselníku')

rg['IDX'] = rg.index
df = pd.merge(df.rename(columns={c: 'KOD'}), rg[['KOD', 'IDX']].rename(columns={'IDX': c}),
              how='left').drop(columns=['KOD'])
rg_dict = rg[['IDX', 'TXT']].set_index('IDX')['TXT'].to_dict()
value_labels[c] = rg_dict

###  OLD AND UNUSED MESS  ###
# region
foo21 = pedf_app(ff21, 'ss_typ_g')
foo21

foo21 = pedf_app(ff21, 'ss_kraj')
ff21.columns

ff21['ss_typ'].value_counts()
ff21[['ss_typ', 'rmat']].value_counts()

ff21[ff21['ss_typ'] == 'Gymnázium']['ss_gym_delka'].value_counts()

foo21

foo17 = pedf_app(ff17, 'ss_kraj')[['ss_kraj', 'pedf_rel']].rename(columns={'pedf_rel': 2017})
foo21 = pedf_app(ff21, 'ss_kraj')[['ss_kraj', 'pedf_rel']].rename(columns={'pedf_rel': 2021})
foo = pd.merge(foo21, foo17)

fig, ax = plt.subplots()
sns.barplot(foo)
ax.show()
foo.plot().show()

doc.add_paragraph('')

fig, ax = plt.subplots()
ax = sns.histplot(data=ff, y='prihl', discrete=True, binwidth=0.8, shrink=0.8, ec=None)
ax.set(xlabel='Četnost', ylabel='Počet přihlášek', title='Počet přihlášek jednotlivých uchazečů')
ax.set_ylim(0, 13)
ax.invert_yaxis()
fig.tight_layout()

figpath = savefig(fig, f'cetnost_prihlasek_{y}')

doc.add_picture(figpath, width=CHART_WIDTH)
doc.add_paragraph('')

doc.show()

latent_styles = doc.styles.latent_styles
len(latent_styles)

table_style = [ls for ls in latent_styles if ls.name == 'Grid Table Light'][0]
doc.styles.add_style('Table Grid Light', docx.enum.style.WD_STYLE_TYPE.TABLE)

for ls in latent_styles:
    print(ls.name)

table = doc.add_table(rows=1, cols=3, style='Plain Table 1')
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Kraj'
hdr_cells[1].text = 'Přihlášek na PedF'
hdr_cells[2].text = 'Podíl na PedF'

for _, row in foo.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row['ss_kraj']
    row_cells[1].text = f'{row.prihl_bool:,}'
    row_cells[2].text = f'{row.pedf_rel} %'

doc.show()

doc.add_heading('Úspěšnost uchazečů o pedagogické fakulty')

for y in [21, 17]:
    ff = eval(f'ff{y}')

    tot_uchazecu = ff['prihl_pedf'].shape[0]
    tot_chlapcu = (ff['gender'] == 'Chlapec').sum()
    tot_divek = (ff['gender'] == 'Dívka').sum()

    doc.add_heading(f'20{y}:', 2)

pedf_app(ff21, 'gender')

pedf_app(ff21, 'ss_kraj')

doc.show()

df21[df21['pedf']][['fak_nazev', 'vs_nazev']].drop_duplicates().sort_values('vs_nazev')

sns.histplot(df21['pocet'], discrete=True, binwidth=0.8, ec=None).show()

doc.add_heading('2017:', 2)
doc.add_paragraph('')

doc.show()

os.getcwd()

font_size_word()
sns.histplot(data=ff21, y='prihl', hue='gender', multiple='stack', discrete=True, binwidth=0.8, ec=None).show()

ax = sns.histplot(data=ff21, y='prihl', discrete=True, binwidth=0.8, shrink=0.8, ec=None)
ax.set(xlabel='Četnost', ylabel='Počet přihlášek', title='Počet přihlášek jednotlivých uchazečů')
ax.set_ylim(0, 13)
ax.invert_yaxis()
ax.get_figure().tight_layout()

figpath = savefig(ax, 'cetnost_prihlasek_21')

doc.add_picture(figpath, width=Mm(140))
doc.show()

ax.get_figure().savefig

ax.get_figure()

ax.show()

ff21['prihl'].value_counts()

plt.rcParams['figure.titlesize'] = 30
ax.show()

eval(f'df{17}')

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

p = doc.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

doc.add_heading('Heading, level 1', level=1)

doc.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
doc.add_paragraph(
    'first item in ordered list', style='List Number'
)

for style in doc.styles:
    print("style.name == %s" % style.name)

### Learning how to use docx...


df.show()

document = Document()
document.add_heading('Document title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save(f'{foo}/demo.docx')

foo = create_stamped_temp('docx')
foo

os.startfile(f'{foo}/demo.docx')

# how to set a style?


doc = Document()
styles = doc.styles

style = styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(14)

doc.add_heading('Document title', 0)

doc.show()
docx_document_show(doc)
doc.save(tempfile.gettempdir() + '/foo.docx')

p = doc.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

doc.add_heading('Heading, level 1', level=1)

doc.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
doc.add_paragraph(
    'first item in ordered list', style='List Number'
)

doc.show()

foo = create_stamped_temp('docx')
doc.save(f'{foo}/demo.docx')
os.startfile(f'{foo}/demo.docx')

doc = Document('../IDEA_SABLONA.docx')
doc.show()

type(doc)

doc = Document()

doc = Document('empty.docx')
doc._body.clear_content()

doc.add_heading('Document title', 0)

p = doc.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

doc.add_heading('Heading, level 1', level=1)

doc.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
doc.add_paragraph(
    'first item in ordered list', style='List Number'
)

stamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
path = tempfile.gettempdir() + '/docs'
os.makedirs(path, exist_ok=True)
path = '{}/{}.{}'.format(path, stamp, format)

path

self.save(path, **kwargs)

doc.element.xml

doc.show()
doc

docx_document_show(doc)

os.environ['TEMP'] = 'D:/temp'

tempfile.tempdir = 'D:/temp'
tempfile.gettempdir()

Document.show = docx_document_show
doc.show()

Document
Document()

# endregion
# endregion

# endregion


# CZ
R_cz = 0.2928 * 28443
P_cz = R_cz / 5.2
R_cz - P_cz

R_co = 0.5562 * 13265
P_co = R_co / 60.4
R_co - P_co

# COLOMBIA
