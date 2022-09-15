###  IMPORTS  ###
#region
import os
import sys
from venv import create
import pyreadstat
import pandas as pd
import numpy as np
import re
from statsmodels.stats.weightstats import DescrStatsW
from sklearn.linear_model import LinearRegression
import statsmodels.api as sm
import statsmodels.formula.api as smf
import scipy
import seaborn as sns
import matplotlib as mpl
import matplotlib.pyplot as plt
plt.rcParams['figure.figsize'] = 12, 6
import dbf

from docx import Document
from docx.shared import Mm, Pt

from libs.utils import *
from libs.plots import *
from libs.extensions import *
from libs import uchazec

import importlib

# importlib.reload(uchazec)
# importlib.invalidate_caches()
# import locale
# locale.setlocale(locale.LC_ALL, 'cs_CS')
#endregion

###  SETUP THE ENVIRONMENT  ###
#region
# font sizes for importing to word
def font_size(small=9, medium=11, big='large'):
    plt.rc('font', size=small)          # controls default text sizes
    plt.rc('axes', titlesize=big)     # fontsize of the axes title
    plt.rc('axes', labelsize=medium)    # fontsize of the x and y labels
    plt.rc('xtick', labelsize=small)    # fontsize of the tick labels
    plt.rc('ytick', labelsize=small)    # fontsize of the tick labels
    plt.rc('legend', fontsize=small)    # legend fontsize
    plt.rc('figure', titlesize=big)  # fontsize of the figure title    

def font_size_word():
    font_size(small=12, medium=14, big=18)
    
def savefig(ax_or_fig, title):
    fig_path = f'D:\\projects\\idea\\code\\output\\uchazec\\{title}.png'
    fig = ax_or_fig if isinstance(ax_or_fig, plt.Figure) else ax_or_fig.get_figure()
    fig.savefig(fig_path)
    return fig_path

font_size_word()
CHART_WIDTH = Mm(158)
# plt.style.use('dark_background')
#endregion

###  LOAD AND PREPARE DATA  ###
#region
df17 = pd.read_parquet('temp/uchazec/uch17.parquet')
df21 = pd.read_parquet('temp/uchazec/uch21.parquet')

df17, rep17 = uchazec.filter_data(df17)
df21, rep21 = uchazec.filter_data(df21)

df17 = uchazec.add_variables(df17)
df21 = uchazec.add_variables(df21)

ff17 = uchazec.get_per_app(df17)
ff21 = uchazec.get_per_app(df21)

def add_ss_typ_g(ff):
    ff['ss_typ_g'] = np.where(ff['ss_typ'].isin(['SOŠ', 'Není SŠ', 'Jiné']), ff['ss_typ'], np.where(ff['ss_typ'] == 'Gymnázium', np.where(ff['ss_gym_delka'] == 4, 'Gym 4-leté', np.where(ff['ss_gym_delka'] == 6, 'Gym 6-leté', np.where(ff['ss_gym_delka'] == 8, 'Gym 8-leté', 'Jiné'))), 'Jiné'))
    ss_typ_g_cat = ['Gym 4-leté', 'Gym 6-leté', 'Gym 8-leté', 'SOŠ', 'Není SŠ', 'Jiné']
    ff['ss_typ_g'] = pd.Categorical(ff['ss_typ_g'], categories=ss_typ_g_cat, ordered=True)
    return ff

df21 = add_ss_typ_g(df21)
df17 = add_ss_typ_g(df17)
ff21 = add_ss_typ_g(ff21)
ff17 = add_ss_typ_g(ff17)

def pedf_app(ff, c):
    foo = ff.groupby(c)[['prihl_pedf_bool', 'prihl_bool']].sum().reset_index()
    foo['pedf_rel'] = np.round(100 * foo['prihl_pedf_bool'] / foo['prihl_bool'], 1)
    return foo

pedf_vs = df21[df21['pedf']][['fak_nazev', 'vs_nazev']].drop_duplicates().sort_values('vs_nazev').reset_index(drop=False)

# df = df21
# df.head(n=1000).show()
#endregion

###  GENERATE REPORT  ###
#region

# EMPTY DOC
#region
doc = Document('template.docx')
doc._body.clear_content()

section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(25.4)
section.right_margin = Mm(25.4)
section.top_margin = Mm(25.4)
section.bottom_margin = Mm(25.4)

doc.add_heading('IDEA: Uchazeč', 0)
#endregion

# 1. PŘÍPRAVA DAT
#region
doc.add_heading('1. Příprava dat', 1)
doc.add_paragraph('Analýza je založena na anonymizovaných datech z databáze Uchazeč, která obsahuje informaci o všech přihlášených na vysoké školy pro roky 2017 a 2021.')

for y in [17, 21]:
    doc.add_heading(f'20{y}:', 2)
    for r in eval(f'rep{y}'):
        doc.add_paragraph(r, style='Bullet List')

doc.add_paragraph('')
doc.add_paragraph('V datech jsou i nějaké další problémy: například stejný student může mít více přihlášek, které se v některých údajích o SŠ mohou lišit – odlišný typ SŠ (například SOŠ + Není SŠ, odlišná lokalita SŠ, případně chybějící údaje). Při agregaci na úroveň jednotlivých respondentů bereme v úvahu údaj z prvního záznamu o přihlášce uchazeče.')
#endregion

# 2. PŘIHLÁŠKY NA VYSOKÉ ŠKOLY
#region
doc.add_heading('2. Přihlášky na vysoké školy', 1)

for y in [17, 21]:
    ff = eval(f'ff{y}')
    tot_prihlasek = eval(f'df{y}').shape[0]
    tot_uchazecu = ff.shape[0]
    tot_chlapcu = (ff['gender'] == 'Chlapec').sum()
    tot_divek = (ff['gender'] == 'Dívka').sum()
    avg_prihlasek = ff['prihl'].mean()
    max_prihlasek = ff['prihl'].max()

    doc.add_heading(f'20{y}:', 2)
    doc.add_paragraph(f'V roce 20{y} bylo podáno {tot_prihlasek:,} přihlášek na vysoké školy od {tot_uchazecu:,} českých uchazečů, z toho se hlásilo {tot_chlapcu:,} chlapců ({100 * tot_chlapcu / tot_uchazecu:.3g} %) a {tot_divek:,} dívek ({100 * tot_divek / tot_uchazecu:.3g} %). Průměrný uchazeč si podal {avg_prihlasek:.1f} přihlášek, maximální počet podaných přihlášek byl {max_prihlasek:.0f}. Četnost počtu přihlášek jednotlivých uchazečů ukazuje následující graf.')
    doc.add_paragraph('')

    fig, ax = plt.subplots()
    ax = sns.histplot(data=ff, y='prihl', discrete=True, binwidth=0.8, shrink=0.8, ec=None)
    ax.set(xlabel='Četnost', ylabel='Počet přihlášek', title='Počet přihlášek jednotlivých uchazečů')
    ax.set_ylim(0, 13)
    ax.invert_yaxis()
    fig.tight_layout()

    figpath = savefig(fig, f'cetnost_prihlasek_{y}')

    doc.add_picture(figpath, width=CHART_WIDTH)
    doc.add_paragraph('')
#endregion

# 3. PROFIL UCHAZEČŮ O PEDAGOGICKÉ FAKULTY
#region
doc.add_heading('3. Profil uchazečů o pedagogické fakulty', 1)
doc.add_paragraph('Uchazeči se mohli hlásit na některou z osmi pedagogických fakult v České republice:')
doc.add_paragraph('')

for _, row in pedf_vs.iterrows():
    doc.add_paragraph(row['vs_nazev'], style='Bullet List')

for y in [17, 21]:
    ff = eval(f'ff{y}')

    tot_uchazecu = ff.shape[0]
    tot_chlapcu = (ff['gender'] == 'Chlapec').sum()
    tot_divek = (ff['gender'] == 'Dívka').sum()

    pedf_uchazecu = ff['prihl_pedf'].sum()
    pedf_chlapcu = ff[ff['gender'] == 'Chlapec']['prihl_pedf_bool'].sum()
    pedf_divek = ff[ff['gender'] == 'Dívka']['prihl_pedf_bool'].sum()

    doc.add_heading(f'20{y}:', 2)

    doc.add_paragraph(f'V roce 20{y} se na pedagogické fakulty hlásilo {pedf_uchazecu:,.0f} ({100 * pedf_uchazecu / tot_uchazecu:.3g} %). Z toho bylo {pedf_divek:,} dívek ({100 * pedf_divek / tot_divek:.3g} % všech dívek) a {pedf_chlapcu:,} chlapců ({100 * pedf_chlapcu / tot_chlapcu:.3g} % všech chlapců). Následující tabulka shrnuje, jaký podíl uchazečů se hlásil na pedagogické fakulty z jednotlivých krajů (procenta vyjadřují podíl ze všech uchazečů daného kraje).')
    doc.add_paragraph('')

    foo = pedf_app(ff, 'ss_kraj')
    table = doc.add_table(rows=1, cols=3, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kraj'
    hdr_cells[1].text = 'Uchazeči o PedF'
    hdr_cells[2].text = 'Podíl PedF'

    for _, row in foo.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['ss_kraj']
        row_cells[1].text = f'{row.prihl_pedf_bool:,}'
        row_cells[2].text = f'{row.pedf_rel} %'

    doc.add_paragraph('')
    p = doc.add_paragraph('Další tabulka ukazuje podobné srovnání podle typu střední školy.')
    if y == 17:
        p.add_run(' ("Není SŠ" zpravidla označuje uchazeče, které se nehlásí přímo ze střední školy, ale není tomu tak vždy – v některých případech byly takto označeni i aktuální maturanti, naopak u některých starších uchazečů byl typ střední školy vyplněný správně; malý počet uchazečů s jiným typem střední školy byl vynechán.)')
    if y == 21:
        p.add_run(' Celkově se počet absolventů gymnázií hlásících se na vysoké školy změnil jen mírně, oproti tomu počet absolventů středních odborných škol, kteří se hlásí na vysoké školy, znatelně narostl. Nejvíce studentů se na pedagogické fakulty hlásí ze středních odborných škol, v relativním vyjádření je to pak největší podíl z absolventů čtyřletých gymnázií.')
    doc.add_paragraph('')

    foo = pedf_app(ff, 'ss_typ_g')
    table = doc.add_table(rows=1, cols=4, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Typ školy'
    hdr_cells[1].text = 'Uchazeči o PedF'
    hdr_cells[2].text = 'Uchazeči celkem'
    hdr_cells[3].text = 'Podíl PedF'

    for _, row in foo[foo['ss_typ_g'] != 'Jiné'].iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['ss_typ_g']
        row_cells[1].text = f'{row.prihl_pedf_bool:,}'
        row_cells[2].text = f'{row.prihl_bool:,}'
        row_cells[3].text = f'{row.pedf_rel} %'


doc.add_heading('Srovnání vývoje', 2)
doc.add_paragraph('Následující graf srovnává podíl uchazečů o pedagogické fakulty z jednotlivých krajů v letech 2017 a 2021. Je patrné, že podíl zájemců o pedagogické fakulty se v roce 2021 výrazně zvýšil téměř ve všech krajích.')
doc.add_paragraph('')

foo17 = pedf_app(ff17, 'ss_kraj')[['ss_kraj', 'pedf_rel']]
foo17['year'] = '2017'
foo21 = pedf_app(ff21, 'ss_kraj')[['ss_kraj', 'pedf_rel']]
foo21['year'] = '2021'
foo = pd.concat([foo21, foo17], ignore_index=True)

fig, ax = plt.subplots()
sns.scatterplot(data=foo, y='ss_kraj', x='pedf_rel', hue='year', s=60)
ax.set(xlabel='Podíl uchazečů', ylabel=None, title='Podíl uchazečů o pedagogické fakulty')
plt.legend(title='Rok')
plt.plot(foo21.pedf_rel, foo21.ss_kraj)
plt.plot(foo17.pedf_rel, foo17.ss_kraj)
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'uchazeci_kraje_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_paragraph('')
doc.add_paragraph('Další graf srovnává vývoj podílu uchazečů mezi lety 2017 a 2021 podle typu střední školy. Je patrné, že nejvíce zájemců o střední školy přibylo mezi absolventy 4-letých gymnázií, případně také z víceletých gymnázií.')
doc.add_paragraph('')

foo17 = pedf_app(ff17, 'ss_typ_g')[['ss_typ_g', 'pedf_rel']]
foo17['year'] = '2017'
foo17 = foo17[foo17['ss_typ_g'] != 'Jiné']
foo21 = pedf_app(ff21, 'ss_typ_g')[['ss_typ_g', 'pedf_rel']]
foo21['year'] = '2021'
foo21 = foo21[foo21['ss_typ_g'] != 'Jiné']
foo = pd.concat([foo21, foo17], ignore_index=True)

fig, ax = plt.subplots()
sns.scatterplot(data=foo, y='ss_typ_g', x='pedf_rel', hue='year', s=60)
ax.set(xlabel='Podíl uchazečů', ylabel=None, title='Podíl uchazečů o pedagogické fakulty')
plt.legend(title='Rok')
plt.plot(foo21.pedf_rel, foo21.ss_typ_g)
plt.plot(foo17.pedf_rel, foo17.ss_typ_g)
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'uchazeci_typ_ss_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_heading('Podíl z přihlášek', 2)

avg_prihl_17 = ff17['prihl'].mean()
avg_prihl_pedf_17 = ff17['prihl_pedf'].mean()
avg_prihl_nepedf_17 = ff17['prihl_nepedf'].mean()
avg_prihl_21 = ff21['prihl'].mean()
avg_prihl_pedf_21 = ff21['prihl_pedf'].mean()
avg_prihl_nepedf_21 = ff21['prihl_nepedf'].mean()

doc.add_paragraph(f'Data ukazují, že na pedagogické fakulty se hlásí vyšší podíl uchazečů. Zároveň si v roce 2021 průměrný uchazeč podává více přihlášek: {avg_prihl_21:.2g} oproti {avg_prihl_17:.2g} přihlášce v roce 2017. Z toho si v roce 2017 podal průměrný uchazeč {avg_prihl_pedf_17:.2g} přihlášek na pedagogické fakulty a {avg_prihl_nepedf_17:.2g} na ostatní fakulty. V roce 2021 to bylo {avg_prihl_pedf_21:.2g} přihlášek na pedagogické fakulty (nárůst o {100 * avg_prihl_pedf_21 / avg_prihl_pedf_17 - 100:.3g} %) a {avg_prihl_nepedf_21:.2g} na ostatní fakulty (nárůst o {100 * avg_prihl_nepedf_21 / avg_prihl_nepedf_17 - 100:.3g} %).')

doc.add_paragraph(f'Pro ověření výsledků ukazují následující tabulky a grafy obdobná srovnání jako předchozí podíly podle krajů a typů střední školy, avšak vyjádřené jako podíl všech přihlášek, nikoli jako podíl uchazečů.')

df17.groupby('ss_kraj')[['ones', 'pedf']].sum().reset_index()


for y in [17, 21]:
    df = eval(f'df{y}')

    tot_prihlasek = df.shape[0]
    pedf_prihlasek = df['pedf'].sum()

    doc.add_heading(f'20{y}:', 3)

    doc.add_paragraph(f'V roce 20{y} bylo na pedagogické fakulty podáno {pedf_prihlasek:,.0f} z celkového počtu {tot_prihlasek:,.0f} přihlášek na vysoké školy od českých uchazečů (podíl pedagogických fakult tedy byl {100 * pedf_prihlasek / tot_prihlasek:.3g} %). Následující tabulka shrnuje, jaký byl podíl přihlášek na pedagogické fakulty ze všech přihlášek podaných z daného kraje.')
    doc.add_paragraph('')

    foo = df.groupby('ss_kraj')[['ones', 'pedf']].sum().reset_index()
    foo['pedf_rel'] = 100 * foo['pedf'] / foo['ones']

    table = doc.add_table(rows=1, cols=3, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kraj'
    hdr_cells[1].text = 'Přihlášky na PedF'
    hdr_cells[2].text = 'Podíl PedF'

    for _, row in foo.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['ss_kraj']
        row_cells[1].text = f'{row.pedf:,}'
        row_cells[2].text = f'{row.pedf_rel:.1f} %'

    doc.add_paragraph('')
    doc.add_paragraph('Další tabulka obdobně srovnává typy středních škol.')
    doc.add_paragraph('')

    foo = df.groupby('ss_typ_g')[['ones', 'pedf']].sum().reset_index()
    foo['pedf_rel'] = 100 * foo['pedf'] / foo['ones']

    table = doc.add_table(rows=1, cols=4, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Typ školy'
    hdr_cells[1].text = 'Přihlášky na PedF'
    hdr_cells[2].text = 'Přihlášky celkem'
    hdr_cells[3].text = 'Podíl PedF'

    for _, row in foo[foo['ss_typ_g'] != 'Jiné'].iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['ss_typ_g']
        row_cells[1].text = f'{row.pedf:,}'
        row_cells[2].text = f'{row.ones:,}'
        row_cells[3].text = f'{row.pedf_rel:.1f} %'


doc.add_heading('Vývoj mezi lety 2017 a 2021', 3)

doc.add_paragraph('Graf srovnávající vývoj podílu přihlášek na pedagogické fakulty v jednotlivých krajích potvrzuje, že vyšší zájem o pedagogické fakulty není daný jenom vyšším množstvím přihlášek. Podíl přihlášek na pedagogické fakulty se mezi lety 2017 a 2021 zvýšil téměř ve všech krajích.')
doc.add_paragraph('')

foo17 = df17.groupby('ss_kraj')[['ones', 'pedf']].sum().reset_index()
foo17['pedf_rel'] = 100 * foo17['pedf'] / foo17['ones']
foo17['year'] = '2017'
foo21 = df21.groupby('ss_kraj')[['ones', 'pedf']].sum().reset_index()
foo21['pedf_rel'] = 100 * foo21['pedf'] / foo21['ones']
foo21['year'] = '2021'

foo = pd.concat([foo21, foo17], ignore_index=True)

fig, ax = plt.subplots()
sns.scatterplot(data=foo, y='ss_kraj', x='pedf_rel', hue='year', s=60)
ax.set(xlabel='Podíl přihlášek', ylabel=None, title='Podíl přihlášek na pedagogické fakulty')
plt.legend(title='Rok')
plt.plot(foo21.pedf_rel, foo21.ss_kraj)
plt.plot(foo17.pedf_rel, foo17.ss_kraj)
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'prihlasky_kraje_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_paragraph('')
doc.add_paragraph('Další graf srovnává vývoj podílu přihlášek na pedagogické fakulty mezi lety 2017 a 2021 podle typu střední školy. U absolventů SOŠ se podíl přihlášek na pedagogické fakulty prakticky nezměnil, oproti tomu u absolventů gymnázií podíl přihlášek na pedagogické fakulty narostl (zejména u čtyřletých gymnázií), podíl se zvýšil také u uchazečů, kteří se nehlásí ze střední školy.')
doc.add_paragraph('')

foo17 = df17.groupby('ss_typ_g')[['ones', 'pedf']].sum().reset_index()
foo17['pedf_rel'] = 100 * foo17['pedf'] / foo17['ones']
foo17['year'] = '2017'
foo17 = foo17[foo17['ss_typ_g'] != 'Jiné']
foo21 = df21.groupby('ss_typ_g')[['ones', 'pedf']].sum().reset_index()
foo21['pedf_rel'] = 100 * foo21['pedf'] / foo21['ones']
foo21['year'] = '2021'
foo21 = foo21[foo21['ss_typ_g'] != 'Jiné']
foo = pd.concat([foo21, foo17], ignore_index=True)

fig, ax = plt.subplots()
sns.scatterplot(data=foo, y='ss_typ_g', x='pedf_rel', hue='year', s=60)
ax.set(xlabel='Podíl přihlášek', ylabel=None, title='Podíl přihlášek na pedagogické fakulty')
plt.legend(title='Rok')
plt.plot(foo21.pedf_rel, foo21.ss_typ_g)
plt.plot(foo17.pedf_rel, foo17.ss_typ_g)
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'prihlasky_typ_ss_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

#endregion

# 4. ÚSPĚŠNOST UCHAZEČŮ O PEDAGOGICKÉ FAKULTY
#region
doc.add_heading('4. Úspěšnost uchazečů o pedagogické fakulty', 1)

for y in [17, 21]:
    ff = eval(f'ff{y}')

    avg_prijat = ff['prijat_bool'].mean()
    avg_prijat_divka = ff[ff['gender'] == 'Dívka']['prijat_bool'].mean()
    avg_prijat_chlapec = ff[ff['gender'] == 'Chlapec']['prijat_bool'].mean()

    ff = ff[ff['prihl_pedf_bool']]
    avg_prijat_pedf = ff['prijat_pedf_bool'].mean()
    avg_prijat_pedf_divka = ff[ff['gender'] == 'Dívka']['prijat_pedf_bool'].mean()
    avg_prijat_pedf_chlapec = ff[ff['gender'] == 'Chlapec']['prijat_pedf_bool'].mean()

    doc.add_heading(f'20{y}:', 2)

    doc.add_paragraph(f'V roce 20{y} bylo přijato {100 * avg_prijat:.1f} % uchazečů, kteří se hlásili na vysoké školy (u dívek {100 * avg_prijat_divka:.1f} %, u chlapců {100 * avg_prijat_chlapec:.1f} % – rozdíl je pravděpodobně dán odlišným výběrem oborů).')
    doc.add_paragraph('')

    doc.add_paragraph(f'Na pedagogické fakulty bylo přijato {100 * avg_prijat_pedf:.1f} % z uchazečů (u dívek {100 * avg_prijat_pedf_divka:.1f} %, u chlapců {100 * avg_prijat_pedf_chlapec:.1f} %).')
    doc.add_paragraph('')

    cols = ['prihl_nepedf', 'prihl_pedf', 'prijat_nepedf', 'prijat_pedf', 'zapis_nepedf', 'zapis_pedf']
    foo = (ff[cols] > 0).value_counts().rename('count').reset_index().query('prihl_nepedf & prihl_pedf & prijat_nepedf & prijat_pedf').drop(columns=['prihl_nepedf', 'prihl_pedf', 'prijat_nepedf', 'prijat_pedf']).reset_index(drop=True)
    ano_ne = {True: 'Ano', False: 'Ne'}
    foo['zapis_pedf'] = foo['zapis_pedf'].map(ano_ne)
    foo['zapis_nepedf'] = foo['zapis_nepedf'].map(ano_ne)

    tot = foo['count'].sum()

    doc.add_paragraph(f'Celkem {tot:,} uchazečů se hlásili a zároveň byli přijati na pedagogickou fakultu i na jinou fakultu. U těchto uchazečů je možné sledovat, na kterou z fakult se nakonec zapsali:')

    table = doc.add_table(rows=1, cols=4, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Zápis PedF'
    hdr_cells[1].text = 'Zápis jinde'
    hdr_cells[2].text = 'Počet'
    hdr_cells[3].text = 'Podíl'

    for _, row in foo.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['zapis_pedf']
        row_cells[1].text = row['zapis_nepedf']
        row_cells[2].text = f'{row["count"]:,}'
        row_cells[3].text = f'{100 * row["count"] / tot:.1f} %'




# for y in [21, 17]:
#     ff = eval(f'ff{y}')

#     tot_uchazecu = ff['prihl_pedf'].shape[0]
#     tot_chlapcu = (ff['gender'] == 'Chlapec').sum()
#     tot_divek = (ff['gender'] == 'Dívka').sum()


#     doc.add_heading(f'20{y}:', 2)
#endregion

doc.show()

# TO DO?
# - obdoba krajů pro typy střední školy
# - úspěšnost uchazečů
# - ???

#endregion

###  OLD AND UNUSED MESS  ###
#region
foo21 = pedf_app(ff21, 'ss_typ_g')
foo21


foo21 = pedf_app(ff21, 'ss_kraj')
ff21.columns


ff21['ss_typ'].value_counts()
ff21[['ss_typ', 'rmat']].value_counts()

ff21[ff21['ss_typ'] == 'Gymnázium']['ss_gym_delka'].value_counts()

foo21


foo17 = pedf_app(ff17, 'ss_kraj')[['ss_kraj', 'pedf_rel']].rename(columns={'pedf_rel': 2017})
foo21 = pedf_app(ff21, 'ss_kraj')[['ss_kraj', 'pedf_rel']].rename(columns={'pedf_rel': 2021})
foo = pd.merge(foo21, foo17)

fig, ax = plt.subplots()
sns.barplot(foo)
ax.show()
foo.plot().show()

doc.add_paragraph('')

fig, ax = plt.subplots()
ax = sns.histplot(data=ff, y='prihl', discrete=True, binwidth=0.8, shrink=0.8, ec=None)
ax.set(xlabel='Četnost', ylabel='Počet přihlášek', title='Počet přihlášek jednotlivých uchazečů')
ax.set_ylim(0, 13)
ax.invert_yaxis()
fig.tight_layout()

figpath = savefig(fig, f'cetnost_prihlasek_{y}')

doc.add_picture(figpath, width=CHART_WIDTH)
doc.add_paragraph('')



doc.show()

latent_styles = doc.styles.latent_styles
len(latent_styles)

table_style = [ls for ls in latent_styles if ls.name == 'Grid Table Light'][0]
doc.styles.add_style('Table Grid Light', docx.enum.style.WD_STYLE_TYPE.TABLE)

for ls in latent_styles:
    print(ls.name)

table = doc.add_table(rows=1, cols=3, style='Plain Table 1')
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Kraj'
hdr_cells[1].text = 'Přihlášek na PedF'
hdr_cells[2].text = 'Podíl na PedF'

for _, row in foo.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row['ss_kraj']
    row_cells[1].text = f'{row.prihl_bool:,}'
    row_cells[2].text = f'{row.pedf_rel} %'

doc.show()


doc.add_heading('Úspěšnost uchazečů o pedagogické fakulty')

for y in [21, 17]:
    ff = eval(f'ff{y}')

    tot_uchazecu = ff['prihl_pedf'].shape[0]
    tot_chlapcu = (ff['gender'] == 'Chlapec').sum()
    tot_divek = (ff['gender'] == 'Dívka').sum()


    doc.add_heading(f'20{y}:', 2)



pedf_app(ff21, 'gender')

pedf_app(ff21, 'ss_kraj')


doc.show()

df21[df21['pedf']][['fak_nazev', 'vs_nazev']].drop_duplicates().sort_values('vs_nazev')


sns.histplot(df21['pocet'], discrete=True, binwidth=0.8, ec=None).show()

doc.add_heading('2017:', 2)
doc.add_paragraph('')



doc.show()



os.getcwd()

font_size_word()
sns.histplot(data=ff21, y='prihl', hue='gender', multiple='stack', discrete=True, binwidth=0.8, ec=None).show()

ax = sns.histplot(data=ff21, y='prihl', discrete=True, binwidth=0.8, shrink=0.8, ec=None)
ax.set(xlabel='Četnost', ylabel='Počet přihlášek', title='Počet přihlášek jednotlivých uchazečů')
ax.set_ylim(0, 13)
ax.invert_yaxis()
ax.get_figure().tight_layout()

figpath = savefig(ax, 'cetnost_prihlasek_21')

doc.add_picture(figpath, width=Mm(140))
doc.show()

ax.get_figure().savefig

ax.get_figure()

ax.show()

ff21['prihl'].value_counts()

plt.rcParams['figure.titlesize'] = 30
ax.show()

eval(f'df{17}')

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc


p = doc.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

doc.add_heading('Heading, level 1', level=1)

doc.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
doc.add_paragraph(
    'first item in ordered list', style='List Number'
)

for style in doc.styles:
    print("style.name == %s" % style.name)






















### Learning how to use docx...


df.show()

document = Document()
document.add_heading('Document title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save(f'{foo}/demo.docx')

foo = create_stamped_temp('docx')
foo

os.startfile(f'{foo}/demo.docx')

# how to set a style?


doc = Document()
styles = doc.styles

style = styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(14)

doc.add_heading('Document title', 0)

doc.show()
docx_document_show(doc)
doc.save(tempfile.gettempdir() + '/foo.docx')

p = doc.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

doc.add_heading('Heading, level 1', level=1)

doc.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
doc.add_paragraph(
    'first item in ordered list', style='List Number'
)



doc.show()

foo = create_stamped_temp('docx')
doc.save(f'{foo}/demo.docx')
os.startfile(f'{foo}/demo.docx')

doc = Document('../IDEA_SABLONA.docx')
doc.show()

type(doc)

doc = Document()


doc = Document('empty.docx')
doc._body.clear_content()


doc.add_heading('Document title', 0)

p = doc.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

doc.add_heading('Heading, level 1', level=1)

doc.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
doc.add_paragraph(
    'first item in ordered list', style='List Number'
)

stamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
path = tempfile.gettempdir() + '/docs'
os.makedirs(path, exist_ok=True)
path = '{}/{}.{}'.format(path, stamp, format)

path

self.save(path, **kwargs)

doc.element.xml

doc.show()
doc

docx_document_show(doc)

os.environ['TEMP'] = 'D:/temp'

tempfile.tempdir = 'D:/temp'
tempfile.gettempdir()

Document.show = docx_document_show
doc.show()

Document
Document()

#endregion
