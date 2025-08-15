###  IMPORTS  ###
#region
import os
import sys
from venv import create
import pyreadstat
import pandas as pd
import numpy as np
import re
import requests
from urllib.request import urlopen
from io import StringIO
from statsmodels.stats.weightstats import DescrStatsW
from sklearn.linear_model import LinearRegression
import statsmodels.api as sm
import statsmodels.formula.api as smf
import scipy
import seaborn as sns
import matplotlib as mpl
import matplotlib.pyplot as plt
plt.rcParams['figure.figsize'] = 12, 6
import dbf
import json

from docx import Document
from docx.shared import Mm, Pt

from libs.utils import *
from libs.plots import *
from libs.extensions import *
from libs import uchazec

import importlib

# importlib.reload(uchazec)
# importlib.invalidate_caches()
# import locale
# locale.setlocale(locale.LC_ALL, 'cs_CS')
#endregion

###  SETUP THE ENVIRONMENT  ###
#region
# font sizes for importing to word
def font_size(small=9, medium=11, big='large'):
    plt.rc('font', size=small)          # controls default text sizes
    plt.rc('axes', titlesize=big)     # fontsize of the axes title
    plt.rc('axes', labelsize=medium)    # fontsize of the x and y labels
    plt.rc('xtick', labelsize=small)    # fontsize of the tick labels
    plt.rc('ytick', labelsize=small)    # fontsize of the tick labels
    plt.rc('legend', fontsize=small)    # legend fontsize
    plt.rc('figure', titlesize=big)  # fontsize of the figure title    

def font_size_word():
    font_size(small=12, medium=14, big=18)
    
def savefig(ax_or_fig, title):
    fig_path = f'D:\\projects\\idea\\code\\output\\uchazec\\{title}.png'
    fig = ax_or_fig if isinstance(ax_or_fig, plt.Figure) else ax_or_fig.get_figure()
    fig.savefig(fig_path)
    return fig_path

font_size_word()
CHART_WIDTH = Mm(158)
# plt.style.use('dark_background')
#endregion

###  LOAD AND PREPARE DATA  ###
#region

# for year in [17, 21]:
#     print(f'loading 20{year}')

#     # load and process data
#     df, variable_labels, value_labels = uchazec.loader(year=year)

#     # save data for stata -- this is apparently broken in windows - crashes on value labels...
#     # df.to_stata(f'{uchazec.data_root}/uchazec/uch{year}.dta', write_index=False, version=118, variable_labels=variable_labels, value_labels=value_labels)
    
#     # apply value labels and save for python
#     for c in df.columns:
#         if c in value_labels.keys():
#             df[c] = df[c].map(value_labels[c]).astype('category')
#     df.to_parquet(f'temp/uchazec/uch{year}.parquet')

regenerate_data = False

if regenerate_data:

    print('Regenerating from cache.')
    df17 = pd.read_parquet('temp/uchazec/uch17.parquet')
    df21 = pd.read_parquet('temp/uchazec/uch21.parquet')

    df21 = uchazec.add_isced(df21)
    df17 = uchazec.add_isced(df17)

    isc17 = uchazec.get_per_isced(df17)
    isc21 = uchazec.get_per_isced(df21)

    df17, rep17 = uchazec.filter_data(df17)
    df21, rep21 = uchazec.filter_data(df21)

    df17 = uchazec.add_variables(df17)
    df21 = uchazec.add_variables(df21)

    ff17 = uchazec.get_per_app(df17)
    ff21 = uchazec.get_per_app(df21)

    df21 = uchazec.add_ss_typ_g(df21)
    df17 = uchazec.add_ss_typ_g(df17)
    ff21 = uchazec.add_ss_typ_g(ff21)
    ff17 = uchazec.add_ss_typ_g(ff17)

    def pedf_app(ff, c):
        foo = ff.groupby(c)[['prihl_pedf_bool', 'prihl_bool']].sum().reset_index()
        foo['pedf_rel'] = np.round(100 * foo['prihl_pedf_bool'] / foo['prihl_bool'], 1)
        return foo

    pedf_vs = df21[df21['pedf']][['fak_nazev', 'vs_nazev']].drop_duplicates().sort_values('vs_nazev').reset_index(drop=False)

    print('Storing final datasets.')
    df17.to_parquet('temp/uchazec/final/uch17.parquet')
    df21.to_parquet('temp/uchazec/final/uch21.parquet')
    ff17.to_parquet('temp/uchazec/final/ff17.parquet')
    ff21.to_parquet('temp/uchazec/final/ff21.parquet')
    isc17.to_parquet('temp/uchazec/final/isc17.parquet')
    isc21.to_parquet('temp/uchazec/final/isc21.parquet')
    pedf_vs.to_parquet('temp/uchazec/final/pedf_vs.parquet')
    with open('temp/uchazec/final/rep17.json', 'w') as f:
        json.dump(rep17, f)
    with open('temp/uchazec/final/rep21.json', 'w') as f:
        json.dump(rep21, f)        


print('Reading final datasets.')
df17 = pd.read_parquet('temp/uchazec/final/uch17.parquet')
df21 = pd.read_parquet('temp/uchazec/final/uch21.parquet')
ff17 = pd.read_parquet('temp/uchazec/final/ff17.parquet')
ff21 = pd.read_parquet('temp/uchazec/final/ff21.parquet')
isc17 = pd.read_parquet('temp/uchazec/final/isc17.parquet')
isc21 = pd.read_parquet('temp/uchazec/final/isc21.parquet')
pedf_vs = pd.read_parquet('temp/uchazec/final/pedf_vs.parquet')
with open('temp/uchazec/final/rep17.json', 'r') as f:
    rep17 = json.load(f)
with open('temp/uchazec/final/rep21.json', 'r') as f:
    rep21 = json.load(f)




# df = df21
# df.head(n=1000).show()
#endregion

###  GENERATE REPORT  ###
#region

# EMPTY DOC
#region
doc = Document('template.docx')
doc._body.clear_content()

section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(25.4)
section.right_margin = Mm(25.4)
section.top_margin = Mm(25.4)
section.bottom_margin = Mm(25.4)

doc.add_heading('IDEA: Uchazeč', 0)

# table of contents field
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc.add_paragraph('')
paragraph = doc.add_paragraph()
run = paragraph.add_run()
fldChar = OxmlElement('w:fldChar')  # creates a new element
fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
instrText.text = 'TOC \\o "1-2" \\h \\z \\u'   # change 1-3 depending on heading levels you need

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:t')
fldChar3.text = "Right-click to update field."
fldChar2.append(fldChar3)

fldChar4 = OxmlElement('w:fldChar')
fldChar4.set(qn('w:fldCharType'), 'end')

r_element = run._r
r_element.append(fldChar)
r_element.append(instrText)
r_element.append(fldChar2)
r_element.append(fldChar4)
p_element = paragraph._p

def add_hyperlink(paragraph, url, text, color='0000FF', underline=True):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Add ? Remove underlining if it is requested
    if underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      # u.set(docx.oxml.shared.qn('w:val'), 'none')
      u.set(docx.oxml.shared.qn('w:val'), 'single')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

#endregion

# 1. PŘÍPRAVA DAT
#region
doc.add_heading('1. Příprava dat', 1)
doc.add_paragraph('Analýza je založena na anonymizovaných datech z databáze Uchazeč, která obsahuje informaci o všech přihlášených na vysoké školy pro roky 2017 a 2021.')

for y in [17, 21]:
    doc.add_heading(f'Data 20{y}', 2)
    for r in eval(f'rep{y}'):
        doc.add_paragraph(r, style='Bullet List')

doc.add_heading('Limity dat', 2)
doc.add_paragraph('V datech se vyskytují různé problémy, které komplikují analýzu:')
doc.add_paragraph('')
doc.add_paragraph('Stejný student má často více přihlášek, které se v různých údajích (například o SŠ, odkud se hlásí) mohou lišit. Tedy například jedna přihláška uvádí "SOŠ", druhá "Není SŠ", případně odlišná lokalita SŠ nebo chybějící údaje). Při agregaci na úroveň jednotlivých respondentů bereme pro jednoduchost v úvahu údaj z prvního záznamu o přihlášce uchazeče.', style='Bullet List')
doc.add_paragraph('')
p = doc.add_paragraph('Informace o výsledku přijímacího řízení je v datech z roku 2017 kódována mírně odlišným způsobem a některé hodnoty již ', style='Bullet List')
add_hyperlink(p, 'http://stistko.uiv.cz/katalog/cslnk.asp?idc=MCPR&ciselnik=V%FDsledek+p%F8ij%EDmac%EDho+%F8%EDzen%ED&aap=on', 'aktuální číselník')
p.add_run(' neobsahuje. Lze je dohledat ve ')
add_hyperlink(p, 'http://stistko.uiv.cz/katalog/ciselnik11x.asp?idc=MCPR&ciselnik=V%FDsledek+p%F8ij%EDmac%EDho+%F8%EDzen%ED&aap=on&poznamka=', 'starších souborech')
p.add_run(', tím se opět komplikuje napojování dat, protože starší číselníky nejsou dostupné jako standardizovaný XML soubor.')
doc.add_paragraph('')
doc.add_paragraph('Studijní programy a obory jsou v datech pro rok 2017 kódovány odlišným způsobem než v datech pro rok 2021, proto není možné uchazeče o pedagogické studium snadno identifikovat na úrovni oborů nebo programů. Proto se analýza zaměřuje zatím na uchazeče o studium na pedagogických fakultách, ačkoli takové vymezení není zcela přesné.', style='Bullet List')
#endregion

# 2. PŘIHLÁŠKY NA VYSOKÉ ŠKOLY
#region
doc.add_heading('2. Přihlášky na vysoké školy', 1)

for y in [17, 21]:
    ff = eval(f'ff{y}')
    tot_prihlasek = eval(f'df{y}').shape[0]
    tot_uchazecu = ff.shape[0]
    tot_chlapcu = (ff['gender'] == 'Chlapec').sum()
    tot_divek = (ff['gender'] == 'Dívka').sum()
    avg_prihlasek = ff['prihl'].mean()
    max_prihlasek = ff['prihl'].max()

    doc.add_heading(f'Přihlášky 20{y}', 2)
    doc.add_paragraph(f'V roce 20{y} bylo podáno {tot_prihlasek:,} přihlášek na vysoké školy od {tot_uchazecu:,} českých uchazečů, z toho se hlásilo {tot_chlapcu:,} chlapců ({100 * tot_chlapcu / tot_uchazecu:.3g} %) a {tot_divek:,} dívek ({100 * tot_divek / tot_uchazecu:.3g} %). Průměrný uchazeč si podal {avg_prihlasek:.1f} přihlášek, maximální počet podaných přihlášek byl {max_prihlasek:.0f}. Četnost počtu přihlášek jednotlivých uchazečů ukazuje následující graf.')
    doc.add_paragraph('')

    fig, ax = plt.subplots()
    ax = sns.histplot(data=ff, y='prihl', discrete=True, binwidth=0.8, shrink=0.8, ec=None)
    ax.set(xlabel='Četnost', ylabel='Počet přihlášek', title='Počet přihlášek jednotlivých uchazečů')
    ax.set_ylim(0, 13)
    ax.invert_yaxis()
    fig.tight_layout()

    figpath = savefig(fig, f'cetnost_prihlasek_{y}')

    doc.add_picture(figpath, width=CHART_WIDTH)
    doc.add_paragraph('')
#endregion

# 3. PROFIL UCHAZEČŮ O PEDAGOGICKÉ FAKULTY
#region
doc.add_heading('3. Profil uchazečů o pedagogické fakulty', 1)
doc.add_paragraph('Uchazeči se mohli hlásit na některou z osmi pedagogických fakult v České republice:')
doc.add_paragraph('')

for _, row in pedf_vs.iterrows():
    doc.add_paragraph(row['vs_nazev'], style='Bullet List')

for y in [17, 21]:
    ff = eval(f'ff{y}')

    tot_uchazecu = ff.shape[0]
    tot_chlapcu = (ff['gender'] == 'Chlapec').sum()
    tot_divek = (ff['gender'] == 'Dívka').sum()

    pedf_uchazecu = ff['prihl_pedf'].sum()
    pedf_chlapcu = ff[ff['gender'] == 'Chlapec']['prihl_pedf_bool'].sum()
    pedf_divek = ff[ff['gender'] == 'Dívka']['prihl_pedf_bool'].sum()

    doc.add_heading(f'Uchazeči 20{y}', 2)

    doc.add_paragraph(f'V roce 20{y} se na pedagogické fakulty hlásilo {pedf_uchazecu:,.0f} ({100 * pedf_uchazecu / tot_uchazecu:.3g} %). Z toho bylo {pedf_divek:,} dívek ({100 * pedf_divek / tot_divek:.3g} % všech dívek) a {pedf_chlapcu:,} chlapců ({100 * pedf_chlapcu / tot_chlapcu:.3g} % všech chlapců). Následující tabulka shrnuje, jaký podíl uchazečů se hlásil na pedagogické fakulty z jednotlivých krajů (procenta vyjadřují podíl ze všech uchazečů daného kraje).')
    doc.add_paragraph('')

    foo = pedf_app(ff, 'ss_kraj')
    table = doc.add_table(rows=1, cols=3, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kraj'
    hdr_cells[1].text = 'Uchazeči o PedF'
    hdr_cells[2].text = 'Podíl PedF'

    for _, row in foo.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['ss_kraj']
        row_cells[1].text = f'{row.prihl_pedf_bool:,}'
        row_cells[2].text = f'{row.pedf_rel} %'

    doc.add_paragraph('')
    p = doc.add_paragraph('Další tabulka ukazuje podobné srovnání podle typu střední školy.')
    if y == 17:
        p.add_run(' ("Není SŠ" zpravidla označuje uchazeče, které se nehlásí přímo ze střední školy, ale není tomu tak vždy – v některých případech byly takto označeni i aktuální maturanti, naopak u některých starších uchazečů byl typ střední školy vyplněný správně; malý počet uchazečů s jiným typem střední školy byl vynechán.)')
    if y == 21:
        p.add_run(' Celkově se počet absolventů gymnázií hlásících se na vysoké školy změnil jen mírně, oproti tomu počet absolventů středních odborných škol, kteří se hlásí na vysoké školy, znatelně narostl. Nejvíce studentů se na pedagogické fakulty hlásí ze středních odborných škol, v relativním vyjádření je to pak největší podíl z absolventů čtyřletých gymnázií.')
    doc.add_paragraph('')

    foo = pedf_app(ff, 'ss_typ_g')
    table = doc.add_table(rows=1, cols=4, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Typ školy'
    hdr_cells[1].text = 'Uchazeči o PedF'
    hdr_cells[2].text = 'Uchazeči celkem'
    hdr_cells[3].text = 'Podíl PedF'

    for _, row in foo[foo['ss_typ_g'] != 'Jiné'].iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['ss_typ_g']
        row_cells[1].text = f'{row.prihl_pedf_bool:,}'
        row_cells[2].text = f'{row.prihl_bool:,}'
        row_cells[3].text = f'{row.pedf_rel} %'


doc.add_heading('Srovnání vývoje', 2)
doc.add_paragraph('Následující graf srovnává podíl uchazečů o pedagogické fakulty z jednotlivých krajů v letech 2017 a 2021. Je patrné, že podíl zájemců o pedagogické fakulty se v roce 2021 výrazně zvýšil téměř ve všech krajích.')
doc.add_paragraph('')

foo17 = pedf_app(ff17, 'ss_kraj')[['ss_kraj', 'pedf_rel']]
foo17['year'] = '2017'
foo21 = pedf_app(ff21, 'ss_kraj')[['ss_kraj', 'pedf_rel']]
foo21['year'] = '2021'
foo = pd.concat([foo21, foo17], ignore_index=True)

fig, ax = plt.subplots()
sns.scatterplot(data=foo, y='ss_kraj', x='pedf_rel', hue='year', s=60)
ax.set(xlabel='Podíl uchazečů', ylabel=None, title='Podíl uchazečů o pedagogické fakulty')
plt.legend(title='Rok')
plt.plot(foo21.pedf_rel, foo21.ss_kraj)
plt.plot(foo17.pedf_rel, foo17.ss_kraj)
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'uchazeci_kraje_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_paragraph('')
doc.add_paragraph('Další graf srovnává vývoj podílu uchazečů mezi lety 2017 a 2021 podle typu střední školy. Je patrné, že nejvíce zájemců o střední školy přibylo mezi absolventy 4-letých gymnázií, případně také z víceletých gymnázií.')
doc.add_paragraph('')

foo17 = pedf_app(ff17, 'ss_typ_g')[['ss_typ_g', 'pedf_rel']]
foo17['year'] = '2017'
foo17 = foo17[foo17['ss_typ_g'] != 'Jiné']
foo21 = pedf_app(ff21, 'ss_typ_g')[['ss_typ_g', 'pedf_rel']]
foo21['year'] = '2021'
foo21 = foo21[foo21['ss_typ_g'] != 'Jiné']
foo = pd.concat([foo21, foo17], ignore_index=True)

fig, ax = plt.subplots()
sns.scatterplot(data=foo, y='ss_typ_g', x='pedf_rel', hue='year', s=60)
ax.set(xlabel='Podíl uchazečů', ylabel=None, title='Podíl uchazečů o pedagogické fakulty')
plt.legend(title='Rok')
plt.plot(foo21.pedf_rel, foo21.ss_typ_g)
plt.plot(foo17.pedf_rel, foo17.ss_typ_g)
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'uchazeci_typ_ss_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_heading('Podíl z přihlášek', 2)

avg_prihl_17 = ff17['prihl'].mean()
avg_prihl_pedf_17 = ff17['prihl_pedf'].mean()
avg_prihl_nepedf_17 = ff17['prihl_nepedf'].mean()
avg_prihl_21 = ff21['prihl'].mean()
avg_prihl_pedf_21 = ff21['prihl_pedf'].mean()
avg_prihl_nepedf_21 = ff21['prihl_nepedf'].mean()

doc.add_paragraph(f'Data ukazují, že na pedagogické fakulty se hlásí vyšší podíl uchazečů. Zároveň si v roce 2021 průměrný uchazeč podává více přihlášek: {avg_prihl_21:.2g} oproti {avg_prihl_17:.2g} přihlášce v roce 2017. Z toho si v roce 2017 podal průměrný uchazeč {avg_prihl_pedf_17:.2g} přihlášek na pedagogické fakulty a {avg_prihl_nepedf_17:.2g} na ostatní fakulty. V roce 2021 to bylo {avg_prihl_pedf_21:.2g} přihlášek na pedagogické fakulty (nárůst o {100 * avg_prihl_pedf_21 / avg_prihl_pedf_17 - 100:.3g} %) a {avg_prihl_nepedf_21:.2g} na ostatní fakulty (nárůst o {100 * avg_prihl_nepedf_21 / avg_prihl_nepedf_17 - 100:.3g} %).')

doc.add_paragraph(f'Pro ověření výsledků ukazují následující tabulky a grafy obdobná srovnání jako předchozí podíly podle krajů a typů střední školy, avšak vyjádřené jako podíl všech přihlášek, nikoli jako podíl uchazečů.')

df17.groupby('ss_kraj')[['ones', 'pedf']].sum().reset_index()


for y in [17, 21]:
    df = eval(f'df{y}')

    tot_prihlasek = df.shape[0]
    pedf_prihlasek = df['pedf'].sum()

    doc.add_heading(f'20{y}:', 3)

    doc.add_paragraph(f'V roce 20{y} bylo na pedagogické fakulty podáno {pedf_prihlasek:,.0f} z celkového počtu {tot_prihlasek:,.0f} přihlášek na vysoké školy od českých uchazečů (podíl pedagogických fakult tedy byl {100 * pedf_prihlasek / tot_prihlasek:.3g} %). Následující tabulka shrnuje, jaký byl podíl přihlášek na pedagogické fakulty ze všech přihlášek podaných z daného kraje.')
    doc.add_paragraph('')

    foo = df.groupby('ss_kraj')[['ones', 'pedf']].sum().reset_index()
    foo['pedf_rel'] = 100 * foo['pedf'] / foo['ones']

    table = doc.add_table(rows=1, cols=3, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kraj'
    hdr_cells[1].text = 'Přihlášky na PedF'
    hdr_cells[2].text = 'Podíl PedF'

    for _, row in foo.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['ss_kraj']
        row_cells[1].text = f'{row.pedf:,}'
        row_cells[2].text = f'{row.pedf_rel:.1f} %'

    doc.add_paragraph('')
    doc.add_paragraph('Další tabulka obdobně srovnává typy středních škol.')
    doc.add_paragraph('')

    foo = df.groupby('ss_typ_g')[['ones', 'pedf']].sum().reset_index()
    foo['pedf_rel'] = 100 * foo['pedf'] / foo['ones']

    table = doc.add_table(rows=1, cols=4, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Typ školy'
    hdr_cells[1].text = 'Přihlášky na PedF'
    hdr_cells[2].text = 'Přihlášky celkem'
    hdr_cells[3].text = 'Podíl PedF'

    for _, row in foo[foo['ss_typ_g'] != 'Jiné'].iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['ss_typ_g']
        row_cells[1].text = f'{row.pedf:,}'
        row_cells[2].text = f'{row.ones:,}'
        row_cells[3].text = f'{row.pedf_rel:.1f} %'


doc.add_heading('Vývoj mezi lety 2017 a 2021', 3)

doc.add_paragraph('Graf srovnávající vývoj podílu přihlášek na pedagogické fakulty v jednotlivých krajích potvrzuje, že vyšší zájem o pedagogické fakulty není daný jenom vyšším množstvím přihlášek. Podíl přihlášek na pedagogické fakulty se mezi lety 2017 a 2021 zvýšil téměř ve všech krajích.')
doc.add_paragraph('')

foo17 = df17.groupby('ss_kraj')[['ones', 'pedf']].sum().reset_index()
foo17['pedf_rel'] = 100 * foo17['pedf'] / foo17['ones']
foo17['year'] = '2017'
foo21 = df21.groupby('ss_kraj')[['ones', 'pedf']].sum().reset_index()
foo21['pedf_rel'] = 100 * foo21['pedf'] / foo21['ones']
foo21['year'] = '2021'

foo = pd.concat([foo21, foo17], ignore_index=True)

fig, ax = plt.subplots()
sns.scatterplot(data=foo, y='ss_kraj', x='pedf_rel', hue='year', s=60)
ax.set(xlabel='Podíl přihlášek', ylabel=None, title='Podíl přihlášek na pedagogické fakulty')
plt.legend(title='Rok')
plt.plot(foo21.pedf_rel, foo21.ss_kraj)
plt.plot(foo17.pedf_rel, foo17.ss_kraj)
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'prihlasky_kraje_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_paragraph('')
doc.add_paragraph('Další graf srovnává vývoj podílu přihlášek na pedagogické fakulty mezi lety 2017 a 2021 podle typu střední školy. U absolventů SOŠ se podíl přihlášek na pedagogické fakulty prakticky nezměnil, oproti tomu u absolventů gymnázií podíl přihlášek na pedagogické fakulty narostl (zejména u čtyřletých gymnázií), podíl se zvýšil také u uchazečů, kteří se nehlásí ze střední školy.')
doc.add_paragraph('')

foo17 = df17.groupby('ss_typ_g')[['ones', 'pedf']].sum().reset_index()
foo17['pedf_rel'] = 100 * foo17['pedf'] / foo17['ones']
foo17['year'] = '2017'
foo17 = foo17[foo17['ss_typ_g'] != 'Jiné']
foo21 = df21.groupby('ss_typ_g')[['ones', 'pedf']].sum().reset_index()
foo21['pedf_rel'] = 100 * foo21['pedf'] / foo21['ones']
foo21['year'] = '2021'
foo21 = foo21[foo21['ss_typ_g'] != 'Jiné']
foo = pd.concat([foo21, foo17], ignore_index=True)

fig, ax = plt.subplots()
sns.scatterplot(data=foo, y='ss_typ_g', x='pedf_rel', hue='year', s=60)
ax.set(xlabel='Podíl přihlášek', ylabel=None, title='Podíl přihlášek na pedagogické fakulty')
plt.legend(title='Rok')
plt.plot(foo21.pedf_rel, foo21.ss_typ_g)
plt.plot(foo17.pedf_rel, foo17.ss_typ_g)
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'prihlasky_typ_ss_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

#endregion

# 4. ÚSPĚŠNOST UCHAZEČŮ O PEDAGOGICKÉ FAKULTY
#region
doc.add_heading('4. Úspěšnost uchazečů o pedagogické fakulty', 1)

for y in [17, 21]:
    ff = eval(f'ff{y}')

    avg_prijat = ff['prijat_bool'].mean()
    avg_prijat_divka = ff[ff['gender'] == 'Dívka']['prijat_bool'].mean()
    avg_prijat_chlapec = ff[ff['gender'] == 'Chlapec']['prijat_bool'].mean()

    ff = ff[ff['prihl_pedf_bool']]
    avg_prijat_pedf = ff['prijat_pedf_bool'].mean()
    avg_prijat_pedf_divka = ff[ff['gender'] == 'Dívka']['prijat_pedf_bool'].mean()
    avg_prijat_pedf_chlapec = ff[ff['gender'] == 'Chlapec']['prijat_pedf_bool'].mean()
    tot_prijat_pedf = ff['prijat_pedf_bool'].sum()

    doc.add_heading(f'Rok 20{y}', 2)

    doc.add_paragraph(f'V roce 20{y} bylo přijato {100 * avg_prijat:.1f} % uchazečů, kteří se hlásili na vysoké školy (u dívek {100 * avg_prijat_divka:.1f} %, u chlapců {100 * avg_prijat_chlapec:.1f} % – rozdíl je pravděpodobně dán odlišným výběrem oborů).')
    doc.add_paragraph('')

    doc.add_paragraph(f'Na pedagogické fakulty bylo přijato {100 * avg_prijat_pedf:.1f} % z uchazečů (u dívek {100 * avg_prijat_pedf_divka:.1f} %, u chlapců {100 * avg_prijat_pedf_chlapec:.1f} %), celkem se jednalo o {tot_prijat_pedf:,.0f} osob.')
    doc.add_paragraph('')

    cols = ['prihl_nepedf', 'prihl_pedf', 'prijat_nepedf', 'prijat_pedf', 'zapis_nepedf', 'zapis_pedf']
    foo = (ff[cols] > 0).value_counts().rename('count').reset_index().query('prihl_nepedf & prihl_pedf & prijat_nepedf & prijat_pedf').drop(columns=['prihl_nepedf', 'prihl_pedf', 'prijat_nepedf', 'prijat_pedf']).reset_index(drop=True)
    ano_ne = {True: 'Ano', False: 'Ne'}
    foo['zapis_pedf'] = foo['zapis_pedf'].map(ano_ne)
    foo['zapis_nepedf'] = foo['zapis_nepedf'].map(ano_ne)

    tot = foo['count'].sum()

    doc.add_paragraph(f'Celkem {tot:,} uchazečů se hlásilo a zároveň bylo přijato na pedagogickou fakultu i na jinou fakultu. U těchto uchazečů je možné sledovat, na kterou z fakult se skutečně zapsali, jak shrnuje následující tabulka.')
    doc.add_paragraph('')

    table = doc.add_table(rows=1, cols=4, style='Plain Table 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Zápis PedF'
    hdr_cells[1].text = 'Zápis jinde'
    hdr_cells[2].text = 'Počet'
    hdr_cells[3].text = 'Podíl'

    for _, row in foo.sort_values(['zapis_pedf', 'zapis_nepedf']).iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['zapis_pedf']
        row_cells[1].text = row['zapis_nepedf']
        row_cells[2].text = f'{row["count"]:,}'
        row_cells[3].text = f'{100 * row["count"] / tot:.1f} %'

doc.add_paragraph('')
doc.add_paragraph('Data naznačují, že přijímací řízení na pedagogické fakulty v roce 2021 bylo náročnější než v roce 2017. Zároveň uchazeči, kteří si při zápisu vybírali mezi pedagogickou a nepedagogickou fakultou, dávali v roce 2021 častěji přednost pedagogické fakultě, oproti tomu v roce 2017 dávali spíše přednost nepedagogické fakultě.')

doc.add_heading('Úspěšnost podle krajů', 2)

ff = ff[ff['prihl_pedf_bool']]
avg_prijat_pedf = ff['prijat_pedf_bool'].mean()

foo17 = ff17[ff17['prihl_pedf_bool']].copy().groupby('ss_kraj')['prijat_pedf_bool'].mean().reset_index()
foo17['year'] = '2017'
foo17['prijat'] = 100 * foo17['prijat_pedf_bool']
foo21 = ff21[ff21['prihl_pedf_bool']].copy().groupby('ss_kraj')['prijat_pedf_bool'].mean().reset_index()
foo21['year'] = '2021'
foo21['prijat'] = 100 * foo21['prijat_pedf_bool']

foo = pd.concat([foo21, foo17], ignore_index=True)

fig, ax = plt.subplots()
sns.scatterplot(data=foo, y='ss_kraj', x='prijat', hue='year', s=60)
ax.set(xlabel='Podíl přijatých', ylabel=None, title='Podíl přijatých na pedagogické fakulty')
plt.legend(title='Rok')
plt.plot(foo21.prijat, foo21.ss_kraj)
plt.plot(foo17.prijat, foo17.ss_kraj)
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'prijati_pedf_kraje_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_heading('Úspěšnost podle typu střední školy', 2)

foo17 = ff17[ff17['prihl_pedf_bool']].copy().groupby('ss_typ_g')['prijat_pedf_bool'].mean().reset_index()
foo17['year'] = '2017'
foo17 = foo17[foo17['ss_typ_g'] != 'Jiné']
foo17['prijat'] = 100 * foo17['prijat_pedf_bool']
foo21 = ff21[ff21['prihl_pedf_bool']].copy().groupby('ss_typ_g')['prijat_pedf_bool'].mean().reset_index()
foo21['year'] = '2021'
foo21 = foo21[foo21['ss_typ_g'] != 'Jiné']
foo21['prijat'] = 100 * foo21['prijat_pedf_bool']
foo = pd.concat([foo21, foo17], ignore_index=True)

fig, ax = plt.subplots()
sns.scatterplot(data=foo, y='ss_typ_g', x='prijat', hue='year', s=60)
ax.set(xlabel='Podíl přijatých', ylabel=None, title='Podíl přijatých na pedagogické fakulty')
plt.legend(title='Rok')
plt.plot(foo21.prijat, foo21.ss_typ_g)
plt.plot(foo17.prijat, foo17.ss_typ_g)
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'prijati_pedf_typ_ss_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

doc.add_heading('Náročnost přijetí podle fakult', 2)

foo17 = df17[df17['pedf']].copy()
foo17['vs_nazev'] = foo17['vs_nazev'].cat.remove_unused_categories()
foo17 = foo17.groupby('vs_nazev')['vypr_flag'].mean().reset_index()
foo17['year'] = '2017'
foo17['prijat'] = 100 * foo17['vypr_flag']

foo21 = df21[df21['pedf']].copy()
foo21['vs_nazev'] = foo21['vs_nazev'].cat.remove_unused_categories()
foo21 = foo21.groupby('vs_nazev')['vypr_flag'].mean().reset_index()
foo21['year'] = '2021'
foo21['prijat'] = 100 * foo21['vypr_flag']

foo = pd.concat([foo21, foo17], ignore_index=True)

fig, ax = plt.subplots()
sns.scatterplot(data=foo, y='vs_nazev', x='prijat', hue='year', s=60)
ax.set(xlabel='Podíl přijatých', ylabel=None, title='Podíl přijatých na pedagogické fakulty')
plt.legend(title='Rok')
plt.plot(foo21.prijat, foo21.vs_nazev)
plt.plot(foo17.prijat, foo17.vs_nazev)
ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
fig.tight_layout()

figpath = savefig(fig, f'prijati_pedf_vs_17_vs_21')
doc.add_picture(figpath, width=CHART_WIDTH)

#endregion

doc.show()

# TO DO?
# - obdoba krajů pro typy střední školy
# - úspěšnost uchazečů
# - ???


#endregion

###  NÁVRHY OD DANA  ###
#region
# 1.	Muzeme fakulty rozdelit podle zakladnich oborovych skupin (mediciny, humanitni, prirodovedni, technicke, humanitni, spolecenske]
# 2.	Muzeme fakulty rozdelit podle miry prevysu poptavky na zaklade oboru (nejvyssi v pravech a medicine a psychologii, mensi v prirodnich vedach humanitnich a spolecenskych, nejmensi v technickych)
# 3.	Muzeme rozdelit stredni skoly podle aplikacnich ambici jejich uchazcu hlasit se na VS s velkym previsem
# 4.	Muzeme rozdelit stredni skoly podle uspesnosti jejich uchazecu na vysoce poptavanych VS
# 5.	Muzeme rozdelit stredni skoly podle podilu uchazecu, kteri se nedostali na zadnou skolu
# 6.	Muzeme rozdelit stredni skoly podle podilu holek mezi uchazeci
# 7.	Muzeme se podivat, kam se zapsali uchazeci o pedf pokud byli prijati zaroven na jinou skolu (podle previsu ci oborove skupiny)

# ad 1:
# - podívej se, co vlastně všechno máme za fakulty, jestli to zvládneme nějak snadno manuálně přiřadit
# - nešla by použít informace ze studijních programů? jejich začátek tuším byl obor?

# - 2021: studijní programy jsou pravděpodobně konstruovány pomocí ISCED-F-2013
# - 2017: existuje převodník studijních programů SP <--> ISCED 
# - a zdá se, že ISCED-F-2013 je konstruován hierarchicky, tedy umožňuje snadné sdružování do vyšších jednotek

# vytvoř něco jako ZAPS dataset na kompletních datech
# ok, isc17 a isc21

font_size()

figs17 = uchazec.plot_isced(isc17)
figs21 = uchazec.plot_isced(isc21)

Selector([Selector(figs21, title='Rok 2021'), Selector(figs17, title='Rok 2017')], title='Podíl přijatých a zapsaných podle let a oborů ISCED').show()

# ok, this is pretty simple and seems to work well
# next steps?
# -> novy cistsi soubor?

isc21['selektivita'] = isc21['prijat'] / isc21['total']

df21 = pd.merge(df21, isc21[['aki4', 'selektivita']])

df21.shape

rep21

df21['selektivita'].mean()
np.average(df21['selektivita'], weights=df21['w'])

df21.groupby('gender')[['selektivita', 'vypr_flag']].mean()

df21.groupby('pedf')[['selektivita']].mean()

np.sum(df21['pedf'])


np.mean(df21['ss_izo'].value_counts())

df21[df21.pedf].show()




# teď už bych měl mít komplet ISCED-F v datasetech
df21.show('df21')

i2_zaps = df21[['isced2', 'zaps']].value_counts().unstack()
i2_counts = df21['isced2'].value_counts()

i2_zaps = i2_zaps.sort_index()
i2_counts = i2_counts.sort_index()
i2_zaps['total'] = i2_counts

i2_zaps = df21[['isced2', 'zaps']].value_counts().unstack()
i2_zaps['total'] = i2_zaps.sum(axis=1)
i2_zaps.T / i2_zaps['total']
np.round(100 * i2_zaps / i2_zaps.sum(axis=1)[:,np.newaxis], 1)

zaps = df21[['isced2', 'zaps']].value_counts().unstack()
zaps['total'] = zaps.sum(axis=1)
zaps['prijati'] = zaps['Uchazeč se nezapsal'] + zaps['Uchazeč se zapsal']
zaps['prijati_pct'] = np.round(100 * zaps['prijati'] / zaps['total'], 1)
zaps['zapsani_pct'] = np.round(100 * zaps['Uchazeč se zapsal'] / zaps['prijati'], 1)
zaps.show('zaps21_2')

zaps = df17[['isced2', 'zaps']].value_counts().unstack()
zaps['total'] = zaps.sum(axis=1)
zaps['prijati'] = zaps['Uchazeč se nezapsal'] + zaps['Uchazeč se zapsal']
zaps['prijati_pct'] = np.round(100 * zaps['prijati'] / zaps['total'], 1)
zaps['zapsani_pct'] = np.round(100 * zaps['Uchazeč se zapsal'] / zaps['prijati'], 1)
zaps.show('zaps17_2')

zaps = df21[['isced3', 'aki3', 'isced2', 'aki2', 'zaps']].value_counts().unstack()
zaps['total'] = zaps.sum(axis=1)
zaps['prijati'] = zaps['Uchazeč se nezapsal'] + zaps['Uchazeč se zapsal']
zaps['prijati_pct'] = np.round(100 * zaps['prijati'] / zaps['total'], 1)
zaps['zapsani_pct'] = np.round(100 * zaps['Uchazeč se zapsal'] / zaps['prijati'], 1)
zaps = zaps.drop(columns=['Uchazeč se nezapsal', 'Uchazeč se zapsal', 'Uchazeč nebyl přijat'])
zaps.show('zaps21_3')

zaps = df17[['isced4', 'aki4', 'isced3', 'aki3', 'isced2', 'aki2', 'zaps']].value_counts().unstack()
zaps['total'] = zaps.sum(axis=1)
zaps['prijati'] = zaps['Uchazeč se nezapsal'] + zaps['Uchazeč se zapsal']
zaps['prijati_pct'] = np.round(100 * zaps['prijati'] / zaps['total'], 1)
zaps['zapsani_pct'] = np.round(100 * zaps['Uchazeč se zapsal'] / zaps['prijati'], 1)
zaps = zaps.drop(columns=['Uchazeč se nezapsal', 'Uchazeč se zapsal', 'Uchazeč nebyl přijat'])
#zaps.show('zaps17_4')
zaps = zaps.reset_index()

zaps = df21[['isced4', 'aki4', 'isced3', 'aki3', 'isced2', 'aki2', 'zaps']].value_counts().unstack()
zaps['total'] = zaps.sum(axis=1)
zaps['prijati'] = zaps['Uchazeč se nezapsal'] + zaps['Uchazeč se zapsal']
zaps['prijati_pct'] = np.round(100 * zaps['prijati'] / zaps['total'], 1)
zaps['zapsani_pct'] = np.round(100 * zaps['Uchazeč se zapsal'] / zaps['prijati'], 1)
zaps = zaps.drop(columns=['Uchazeč se nezapsal', 'Uchazeč se zapsal', 'Uchazeč nebyl přijat'])
#zaps.show('zaps21_4')
zaps = zaps.reset_index()

i2s = zaps['isced2'].drop_duplicates().sort_values().values
colors = sns.color_palette(n_colors=10)
colors = {i: c for i, c in zip(i2s, colors)}
norm = plt.Normalize(zaps['total'].min(), zaps['total'].max())

font_size()

len(isc21['isced2'].unique())

xmin, xmax = 0, 110
ymin, ymax = 32, 108

figs21 = []

fig, ax = plt.subplots(figsize=(16, 9))

for i2, idf in zaps.groupby('isced2'):
    sns.scatterplot(x='prijati_pct', y='zapsani_pct', data=idf, label=i2, color=colors[i2], size='total', sizes=(10, 800), size_norm=norm, legend=False, ec=None)
    # labels
    for _, row in idf.iterrows():
        plt.text(row['prijati_pct'], row['zapsani_pct'] + 0.2 + np.sqrt(row['total']) / 80, row['isced4'],
                 color=colors[i2], ha='center', va='bottom')

# xmin, xmax = ax.get_xlim()
# ymin, ymax = ax.get_ylim()
ax.set(xlim=(xmin, xmax), ylim=(ymin, ymax))

ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.yaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.set(xlabel='Podíl přijatých', ylabel='Podíl zapsaných', title='Podíl přijatých a zapsaných podle oborů, 2021')
plt.legend()

figs21.append(Chart(fig, title='Všechny obory'))


for i2, idf in zaps.groupby('isced2'):
    fig, ax = plt.subplots(figsize=(16, 9))
    sns.scatterplot(x='prijati_pct', y='zapsani_pct', data=idf, label=i2, color=colors[i2], size='total', sizes=(10, 800), size_norm=norm, legend=False, ec=None)
    # labels
    for _, row in idf.iterrows():
        plt.text(row['prijati_pct'], row['zapsani_pct'] + 0.2 + np.sqrt(row['total']) / 80, row['isced4'],
                 color=colors[i2], ha='center', va='bottom')

    ax.set(xlim=(xmin, xmax), ylim=(ymin, ymax))

    ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
    ax.yaxis.set_major_formatter(mpl.ticker.PercentFormatter())
    ax.set(xlabel='Podíl přijatých', ylabel='Podíl zapsaných', title=i2)
    plt.legend()

    figs21.append(Chart(fig, title=i2))

sel21 = Selector(figs21, title='Rok 2021')

zaps = df17[['isced4', 'aki4', 'isced3', 'aki3', 'isced2', 'aki2', 'zaps']].value_counts().unstack()
zaps['total'] = zaps.sum(axis=1)
zaps['prijati'] = zaps['Uchazeč se nezapsal'] + zaps['Uchazeč se zapsal']
zaps['prijati_pct'] = np.round(100 * zaps['prijati'] / zaps['total'], 1)
zaps['zapsani_pct'] = np.round(100 * zaps['Uchazeč se zapsal'] / zaps['prijati'], 1)
zaps = zaps.drop(columns=['Uchazeč se nezapsal', 'Uchazeč se zapsal', 'Uchazeč nebyl přijat'])
#zaps.show('zaps17_4')
zaps = zaps.reset_index()

figs17 = []

fig, ax = plt.subplots(figsize=(16, 9))

for i2, idf in zaps.groupby('isced2'):
    sns.scatterplot(x='prijati_pct', y='zapsani_pct', data=idf, label=i2, color=colors[i2], size='total', sizes=(10, 800), size_norm=norm, legend=False, ec=None)
    # labels
    for _, row in idf.iterrows():
        plt.text(row['prijati_pct'], row['zapsani_pct'] + 0.2 + np.sqrt(row['total']) / 80, row['isced4'],
                 color=colors[i2], ha='center', va='bottom')

ax.set(xlim=(xmin, xmax), ylim=(ymin, ymax))

ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.yaxis.set_major_formatter(mpl.ticker.PercentFormatter())
ax.set(xlabel='Podíl přijatých', ylabel='Podíl zapsaných', title='Podíl přijatých a zapsaných podle oborů, 2021')
plt.legend()

figs17.append(Chart(fig, title='Všechny obory'))


for i2, idf in zaps.groupby('isced2'):
    fig, ax = plt.subplots(figsize=(16, 9))
    sns.scatterplot(x='prijati_pct', y='zapsani_pct', data=idf, label=i2, color=colors[i2], size='total', sizes=(10, 800), size_norm=norm, legend=False, ec=None)
    # labels
    for _, row in idf.iterrows():
        plt.text(row['prijati_pct'], row['zapsani_pct'] + 0.2 + np.sqrt(row['total']) / 80, row['isced4'],
                 color=colors[i2], ha='center', va='bottom')

    ax.set(xlim=(xmin, xmax), ylim=(ymin, ymax))

    ax.xaxis.set_major_formatter(mpl.ticker.PercentFormatter())
    ax.yaxis.set_major_formatter(mpl.ticker.PercentFormatter())
    ax.set(xlabel='Podíl přijatých', ylabel='Podíl zapsaných', title=i2)
    plt.legend()

    figs17.append(Chart(fig, title=i2))

sel17 = Selector(figs17, title='Rok 2017')

Selector([sel21, sel17], title='Podíl přijatých a zapsaných podle let a oborů ISCED').show()

ax.show()

i2



# příprava kategorií
df21.head(n=1000).show('df21')
df17.head(n=1000).show('df17')
df21[['fak_nazev', 'vs_nazev']].drop_duplicates().reset_index(drop=True).show()

# akko, akvo, program, regpro?
# program
sys_root = 'D:\\' if sys.platform == 'win32' else '/mnt/d'
data_root = os.path.join(sys_root, 'projects', 'idea', 'data')

# převod SP <--> ISCED
sp_isced = pd.read_excel(f'{data_root}/uchazec/Prevodnik_ISCED.xlsx')
sp_isced['isced'] = sp_isced['isced'].astype('str').str.pad(width=4, side='left', fillchar='0')
sp_isced.show('sp_isced', num_rows=5000)
sp_isced_dict = sp_isced.set_index('kód SP')['isced'].to_dict()

# ISCED-F
aki = {}
aki_xml = {
    2: 'http://stistko.uiv.cz/katalog/textdata/C131443AKI2.xml',
    3: 'http://stistko.uiv.cz/katalog/textdata/C13149AKI3.xml',
    4: 'http://stistko.uiv.cz/katalog/textdata/C13156AKI4.xml'
}

for i in range(2, 5):
    foo = pd.read_xml(aki_xml[i], encoding='cp1250', xpath='./veta')
    foo['KOD'] = foo['KOD'].astype('str').str.pad(width=i, side='left', fillchar='0')
    foo = foo[['KOD', 'TXT']].rename(columns={'KOD': f'aki{i}', 'TXT': f'isced{i}'})
    aki[i] = foo

df21['aki4'] = np.where(df21['program'].str.len() == 5, df21['program'].map(sp_isced_dict), df21['program'].str[1:5])
df21['aki3'] = df21['aki4'].str[:3]
df21['aki2'] = df21['aki4'].str[:2]

df21 = pd.merge(df21, aki[4], how='left')
df21 = pd.merge(df21, aki[3], how='left')
df21 = pd.merge(df21, aki[2], how='left')

df17['aki4'] = np.where(df17['program'].str.len() == 5, df17['program'].map(sp_isced_dict), df17['program'].str[1:5])
df17['aki3'] = df17['aki4'].str[:3]
df17['aki2'] = df17['aki4'].str[:2]

df17 = pd.merge(df17, aki[4], how='left')
df17 = pd.merge(df17, aki[3], how='left')
df17 = pd.merge(df17, aki[2], how='left')


aki[2].show('aki2')
aki[3].show('aki3')
aki[4].show('aki4')

df21[['vs_nazev', 'fak_nazev', 'program', 'aki2', 'aki3', 'aki4', 'isced2', 'isced3', 'isced4']].head(1000).show('df21_isced')

df17[['vs_nazev', 'fak_nazev', 'program', 'aki2', 'aki3', 'aki4', 'isced2', 'isced3', 'isced4']].head(1000).show('df17_isced')

df21 = df21.drop(columns=['aki2', 'aki3', 'aki4', 'isced2', 'isced3', 'isced4'])

aki2 = pd.read_xml(aki2_xml, encoding='cp1250', xpath='./veta')
aki3 = pd.read_xml(aki3_xml, encoding='cp1250', xpath='./veta')
aki4 = pd.read_xml(aki4_xml, encoding='cp1250', xpath='./veta')
aki2['KOD'] = aki2['KOD'].astype('str').str.pad(width=2, side='left', fillchar='0')
aki3['KOD'] = aki3['KOD'].astype('str').str.pad(width=3, side='left', fillchar='0')
aki4['KOD'] = aki4['KOD'].astype('str').str.pad(width=4, side='left', fillchar='0')

aki2.show('aki2')
aki3.show('aki3')
aki4.show('aki4')

df21['aki4'] = df21['program'].str[1:5]

df17['program'].str.len().value_counts()
df17.shape


np.sum(sp_isced['kód SP'].str.len() != 5)

program = pd.read_csv(f'{data_root}/uchazec/uch3-03-6/cisel/program.csv', encoding='cp1250')
program.show('program')

akko_xml = 'http://stistko.uiv.cz/katalog/textdata/C95743AKKO.xml'
akko = pd.read_xml(akko_xml, encoding='cp1250', xpath='./veta')
akko.show('akko')

akko[akko['KOD'].str.len() == 1].show('akko1')
akko[akko['KOD'].str.len() == 2].show('akko2')


akvo_xml = 'http://stistko.uiv.cz/katalog/textdata/C95619AKVO.xml'
akvo = pd.read_xml(akvo_xml, encoding='cp1250', xpath='./veta')
akvo.show('akvo')

regpro = pd.read_csv(f'{data_root}/uchazec/uch3-03-6/cisel/regpro.csv', encoding='cp1250')
regpro.show('regpro')

df17['obor1']


akko1 = akko[akko['KOD'].str.len() == 1][['KOD', 'TXT']].rename(columns={'KOD': 'obor11', 'TXT': 'akko1'})
akko1['akko1'] = akko1['akko1'].str.lower()

foo = df17[['vs_nazev', 'fak_nazev', 'obor1']].value_counts().rename('prihlasky').reset_index()
foo['obor11'] = foo['obor1'].str[0]
foo = pd.merge(foo, akko1)
foo['pocet'] = 1
foo.head()

foo
np.sum(foo.groupby(['vs_nazev', 'fak_nazev', 'akko1'])['pocet'].sum() > 0)

foo.groupby(['vs_nazev', 'fak_nazev', 'akko1']).sum()
bar = foo.groupby(['vs_nazev', 'fak_nazev', 'akko1'])[['prihlasky', 'pocet']].sum().reset_index()
bar = bar[bar['pocet'] > 0].copy()
bar['minus_pocet'] = -bar['pocet']
bar['minus_prihlasky'] = -bar['prihlasky']
bar.sort_values(['vs_nazev', 'fak_nazev', 'minus_pocet', 'minus_prihlasky']).show('bar')



#endregion

###  SOME DEBUGGING  ###

# můžu nahrát ten kompletní číselník?
# http://stistko.uiv.cz/katalog/ciselnika.asp -> MCPR
url_old = 'http://stistko.uiv.cz/katalog/ciselnik11x.asp?idc=MCPR&ciselnik=V%FDsledek+p%F8ij%EDmac%EDho+%F8%EDzen%ED&aap=on&poznamka='
html = urlopen(url_old).read()
cleanr = re.compile('<.*?>')
lines = [l.strip() for l in re.sub(cleanr, '', html.decode('windows-1250')).split('\r\n') if l.count(';') > 4]
text_data = StringIO('\n'.join(lines))
rg = pd.read_csv(text_data, sep=';', index_col=False)
rg.columns = [c.upper() for c in rg.columns]
rg = rg.groupby('KOD').last().reset_index()


# proč je tak často chybějící výsledek přijímacího řízení v datech z roku 2017
# ok, asi změna v kódování v číselníku

year = 17
sys_root = 'D:\\' if sys.platform == 'win32' else '/mnt/d'
data_root = os.path.join(sys_root, 'projects', 'idea', 'data')
path = os.path.join(data_root, 'uchazec', f'0022MUCH{year}P')

foo = pd.read_csv(f'{path}.csv', encoding='cp1250', low_memory=False)
foo['VYPR']
foo.head(1000).show()

np.round(100 * foo['VYPR'].value_counts() / foo.shape[0], 2)

np.round(100 * foo['ZAPS'].value_counts() / foo.shape[0], 2)

url = 'http://stistko.uiv.cz/katalog/textdata/C21925MCPR.xml'

rg = pd.read_xml(url, encoding='cp1250', xpath='./veta')
rg

rg = pd.merge(rg, pd.DataFrame({'KOD': range(10, 20)}), how='outer')
rg['TXT'] = rg['TXT'].fillna('Přijat - pravděpodobně, chybí v číselníku')
rg = pd.merge(rg, pd.DataFrame({'KOD': range(20, 30)}), how='outer')
rg['TXT'] = rg['TXT'].fillna('Nepřijat - pravděpodobně, chybí v číselníku')
rg = pd.merge(rg, pd.DataFrame({'KOD': range(30, 40)}), how='outer')
rg['TXT'] = rg['TXT'].fillna('Přijímací řízení zrušeno - pravděpodobně, chybí v číselníku')


rg['IDX'] = rg.index
df = pd.merge(df.rename(columns={c: 'KOD'}), rg[['KOD', 'IDX']].rename(columns={'IDX': c}), 
                how='left').drop(columns=['KOD'])        
rg_dict = rg[['IDX', 'TXT']].set_index('IDX')['TXT'].to_dict()
value_labels[c] = rg_dict




###  OLD AND UNUSED MESS  ###
#region
foo21 = pedf_app(ff21, 'ss_typ_g')
foo21


foo21 = pedf_app(ff21, 'ss_kraj')
ff21.columns


ff21['ss_typ'].value_counts()
ff21[['ss_typ', 'rmat']].value_counts()

ff21[ff21['ss_typ'] == 'Gymnázium']['ss_gym_delka'].value_counts()

foo21


foo17 = pedf_app(ff17, 'ss_kraj')[['ss_kraj', 'pedf_rel']].rename(columns={'pedf_rel': 2017})
foo21 = pedf_app(ff21, 'ss_kraj')[['ss_kraj', 'pedf_rel']].rename(columns={'pedf_rel': 2021})
foo = pd.merge(foo21, foo17)

fig, ax = plt.subplots()
sns.barplot(foo)
ax.show()
foo.plot().show()

doc.add_paragraph('')

fig, ax = plt.subplots()
ax = sns.histplot(data=ff, y='prihl', discrete=True, binwidth=0.8, shrink=0.8, ec=None)
ax.set(xlabel='Četnost', ylabel='Počet přihlášek', title='Počet přihlášek jednotlivých uchazečů')
ax.set_ylim(0, 13)
ax.invert_yaxis()
fig.tight_layout()

figpath = savefig(fig, f'cetnost_prihlasek_{y}')

doc.add_picture(figpath, width=CHART_WIDTH)
doc.add_paragraph('')



doc.show()

latent_styles = doc.styles.latent_styles
len(latent_styles)

table_style = [ls for ls in latent_styles if ls.name == 'Grid Table Light'][0]
doc.styles.add_style('Table Grid Light', docx.enum.style.WD_STYLE_TYPE.TABLE)

for ls in latent_styles:
    print(ls.name)

table = doc.add_table(rows=1, cols=3, style='Plain Table 1')
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Kraj'
hdr_cells[1].text = 'Přihlášek na PedF'
hdr_cells[2].text = 'Podíl na PedF'

for _, row in foo.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row['ss_kraj']
    row_cells[1].text = f'{row.prihl_bool:,}'
    row_cells[2].text = f'{row.pedf_rel} %'

doc.show()


doc.add_heading('Úspěšnost uchazečů o pedagogické fakulty')

for y in [21, 17]:
    ff = eval(f'ff{y}')

    tot_uchazecu = ff['prihl_pedf'].shape[0]
    tot_chlapcu = (ff['gender'] == 'Chlapec').sum()
    tot_divek = (ff['gender'] == 'Dívka').sum()


    doc.add_heading(f'20{y}:', 2)



pedf_app(ff21, 'gender')

pedf_app(ff21, 'ss_kraj')


doc.show()

df21[df21['pedf']][['fak_nazev', 'vs_nazev']].drop_duplicates().sort_values('vs_nazev')


sns.histplot(df21['pocet'], discrete=True, binwidth=0.8, ec=None).show()

doc.add_heading('2017:', 2)
doc.add_paragraph('')



doc.show()



os.getcwd()

font_size_word()
sns.histplot(data=ff21, y='prihl', hue='gender', multiple='stack', discrete=True, binwidth=0.8, ec=None).show()

ax = sns.histplot(data=ff21, y='prihl', discrete=True, binwidth=0.8, shrink=0.8, ec=None)
ax.set(xlabel='Četnost', ylabel='Počet přihlášek', title='Počet přihlášek jednotlivých uchazečů')
ax.set_ylim(0, 13)
ax.invert_yaxis()
ax.get_figure().tight_layout()

figpath = savefig(ax, 'cetnost_prihlasek_21')

doc.add_picture(figpath, width=Mm(140))
doc.show()

ax.get_figure().savefig

ax.get_figure()

ax.show()

ff21['prihl'].value_counts()

plt.rcParams['figure.titlesize'] = 30
ax.show()

eval(f'df{17}')

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc


p = doc.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

doc.add_heading('Heading, level 1', level=1)

doc.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
doc.add_paragraph(
    'first item in ordered list', style='List Number'
)

for style in doc.styles:
    print("style.name == %s" % style.name)






















### Learning how to use docx...


df.show()

document = Document()
document.add_heading('Document title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save(f'{foo}/demo.docx')

foo = create_stamped_temp('docx')
foo

os.startfile(f'{foo}/demo.docx')

# how to set a style?


doc = Document()
styles = doc.styles

style = styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(14)

doc.add_heading('Document title', 0)

doc.show()
docx_document_show(doc)
doc.save(tempfile.gettempdir() + '/foo.docx')

p = doc.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

doc.add_heading('Heading, level 1', level=1)

doc.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
doc.add_paragraph(
    'first item in ordered list', style='List Number'
)



doc.show()

foo = create_stamped_temp('docx')
doc.save(f'{foo}/demo.docx')
os.startfile(f'{foo}/demo.docx')

doc = Document('../IDEA_SABLONA.docx')
doc.show()

type(doc)

doc = Document()


doc = Document('empty.docx')
doc._body.clear_content()


doc.add_heading('Document title', 0)

p = doc.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

doc.add_heading('Heading, level 1', level=1)

doc.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
doc.add_paragraph(
    'first item in ordered list', style='List Number'
)

stamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
path = tempfile.gettempdir() + '/docs'
os.makedirs(path, exist_ok=True)
path = '{}/{}.{}'.format(path, stamp, format)

path

self.save(path, **kwargs)

doc.element.xml

doc.show()
doc

docx_document_show(doc)

os.environ['TEMP'] = 'D:/temp'

tempfile.tempdir = 'D:/temp'
tempfile.gettempdir()

Document.show = docx_document_show
doc.show()

Document
Document()

#endregion
